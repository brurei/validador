import filecmp
from functools import wraps
from typing import Dict, List, Tuple, Optional

from docx import Document
from flask import Flask, request, redirect, flash, send_from_directory, session, url_for, render_template_string
import os
import hashlib
import difflib
import time
import urllib.parse
import zipfile
import shutil
import json
from werkzeug.utils import secure_filename
import csv
import pandas as pd
from io import BytesIO
from flask import send_file
import fnmatch
import datetime

app = Flask(__name__)
app.secret_key = "comparison_secret_key"  # Necessário para flash messages

# Configuração de upload
UPLOAD_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'uploads')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
# Aumente este valor para permitir arquivos maiores

app.config['MAX_CONTENT_LENGTH'] = 3 * 1024 * 1024 * 1024  # 3GB

# Configurações de usuário (usuário fixo inicialmente)
# Em um ambiente de produção, você deve usar um sistema mais seguro de armazenamento de senhas
USERS = {
    "admin": "8c6976e5b5410415bde908bd4dee15dfb167a9c873fc4bb8a81f6f2ab448a918"  # senha: admin
}


# Função para verificar se o usuário está logado
def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'logged_in' not in session or not session['logged_in']:
            return redirect(url_for('login', next=request.url))
        return f(*args, **kwargs)

    return decorated_function


# Página de login
@app.route('/login', methods=['GET', 'POST'])
def login():
    error = None
    next_url = request.args.get('next', '/')

    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']

        # Calcular o hash SHA-256 da senha fornecida
        password_hash = hashlib.sha256(password.encode()).hexdigest()

        if username in USERS and USERS[username] == password_hash:
            session['logged_in'] = True
            session['username'] = username
            if next_url == '/':
                return redirect('/info')
            else:
                return redirect(next_url)
        else:
            error = 'Credenciais inválidas. Por favor, tente novamente.'

    # Template HTML para a página de login
    login_template = """<!DOCTYPE html>
    <html lang="pt-br">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Login - Sistema de Validação de Arquivos</title>
        <style>
            :root {
                --primary-color: #4CAF50;
                --primary-hover: #43a047;
                --light-gray: #f4f4f4;
                --medium-gray: #e0e0e0;
                --dark-gray: #757575;
                --border-radius: 8px;
                --box-shadow: 0 4px 12px rgba(0, 0, 0, 0.08);
                --transition-speed: 0.3s;
            }

            * {
                box-sizing: border-box;
                margin: 0;
                padding: 0;
            }

            body {
                font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                background-color: #f9f9f9;
                color: #333;
                line-height: 1.6;
                padding: 20px;
                height: 100vh;
                display: flex;
                justify-content: center;
                align-items: center;
            }

            .login-container {
                max-width: 400px;
                width: 100%;
                background: white;
                border-radius: var(--border-radius);
                box-shadow: var(--box-shadow);
                overflow: hidden;
            }

            .login-header {
                background-color: var(--primary-color);
                color: white;
                padding: 20px;
                text-align: center;
            }

            .login-header h1 {
                font-size: 24px;
                font-weight: 600;
            }

            .login-form {
                padding: 30px;
            }

            .form-group {
                margin-bottom: 20px;
            }

            .form-label {
                display: block;
                margin-bottom: 8px;
                font-weight: 500;
                color: #333;
            }

            .form-input {
                width: 100%;
                padding: 12px;
                border: 1px solid var(--medium-gray);
                border-radius: var(--border-radius);
                font-size: 16px;
                transition: border-color var(--transition-speed);
            }

            .form-input:focus {
                outline: none;
                border-color: var(--primary-color);
                box-shadow: 0 0 0 2px rgba(76, 175, 80, 0.2);
            }

            .login-btn {
                display: block;
                width: 100%;
                padding: 14px;
                background-color: var(--primary-color);
                color: white;
                border: none;
                border-radius: var(--border-radius);
                font-size: 16px;
                font-weight: 600;
                cursor: pointer;
                transition: background-color var(--transition-speed);
            }

            .login-btn:hover {
                background-color: var(--primary-hover);
            }

            .error-message {
                background-color: #f8d7da;
                color: #721c24;
                padding: 10px;
                border-radius: var(--border-radius);
                margin-bottom: 20px;
                text-align: center;
            }

            .login-footer {
                text-align: center;
                padding: 15px;
                border-top: 1px solid var(--medium-gray);
                color: var(--dark-gray);
                font-size: 14px;
            }
        </style>
    </head>
    <body>
        <div class="login-container">
            <div class="login-header">
                <h1>Sistema de Validação de Arquivos</h1>
            </div>

            <div class="login-form">
                {% if error %}
                <div class="error-message">
                    {{ error }}
                </div>
                {% endif %}

                <form action="/login" method="post">
                    <input type="hidden" name="next" value="{{ next_url }}">

                    <div class="form-group">
                        <label for="username" class="form-label">Usuário</label>
                        <input type="text" id="username" name="username" class="form-input" required autofocus>
                    </div>

                    <div class="form-group">
                        <label for="password" class="form-label">Senha</label>
                        <input type="password" id="password" name="password" class="form-input" required>
                    </div>

                    <button type="submit" class="login-btn">Entrar</button>
                </form>
            </div>

            <div class="login-footer">
                &copy; 2025 Sistema de Validação de Arquivos
            </div>
        </div>
    </body>
    </html>"""

    return render_template_string(login_template, error=error, next_url=next_url)


# Rota para logout
@app.route('/logout')
def logout():
    session.pop('logged_in', None)
    session.pop('username', None)
    return redirect(url_for('login'))


def apply_filters(file_list, base_dir, filters):
    """
    Aplica filtros avançados a uma lista de arquivos.

    Args:
        file_list: Lista de nomes de arquivos
        base_dir: Diretório base para os arquivos
        filters: Dicionário com os filtros a serem aplicados

    Returns:
        Lista filtrada de nomes de arquivos
    """
    filtered_files = []

    for filename in file_list:
        file_path = os.path.join(base_dir, filename)

        # Pular se o arquivo não existe mais
        if not os.path.exists(file_path):
            continue

        # Filtro de extensão
        if filters.get('file_extensions'):
            extensions = [ext.strip().lower() for ext in filters['file_extensions'].split(',')]
            file_ext = os.path.splitext(filename)[1].lower().replace('.', '')
            if extensions and file_ext not in extensions:
                continue

        # Filtros de tamanho
        file_size_kb = os.path.getsize(file_path) / 1024  # Tamanho em KB

        if filters.get('size_min') and file_size_kb < float(filters['size_min']):
            continue

        if filters.get('size_max') and file_size_kb > float(filters['size_max']):
            continue

        # Filtros de data de modificação
        modified_time = os.path.getmtime(file_path)
        modified_date = datetime.datetime.fromtimestamp(modified_time)

        if filters.get('modified_after'):
            after_date = datetime.datetime.strptime(filters['modified_after'], '%Y-%m-%d')
            if modified_date.date() < after_date.date():
                continue

        if filters.get('modified_before'):
            before_date = datetime.datetime.strptime(filters['modified_before'], '%Y-%m-%d')
            if modified_date.date() > before_date.date():
                continue

        # Filtro de padrão de nome
        if filters.get('name_pattern'):
            if not fnmatch.fnmatch(filename, filters['name_pattern']):
                continue

        # Filtro de exclusão
        if filters.get('exclude_pattern'):
            exclude_patterns = [p.strip() for p in filters['exclude_pattern'].split(',')]
            skip = False
            for pattern in exclude_patterns:
                if fnmatch.fnmatch(filename, pattern):
                    skip = True
                    break
            if skip:
                continue

        # Se passou por todos os filtros, adiciona à lista filtrada
        filtered_files.append(filename)

    return filtered_files


# Modificação na função de comparação em lote para suportar filtros
def compare_file_batches_with_filters(source1_dir, source2_dir, method, filters=None):
    """
    Versão melhorada da função compare_file_batches que suporta filtros avançados
    """
    # Obter lista de arquivos em cada diretório
    files1 = [f for f in os.listdir(source1_dir) if os.path.isfile(os.path.join(source1_dir, f))]
    files2 = [f for f in os.listdir(source2_dir) if os.path.isfile(os.path.join(source2_dir, f))]

    # Aplicar filtros se fornecidos
    if filters:
        files1 = apply_filters(files1, source1_dir, filters)
        files2 = apply_filters(files2, source2_dir, filters)

    results = []

    # Encontrar arquivos comuns entre as duas fontes
    common_files = set(files1).intersection(set(files2))

    # Arquivos exclusivos de cada fonte
    only_in_source1 = set(files1) - set(files2)
    only_in_source2 = set(files2) - set(files1)

    # Comparar arquivos comuns
    for filename in common_files:
        file1_path = os.path.join(source1_dir, filename)
        file2_path = os.path.join(source2_dir, filename)

        result = {
            'filename': filename,
            'exists_in_source1': True,
            'exists_in_source2': True,
            'comparison_result': None,
            'details': None,
            'file_info': {
                'size1': os.path.getsize(file1_path),
                'size2': os.path.getsize(file2_path),
                'modified1': datetime.datetime.fromtimestamp(os.path.getmtime(file1_path)).strftime(
                    '%Y-%m-%d %H:%M:%S'),
                'modified2': datetime.datetime.fromtimestamp(os.path.getmtime(file2_path)).strftime('%Y-%m-%d %H:%M:%S')
            }
        }

        if method == 'hash':
            identical, message = compare_files_hash(file1_path, file2_path)
            result['comparison_result'] = 'identical' if identical else 'different'
            result['details'] = message

        elif method == 'content':
            differences, message = compare_files_content(file1_path, file2_path)
            result['comparison_result'] = 'identical' if not differences else 'different'
            result['details'] = differences if differences else message

            # Processar diferenças para visualizador avançado
            if differences:
                result['processed_diff'] = process_differences_for_advanced_view(file1_path, file2_path, differences)

        elif method == 'binary':
            differences, message = compare_files_binary(file1_path, file2_path)
            result['comparison_result'] = 'identical' if not differences else 'different'
            result['details'] = differences if differences else message

        results.append(result)

    # Adicionar arquivos que existem apenas em uma fonte
    for filename in only_in_source1:
        file_path = os.path.join(source1_dir, filename)
        results.append({
            'filename': filename,
            'exists_in_source1': True,
            'exists_in_source2': False,
            'comparison_result': 'only_in_source1',
            'details': 'Arquivo existe apenas na Fonte 1',
            'file_info': {
                'size1': os.path.getsize(file_path),
                'size2': 0,
                'modified1': datetime.datetime.fromtimestamp(os.path.getmtime(file_path)).strftime('%Y-%m-%d %H:%M:%S'),
                'modified2': 'N/A'
            }
        })

    for filename in only_in_source2:
        file_path = os.path.join(source2_dir, filename)
        results.append({
            'filename': filename,
            'exists_in_source1': False,
            'exists_in_source2': True,
            'comparison_result': 'only_in_source2',
            'details': 'Arquivo existe apenas na Fonte 2',
            'file_info': {
                'size1': 0,
                'size2': os.path.getsize(file_path),
                'modified1': 'N/A',
                'modified2': datetime.datetime.fromtimestamp(os.path.getmtime(file_path)).strftime('%Y-%m-%d %H:%M:%S')
            }
        })

    return results


import re


def process_differences_for_advanced_view(file1_path, file2_path, differences):
    """
    Processa as diferenças para o visualizador avançado lado a lado

    Args:
        file1_path: Caminho do primeiro arquivo
        file2_path: Caminho do segundo arquivo
        differences: Lista de diferenças no formato unificado (unified diff)

    Returns:
        Lista processada de diferenças para visualização lado a lado
    """
    try:
        with open(file1_path, 'r', encoding='utf-8') as f1, open(file2_path, 'r', encoding='utf-8') as f2:
            file1_lines = f1.readlines()
            file2_lines = f2.readlines()
    except UnicodeDecodeError:
        # Retorna lista vazia se não for possível ler como texto
        return []

    processed_diff = []

    # Divide as diferenças em chunks baseados nos headers @@
    chunks = []
    current_chunk = None

    for line in differences:
        if line.startswith('@@'):
            if current_chunk is not None:
                chunks.append(current_chunk)
            current_chunk = {'header': line, 'lines': []}
        elif current_chunk is not None:
            current_chunk['lines'].append(line)

    if current_chunk is not None:
        chunks.append(current_chunk)

    # Processa cada chunk separadamente
    for chunk in chunks:
        # Extrai os números de linha do header @@ -a,b +c,d @@
        header = chunk['header']
        match = re.search(r'@@ -(\d+),(\d+) \+(\d+),(\d+) @@', header)

        if not match:
            continue

        file1_start = int(match.group(1))
        file1_count = int(match.group(2))
        file2_start = int(match.group(3))
        file2_count = int(match.group(4))

        # Rastreia as posições atuais nas linhas dos arquivos
        file1_idx = file1_start - 1  # Índice baseado em 0 para acessar a lista
        file2_idx = file2_start - 1

        # Processa as linhas do chunk
        for line in chunk['lines']:
            if line.startswith('-'):
                # Linha removida (presente apenas no arquivo 1)
                content = line[1:].rstrip('\n')
                processed_diff.append({
                    'type': 'removed',
                    'content': content,
                    'lineNum1': file1_idx + 1,  # +1 para mostrar número baseado em 1
                    'lineNum2': None,
                    'paired': False
                })
                file1_idx += 1
            elif line.startswith('+'):
                # Linha adicionada (presente apenas no arquivo 2)
                content = line[1:].rstrip('\n')
                processed_diff.append({
                    'type': 'added',
                    'content': content,
                    'lineNum1': None,
                    'lineNum2': file2_idx + 1,
                    'paired': False
                })
                file2_idx += 1
            elif line.startswith(' '):
                # Linha de contexto (presente em ambos os arquivos)
                content = line[1:].rstrip('\n')
                processed_diff.append({
                    'type': 'unchanged',
                    'content': content,
                    'lineNum1': file1_idx + 1,
                    'lineNum2': file2_idx + 1
                })
                file1_idx += 1
                file2_idx += 1

    # Tenta agrupar remoções e adições que são modificações da mesma linha
    i = 0
    while i < len(processed_diff) - 1:
        current = processed_diff[i]
        next_item = processed_diff[i + 1]

        # Se uma linha removida é seguida por uma linha adicionada, marque-as como um par
        if current['type'] == 'removed' and next_item['type'] == 'added':
            current['paired'] = True
            next_item['paired'] = True
            next_item['lineNum1'] = current['lineNum1']  # Para alinhamento visual
            i += 2  # Pule ambos os itens
        else:
            i += 1

    return processed_diff


def load_json_file(file_path):
    import json
    with open(file_path, 'r', encoding='utf-8') as f:
        return json.load(f)


# Função para gerar relatório em CSV
def generate_csv_report(comparison_results):
    output = BytesIO()
    fieldnames = ['filename', 'status', 'details', 'exists_in_source1', 'exists_in_source2']

    writer = csv.DictWriter(output, fieldnames=fieldnames)
    writer.writeheader()

    for result in comparison_results:
        status = ''
        details = ''

        if result['comparison_result'] == 'identical':
            status = 'Idêntico'
            details = 'Os arquivos são idênticos'
        elif result['comparison_result'] == 'different':
            status = 'Diferente'
            details = 'Os arquivos são diferentes'
        elif result['comparison_result'] == 'only_in_source1':
            status = 'Somente na Fonte 1'
            details = 'Arquivo existe apenas na Fonte 1'
        elif result['comparison_result'] == 'only_in_source2':
            status = 'Somente na Fonte 2'
            details = 'Arquivo existe apenas na Fonte 2'

        writer.writerow({
            'filename': result['filename'],
            'status': status,
            'details': details,
            'exists_in_source1': 'Sim' if result['exists_in_source1'] else 'Não',
            'exists_in_source2': 'Sim' if result['exists_in_source2'] else 'Não'
        })

    output.seek(0)
    return output


# Função para gerar relatório em Excel
def generate_excel_report(comparison_results):
    # Preparar dados para o DataFrame
    data = []
    for result in comparison_results:
        status = ''
        details = ''

        if result['comparison_result'] == 'identical':
            status = 'Idêntico'
            details = 'Os arquivos são idênticos'
        elif result['comparison_result'] == 'different':
            status = 'Diferente'
            details = 'Os arquivos são diferentes'
        elif result['comparison_result'] == 'only_in_source1':
            status = 'Somente na Fonte 1'
            details = 'Arquivo existe apenas na Fonte 1'
        elif result['comparison_result'] == 'only_in_source2':
            status = 'Somente na Fonte 2'
            details = 'Arquivo existe apenas na Fonte 2'

        data.append({
            'Arquivo': result['filename'],
            'Status': status,
            'Detalhes': details,
            'Presente na Fonte 1': 'Sim' if result['exists_in_source1'] else 'Não',
            'Presente na Fonte 2': 'Sim' if result['exists_in_source2'] else 'Não'
        })

    # Criar DataFrame e exportar para Excel
    df = pd.DataFrame(data)
    output = BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, sheet_name='Relatório de Validação', index=False)

        # Formatação básica
        workbook = writer.book
        worksheet = writer.sheets['Relatório de Validação']

        # Formatar cabeçalho
        header_format = workbook.add_format({
            'bold': True,
            'text_wrap': True,
            'valign': 'top',
            'fg_color': '#D7E4BC',
            'border': 1
        })

        # Aplicar formatação ao cabeçalho
        for col_num, value in enumerate(df.columns.values):
            worksheet.write(0, col_num, value, header_format)

        # Ajustar larguras das colunas
        worksheet.set_column('A:A', 30)  # Arquivo
        worksheet.set_column('B:B', 15)  # Status
        worksheet.set_column('C:C', 40)  # Detalhes
        worksheet.set_column('D:E', 20)  # Presente nas fontes

    output.seek(0)
    return output


# Adicionar rotas para exportação
@app.route('/export-csv/<timestamp>', methods=['GET'])
def export_csv(timestamp):
    """Exporta os resultados da comparação para CSV."""
    try:
        # Recuperar resultados da comparação armazenados temporariamente
        comparison_results = session.get(f'comparison_results_{timestamp}', [])
        if not comparison_results:
            return redirect(f"/?message={urllib.parse.quote('Erro: Dados da comparação não encontrados')}")

        # Gerar o arquivo CSV
        output = generate_csv_report(comparison_results)

        # Enviar o arquivo para download
        return send_file(
            output,
            as_attachment=True,
            download_name=f'comparacao_{timestamp}.csv',
            mimetype='text/csv'
        )

    except Exception as e:
        return redirect(f"/?message={urllib.parse.quote(f'Erro ao exportar CSV: {str(e)}')}")


@app.route('/export-excel/<timestamp>', methods=['GET'])
def export_excel(timestamp):
    """Exporta os resultados da comparação para Excel."""
    try:
        # Recuperar resultados da comparação armazenados temporariamente
        comparison_results = session.get(f'comparison_results_{timestamp}', [])
        if not comparison_results:
            return redirect(f"/?message={urllib.parse.quote('Erro: Dados da comparação não encontrados')}")

        # Gerar o arquivo Excel
        output = generate_excel_report(comparison_results)

        # Enviar o arquivo para download
        return send_file(
            output,
            as_attachment=True,
            download_name=f'comparacao_{timestamp}.xlsx',
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )

    except Exception as e:
        return redirect(f"/?message={urllib.parse.quote(f'Erro ao exportar Excel: {str(e)}')}")


# Funções de comparação
def calculate_hash(file_path, buffer_size=8192):
    """Calcula o hash SHA-256 de um arquivo em blocos."""
    hash_sha256 = hashlib.sha256()
    with open(file_path, 'rb') as f:
        while True:
            data = f.read(buffer_size)
            if not data:
                break
            hash_sha256.update(data)
    return hash_sha256.hexdigest()


def compare_files_hash(file1_path, file2_path):
    """Compara dois arquivos usando hash SHA-256."""
    if not os.path.exists(file1_path):
        return False, f"Erro: O arquivo '{file1_path}' não existe."
    if not os.path.exists(file2_path):
        return False, f"Erro: O arquivo '{file2_path}' não existe."

    hash1 = calculate_hash(file1_path)
    hash2 = calculate_hash(file2_path)

    if hash1 == hash2:
        return True, "Os arquivos são idênticos."
    else:
        return False, "Os arquivos são diferentes."


# MODIFIQUE SUA FUNÇÃO compare_files_content (PARTE ESPECÍFICA):
def compare_files_content(file1_path, file2_path, context_lines=3):
    """Compara o conteúdo de dois arquivos linha por linha e retorna as diferenças."""
    if not os.path.exists(file1_path):
        return [], f"Erro: O arquivo '{file1_path}' não existe."
    if not os.path.exists(file2_path):
        return [], f"Erro: O arquivo '{file2_path}' não existe."

    # Verificação rápida usando filecmp
    if filecmp.cmp(file1_path, file2_path, shallow=False):
        return [], "Os arquivos são idênticos."

    # CORREÇÃO: Verificar se há um layout selecionado na sessão primeiro
    selected_layout = session.get('selected_layout')
    print("PRINTANDO LAYOUT sessão: " + str(selected_layout))

    if selected_layout:
        print("Usando layout selecionado manualmente da sessão")
        return compare_files_content_with_layout(file1_path, file2_path, selected_layout, context_lines)

    # Só fazer detecção automática se nenhum layout foi selecionado
    try:
        is_struct, detected_layout = is_structured_file(file1_path)
    except Exception as e:
        print(f"Erro ao verificar tipo de arquivo: {str(e)}")
        is_struct = False
        detected_layout = None

    print("Antes de validar o arquivo estruturado")

    # Se for arquivo estruturado, processar com o analisador específico
    if is_struct and detected_layout:
        print("Caiu no bloco de arquivo estruturado")

        try:
            # Verificar se é o layout legado
            if detected_layout.get('name') == 'Seca':
                # Usar o layout hardcoded original
                layout = extrair_dicionario_layout()
            else:
                # Usar o layout detectado automaticamente
                layout = detected_layout.get('fields', {})

            # Analisar diferenças estruturadas
            estrutura_diffs = analisar_arquivo_estruturado(file1_path, file2_path, layout)

            # Verificar se o resultado é válido
            if estrutura_diffs is None:
                print("Aviso: analisar_arquivo_estruturado retornou None, usando análise padrão.")
                return super_compare_files(file1_path, file2_path, context_lines)

            # Armazenar para uso na visualização de negócio
            if 'estrutura_diffs' not in session:
                session['estrutura_diffs'] = []
            session['estrutura_diffs'] = estrutura_diffs

            # Armazenar o nome do layout detectado
            session['layout_name'] = detected_layout.get('name', 'Layout Desconhecido')

            # Converter para o formato tradicional de diferenças
            traditional_diffs = []

            # Verificar se há erros na análise
            has_error = any(diff.get('tipo') == 'erro' for diff in estrutura_diffs)
            if has_error:
                error_diff = next(diff for diff in estrutura_diffs if diff.get('tipo') == 'erro')
                return [], f"Erro na análise estruturada: {error_diff.get('mensagem', 'Erro desconhecido')}"

            # Converter diferenças estruturadas para o formato tradicional
            for diff in estrutura_diffs:
                diff_type = diff.get('tipo', '')

                if diff_type == 'campos_alterados':
                    traditional_diffs.append(f"@@ Linha {diff['linha']} - Alterações de campos @@")
                    for campo in diff.get('diferenca', []):
                        traditional_diffs.append(f"- Campo {campo['campo']}: {campo['valor_antigo']}")
                        traditional_diffs.append(f"+ Campo {campo['campo']}: {campo['valor_novo']}")

                elif diff_type == 'linha_alterada':
                    traditional_diffs.append(f"@@ Linha {diff['linha']} @@")
                    traditional_diffs.append(f"- {diff.get('valor_antigo', '')}")
                    traditional_diffs.append(f"+ {diff.get('valor_novo', '')}")

                elif diff_type == 'linha_removida':
                    traditional_diffs.append(f"@@ Linha {diff['linha']} @@")
                    traditional_diffs.append(f"- {diff.get('valor_antigo', '')}")

                elif diff_type == 'linha_adicionada':
                    traditional_diffs.append(f"@@ Linha {diff['linha']} @@")
                    traditional_diffs.append(f"+ {diff.get('valor_novo', '')}")

                elif diff_type == 'info':
                    traditional_diffs.append(f"@@ {diff.get('mensagem', '')} @@")

            return traditional_diffs, f"Análise de arquivo estruturado completa (Layout: {detected_layout.get('name', 'Desconhecido')})."

        except Exception as e:
            import traceback
            traceback.print_exc()
            print(f"Erro ao processar arquivo estruturado: {str(e)}")
            return super_compare_files(file1_path, file2_path, context_lines)

    # Se não for um arquivo estruturado ou a análise falhar, usar a função original
    return super_compare_files(file1_path, file2_path, context_lines)


# Esta função representa o comportamento original de compare_files_content
def super_compare_files(file1_path, file2_path, context_lines=3):
    """
    Versão original da função de comparação, para ser usada quando a análise estruturada falhar
    ou quando o arquivo não for do tipo estruturado.
    """
    try:
        # Ler os conteúdos dos arquivos como bytes
        with open(file1_path, 'rb') as f1, open(file2_path, 'rb') as f2:
            content1 = f1.read()
            content2 = f2.read()

        # Tentar decodificar como texto
        try:
            text1 = content1.decode('utf-8', errors='replace')
            text2 = content2.decode('utf-8', errors='replace')

            # Dividir em linhas
            lines1 = text1.splitlines()
            lines2 = text2.splitlines()

            # Gerar diferenças usando difflib
            diff = list(difflib.unified_diff(
                lines1, lines2,
                fromfile=os.path.basename(file1_path),
                tofile=os.path.basename(file2_path),
                lineterm='', n=context_lines
            ))

            # Se não houver diferenças significativas além dos cabeçalhos
            if len(diff) <= 2:
                differences = []
                # Verificar se o conteúdo é realmente diferente (pode ser apenas formatação de linha)
                if content1 != content2:
                    # Analisar byte a byte para encontrar a primeira diferença
                    differences = analyze_binary_difference(content1, content2)
                return differences, "Análise completa."

            # Retornar as diferenças encontradas
            return diff, "Análise completa."

        except UnicodeDecodeError:
            # Se falhar a decodificação de texto, realizar análise binária
            differences = analyze_binary_difference(content1, content2)
            return differences, "Análise binária."

    except Exception as e:
        return [f"Erro durante a comparação: {str(e)}"], "Falha na análise."


def analyze_binary_difference(content1, content2):
    """Analisa diferenças entre dois conteúdos binários."""
    differences = []

    # Tamanhos diferentes
    size1, size2 = len(content1), len(content2)
    if size1 != size2:
        differences.append(f"@@ Tamanhos diferentes: {size1} vs {size2} bytes @@")

    # Encontrar a primeira posição diferente
    min_size = min(size1, size2)
    different_pos = -1

    for i in range(min_size):
        if content1[i] != content2[i]:
            different_pos = i
            break

    if different_pos >= 0:
        # Mostrar contexto em torno da primeira diferença
        start = max(0, different_pos - 10)
        end_1 = min(size1, different_pos + 30)
        end_2 = min(size2, different_pos + 30)

        differences.append(f"@@ Primeiros caracteres dos arquivos @@")

        # Converter para texto para exibição amigável, com fallback para hex
        try:
            prefix1 = content1[start:different_pos].decode('utf-8', errors='replace')
            change1 = content1[different_pos:end_1].decode('utf-8', errors='replace')
            differences.append(f"-{prefix1}{change1}")
        except:
            # Em caso de erro, mostrar em hex
            differences.append(f"-Hex: {content1[start:end_1].hex()}")

        try:
            prefix2 = content2[start:different_pos].decode('utf-8', errors='replace')
            change2 = content2[different_pos:end_2].decode('utf-8', errors='replace')
            differences.append(f"+{prefix2}{change2}")
        except:
            # Em caso de erro, mostrar em hex
            differences.append(f"+Hex: {content2[start:end_2].hex()}")

    return differences


def compare_files_binary(file1_path, file2_path, max_differences=100, buffer_size=8192):
    """Compara arquivos binários byte a byte em blocos e retorna as posições das diferenças."""
    if not os.path.exists(file1_path):
        return [], f"Erro: O arquivo '{file1_path}' não existe."
    if not os.path.exists(file2_path):
        return [], f"Erro: O arquivo '{file2_path}' não existe."

    diff_positions = []
    position = 0
    diff_count = 0

    with open(file1_path, 'rb') as f1, open(file2_path, 'rb') as f2:
        while True:
            if diff_count >= max_differences:
                diff_positions.append(f"Limite de {max_differences} diferenças atingido. Análise interrompida.")
                break

            block1 = f1.read(buffer_size)
            block2 = f2.read(buffer_size)

            if not block1 and not block2:
                # Fim de ambos os arquivos
                break

            if not block1:
                # Arquivo 1 é menor
                diff_positions.append(f"Posição {position}: EOF no arquivo 1, mais {len(block2)} bytes no arquivo 2")
                diff_count += 1
                break

            if not block2:
                # Arquivo 2 é menor
                diff_positions.append(f"Posição {position}: mais {len(block1)} bytes no arquivo 1, EOF no arquivo 2")
                diff_count += 1
                break

            # Comparar os blocos byte a byte
            for i, (b1, b2) in enumerate(zip(block1, block2)):
                if b1 != b2:
                    if diff_count >= max_differences:
                        break
                    diff_positions.append(
                        f"Posição {position + i}: byte {b1} ({hex(b1)}) no arquivo 1, byte {b2} ({hex(b2)}) no arquivo 2")
                    diff_count += 1

            # Se os blocos têm tamanhos diferentes
            if len(block1) != len(block2):
                diff_positions.append(
                    f"Posição {position + min(len(block1), len(block2))}: Tamanhos de bloco diferentes")
                diff_count += 1

            position += len(block1)

    return diff_positions, "Análise completa."


def extract_zip(zip_path, extract_to):
    """Extrai um arquivo ZIP para um diretório específico."""
    with zipfile.ZipFile(zip_path, 'r') as zip_ref:
        zip_ref.extractall(extract_to)
    return [f for f in os.listdir(extract_to) if os.path.isfile(os.path.join(extract_to, f))]


def compare_file_batches(source1_dir, source2_dir, method):
    """Compara dois diretórios de arquivos, combinando-os pelo nome."""
    # Obter lista de arquivos em cada diretório
    files1 = [f for f in os.listdir(source1_dir) if os.path.isfile(os.path.join(source1_dir, f))]
    files2 = [f for f in os.listdir(source2_dir) if os.path.isfile(os.path.join(source2_dir, f))]

    results = []

    # Encontrar arquivos comuns entre as duas fontes
    common_files = set(files1).intersection(set(files2))

    # Arquivos exclusivos de cada fonte
    only_in_source1 = set(files1) - set(files2)
    only_in_source2 = set(files2) - set(files1)

    # Comparar arquivos comuns
    for filename in common_files:
        file1_path = os.path.join(source1_dir, filename)
        file2_path = os.path.join(source2_dir, filename)

        result = {
            'filename': filename,
            'exists_in_source1': True,
            'exists_in_source2': True,
            'comparison_result': None,
            'details': None
        }

        if method == 'hash':
            identical, message = compare_files_hash(file1_path, file2_path)
            result['comparison_result'] = 'identical' if identical else 'different'
            result['details'] = message

        elif method == 'content':
            differences, message = compare_files_content(file1_path, file2_path)
            result['comparison_result'] = 'identical' if not differences else 'different'
            result['details'] = differences if differences else message

        elif method == 'binary':
            differences, message = compare_files_binary(file1_path, file2_path)
            result['comparison_result'] = 'identical' if not differences else 'different'
            result['details'] = differences if differences else message

        results.append(result)

    # Adicionar arquivos que existem apenas em uma fonte
    for filename in only_in_source1:
        results.append({
            'filename': filename,
            'exists_in_source1': True,
            'exists_in_source2': False,
            'comparison_result': 'only_in_source1',
            'details': 'Arquivo existe apenas na Fonte 1'
        })

    for filename in only_in_source2:
        results.append({
            'filename': filename,
            'exists_in_source1': False,
            'exists_in_source2': True,
            'comparison_result': 'only_in_source2',
            'details': 'Arquivo existe apenas na Fonte 2'
        })

    return results


# SUBSTITUA SUA FUNÇÃO is_structured_file POR ESTA:
def is_structured_file(file_path):
    """
    Detecta se um arquivo é estruturado e qual layout usar.
    Retorna (is_structured: bool, layout_info: dict)
    """
    if not os.path.exists(file_path):
        return False, None

    # Tentar detectar layout automaticamente
    layout_name, layout_data = detect_layout_automatically(file_path)

    if layout_name and layout_data:
        print(f"Layout detectado automaticamente: {layout_name}")
        return True, layout_data

    # Fallback: verificar se é o layout legado SECA
    try:
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            first_line = f.readline().strip()

        # Verificar se é o formato SECA legado
        if first_line.startswith('00SERVICO CALCULADO'):
            return True, {'name': 'Seca', 'fields': extrair_dicionario_layout()}

    except Exception as e:
        print(f"Erro ao verificar arquivo: {e}")

    return False, None


def detect_layout_automatically(file_path: str) -> Tuple[Optional[str], Optional[Dict]]:
    """
    Detecta automaticamente o layout de um arquivo baseado em:
    1. Nome do arquivo
    2. Tamanho dos registros
    3. Padrões de conteúdo
    """
    if not os.path.exists(file_path):
        return None, None

    # Carregar layouts disponíveis
    layouts_cache = load_all_layouts()
    if not layouts_cache:
        print("Nenhum layout encontrado na pasta layouts/")
        return None, None

    # Extrair informações do arquivo
    filename = os.path.basename(file_path)

    # Ler primeiras linhas para análise
    try:
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            lines = [f.readline().rstrip('\n') for _ in range(10)]

        # Filtrar linhas vazias
        non_empty_lines = [line for line in lines if line.strip()]
        if not non_empty_lines:
            return None, None

    except Exception as e:
        print(f"Erro ao ler arquivo {file_path}: {e}")
        return None, None

    # 1. DETECÇÃO POR NOME DO ARQUIVO
    layout_by_filename = detect_by_filename(filename, layouts_cache)
    if layout_by_filename:
        layout_name, layout_data = layout_by_filename
        if validate_layout_with_content(layout_data, non_empty_lines):
            print(f"Layout detectado por nome: {layout_name}")
            return layout_name, layout_data

    # 2. DETECÇÃO POR CONTEÚDO
    layout_by_content = detect_by_content_analysis(non_empty_lines, layouts_cache)
    if layout_by_content:
        layout_name, layout_data = layout_by_content
        print(f"Layout detectado por conteúdo: {layout_name}")
        return layout_name, layout_data

    return None, None

def load_all_layouts() -> Dict:
    """Carrega todos os layouts JSON da pasta layouts/"""
    layouts_cache = {}
    layouts_directory = "layouts"

    if not os.path.exists(layouts_directory):
        print(f"Diretório de layouts não encontrado: {layouts_directory}")
        return {}

    for filename in os.listdir(layouts_directory):
        if filename.endswith('.json'):
            layout_path = os.path.join(layouts_directory, filename)
            try:
                with open(layout_path, 'r', encoding='utf-8') as f:
                    layout = json.load(f)
                    layout_name = filename.replace('.json', '')
                    layouts_cache[layout_name] = layout
                    print(f"Layout carregado: {layout_name}")
            except Exception as e:
                print(f"Erro ao carregar layout {filename}: {e}")

    return layouts_cache


def detect_by_filename(filename: str, layouts_cache: Dict) -> Optional[Tuple[str, Dict]]:
    """Detecta layout baseado no nome do arquivo"""
    filename_lower = filename.lower()

    # Padrões correspondentes aos seus arquivos JSON
    filename_patterns = {
        # ADICIONAR ESTAS ENTRADAS:
        'PASSIVOS_ATUARIAIS_ATIVOS': [
            r'.*ativos.*\.txt',
            r'.*passivos.*ativos.*\.txt',
            r'.*atuariais.*ativos.*\.txt',
            r'.*participantes.*ativos.*\.txt'
        ],
        'PASSIVOS_ATUARIAIS_PROCED': [
            r'.*proced.*\.txt',
            r'.*procedimentos.*\.txt',
            r'.*passivos.*proced.*\.txt',
            r'.*calculos.*atuariais.*\.txt'
        ],
        'PASSIVOS_ATUARIAIS_APOPEN': [
            r'.*apopen.*\.txt',
            r'.*aposentados.*\.txt',
            r'.*pensionistas.*\.txt',
            r'.*aposentados.*pensionistas.*\.txt',
            r'.*beneficiarios.*\.txt'
        ],
        'DMED_DECLARACAO_SERVICOS_MEDICOS': [
            r'.*dmed.*\.txt',
            r'.*declaracao.*servicos.*medicos.*\.txt',
            r'.*declaracao.*saude.*\.txt',
            r'.*servicos.*medicos.*saude.*\.txt'
        ],
        'FOLHA-PETROBRAS-ENVELOPADO': [
            r'.*envelop.*\.txt',
            r'.*p0810u2.*\.txt',
            r'.*smov.*\.txt'
        ],
        'FOLHA-PETROBRAS-MVF': [
            r'.*mvf.*\.txt',
            r'.*p0810d3.*\.txt'
        ],
        'RETORNO_FOLHA_PETROS': [
            r'.*retorno.*petros.*\.txt',
            r'.*p0810w5.*\.txt',
            r'.*folha.*petros.*retorno.*\.txt'
        ],
        'FOLHA-TRANSPETRO': [
            r'.*fotransp.*\.txt',
            r'.*transpetro.*\.txt'
        ],
        'FOLHA-PBIO': [
            r'.*fopbio.*\.txt',
            r'.*pbio.*\.txt'
        ],
        'FOLHA-TERMOBAHIA': [
            r'.*foteba.*\.txt',
            r'.*termobahia.*\.txt'
        ],
        'FOLHA-TBG': [
            r'.*fotbg.*\.txt',
            r'.*tbg.*\.txt'
        ],
        'SECA': [
            r'.*seca.*teste.*\.txt',
            r'.*servico.*calculado.*\.txt'
        ],
        'Faturamento': [
            r'.*faturamento.*\.txt'
        ],
        # ADICIONAR ESTA ENTRADA:
        'DET_NOTA_DEBITO_APS_PETROBRAS': [
            r'.*det.*nota.*debito.*\.txt',
            r'.*aps.*petrobras.*\.txt',
            r'.*nota.*debito.*aps.*\.txt',
            r'.*det.*debito.*petrobras.*\.txt'
        ],
        # ADICIONAR ESTA ENTRADA:
        'DET_NOTA_DEBITO_COLIGADAS': [
            r'.*det.*nota.*debito.*coligadas.*\.txt',
            r'.*nota.*debito.*coligadas.*\.txt',
            r'.*det.*debito.*coligadas.*\.txt',
            r'.*coligadas.*debito.*\.txt'
        ]
    }

    for layout_name, patterns in filename_patterns.items():
        for pattern in patterns:
            if re.match(pattern, filename_lower):
                if layout_name in layouts_cache:
                    return layout_name, layouts_cache[layout_name]

    return None


def detect_by_content_analysis(lines: List[str], layouts_cache: Dict) -> Optional[Tuple[str, Dict]]:
    """Detecta layout analisando o conteúdo das linhas"""
    if not lines:
        return None

    # Calcular tamanho mais comum das linhas
    line_lengths = [len(line) for line in lines]
    most_common_length = max(set(line_lengths), key=line_lengths.count)

    # Analisar primeiros caracteres
    first_chars = [line[:2] if len(line) >= 2 else line for line in lines if line]
    first_char_single = [line[:1] if len(line) >= 1 else '' for line in lines if line]

    print(f"Tamanho mais comum: {most_common_length}")
    print(f"Primeiros 2 caracteres: {set(first_chars)}")

    # REGRAS DE DETECÇÃO POR CONTEÚDO

    # 1. Layouts de 240 posições
    if most_common_length >= 238 and most_common_length <= 242:

        # Retorno Folha Petrobras (00, 01, 99)
        if any(line.startswith(('00', '01', '99')) for line in lines):
            return 'RETORNO-FOLHA-PETROBRAS', layouts_cache.get('RETORNO-FOLHA-PETROBRAS')

        # Folha Petrobras Envelopado (C, 0, R)
        elif any(line.startswith(('C', '0', 'R')) for line in lines):
            return 'FOLHA-PETROBRAS-ENVELOPADO', layouts_cache.get('FOLHA-PETROBRAS-ENVELOPADO')

        # Serviço Calculado (00SERVICO CALCULADO)
        elif any('SERVICO CALCULADO' in line for line in lines):
            return 'SECA', layouts_cache.get('SECA')

        # MVF (contém MOVFL ou LIBER)
        elif any('MOVFL' in line for line in lines) or any('LIBER' in line for line in lines):
            return 'FOLHA-PETROBRAS-MVF', layouts_cache.get('FOLHA-PETROBRAS-MVF')
        # 1.5. Layouts de 200 posições (Petros)
    elif most_common_length >= 198 and most_common_length <= 202:

            # Verificar padrões específicos para layouts Petros
            # Petros usa registros 01 (header), 02 (detalhe), 09 (trailer)
            if any(line.startswith(('01', '02', '09')) for line in lines):

                # Verificar se tem características específicas do envio
                # (campos como ID empresa '019090' na posição 3-8)
                has_empresa_019090 = any('019090' in line for line in lines)

                # Verificar se tem datas no formato DDMMAAAA (envio) vs competência AAAAMM (retorno)
                has_ddmmaaaa_pattern = any(re.search(r'\d{8}', line) for line in lines)

                if has_empresa_019090:
                    return 'ENVIO_FOLHA_PETROS', layouts_cache.get('ENVIO_FOLHA_PETROS')
                else:

                        # Padrão: assumir envio se não identificou claramente como retorno
                        return 'ENVIO_FOLHA_PETROS', layouts_cache.get('ENVIO_FOLHA_PETROS')

            # Se não tem registros 01/02/09, ainda pode ser Petros genérico
            return 'ENVIO_FOLHA_PETROS', layouts_cache.get('ENVIO_FOLHA_PETROS')
            # 1.6. Layouts de 106 posições (Envio Boletos)
    elif most_common_length >= 104 and most_common_length <= 108:

        # Verificar padrões específicos para envio de boletos
        # Boletos usa registros C (header), 0 (detalhe), R (trailer)
        if any(line.startswith(('C', '0', 'R')) for line in lines):

            # Verificar se tem características específicas do envio de boletos
            # Interface YSEFI_AMS nas posições 2-41
            has_ysefi_interface = any('YSEFI_AMS' in line for line in lines)

            if has_ysefi_interface:
                return 'ENVIO_BOLETOS', layouts_cache.get('ENVIO_BOLETOS')

            # Verificar outros padrões típicos de boletos (CPF, valores, etc.)
            has_cpf_pattern = any(re.search(r'0\d{6}\d{8}\d{2}\d{3}\d{18}\d{11}', line) for line in lines)
            if has_cpf_pattern:
                return 'ENVIO_BOLETOS', layouts_cache.get('ENVIO_BOLETOS')

        # Se tem 106 posições mas não identificou padrões específicos
        return 'ENVIO_BOLETOS', layouts_cache.get('ENVIO_BOLETOS')

        # 1.7. Layouts de 170 posições (Det. Nota Débito APS-Petrobras)
    elif most_common_length >= 168 and most_common_length <= 172:

        # Verificar padrões específicos para nota débito APS
        # APS usa registros 05 (header), 10 (detalhe), 15 (trailer)
        if any(line.startswith(('05', '10', '15')) for line in lines):

            # Verificar se tem características específicas do APS
            # Montante com vírgula nas posições específicas
            has_montante_pattern = any(re.search(r'10.{12}\d{11},\d{2}', line) for line in lines)

            if has_montante_pattern:
                return 'DET_NOTA_DEBITO_APS_PETROBRAS', layouts_cache.get('DET_NOTA_DEBITO_APS_PETROBRAS')

            # Verificar outros padrões típicos (centro de custo, códigos)
            has_aps_pattern = any(re.search(r'05.{38}', line) for line in lines)
            if has_aps_pattern:
                return 'DET_NOTA_DEBITO_APS_PETROBRAS', layouts_cache.get('DET_NOTA_DEBITO_APS_PETROBRAS')

        # Se tem 170 posições mas não identificou padrões específicos
        return 'DET_NOTA_DEBITO_APS_PETROBRAS', layouts_cache.get('DET_NOTA_DEBITO_APS_PETROBRAS')

        # 1.8. Layouts de 168 posições (Det. Nota Débito Coligadas)
    elif most_common_length >= 166 and most_common_length <= 170:

        # Verificar padrões específicos para nota débito Coligadas
        # Coligadas usa apenas registros 05 (todos os tipos)
        if any(line.startswith('05') for line in lines):

            # Verificar se tem características específicas das Coligadas
            # Matrícula SAI nas primeiras 8 posições + filler
            has_matricula_pattern = any(re.search(r'^\d{8}.', line) for line in lines)

            # Verificar padrão de totais no trailer
            has_totais_pattern = any('TOTAIS' in line or line.startswith('05') and 'TOT' in line for line in lines)

            if has_matricula_pattern or has_totais_pattern:
                return 'DET_NOTA_DEBITO_COLIGADAS', layouts_cache.get('DET_NOTA_DEBITO_COLIGADAS')

            # Verificar outros padrões (nomes, valores calculados)
            has_valor_calc_pattern = any(re.search(r'.{230}\d', line) for line in lines)
            if has_valor_calc_pattern:
                return 'DET_NOTA_DEBITO_COLIGADAS', layouts_cache.get('DET_NOTA_DEBITO_COLIGADAS')

        # Se tem tamanho próximo a 168 posições, assumir como Coligadas
        return 'DET_NOTA_DEBITO_COLIGADAS', layouts_cache.get('DET_NOTA_DEBITO_COLIGADAS')

        # 1.9. Layout DMED (Declaração Serviços Médicos) - tamanho variável com delimitadores
        # Este layout usa delimitador "|" e não tem tamanho fixo

        # Verificar se tem delimitadores pipe "|" e identificadores DMED
    has_pipe_delimiters = any('|' in line for line in lines)

    if has_pipe_delimiters:
            # Verificar padrões específicos do DMED
            dmed_identifiers = ['Dmed|', 'RESPO|', 'DECPJ|', 'OPPAS|', 'TOP|', 'RTOP|',
                                'DTOP|', 'RDTOP|', 'PSS|', 'RPPSS|', 'BRPPSS|', 'FIMDmed|']

            has_dmed_identifiers = any(any(identifier in line for identifier in dmed_identifiers)
                                       for line in lines)

            if has_dmed_identifiers:
                # Verificar estrutura típica DMED (deve começar com Dmed e terminar com FIMDmed)
                first_line_dmed = any(line.startswith('Dmed|') for line in lines[:3])
                last_line_fimdmed = any(line.startswith('FIMDmed|') for line in lines[-3:])

                if first_line_dmed or last_line_fimdmed:
                    return 'DMED_DECLARACAO_SERVICOS_MEDICOS', layouts_cache.get('DMED_DECLARACAO_SERVICOS_MEDICOS')

                # Verificar outros padrões DMED (CNPJ, CPF, etc.)
                has_cnpj_cpf_pattern = any(re.search(r'(DECPJ|RESPO|TOP|RPPSS).*\|\d{11,14}\|', line)
                                           for line in lines)
                if has_cnpj_cpf_pattern:
                    return 'DMED_DECLARACAO_SERVICOS_MEDICOS', layouts_cache.get('DMED_DECLARACAO_SERVICOS_MEDICOS')

    elif most_common_length >= 200 and most_common_length <= 500:

        # Verificar padrões específicos de dados atuariais em arquivos posicionais

        # ATIVOS - geralmente têm matrícula + dados pessoais + salário + tempo serviço
        # Padrão: matrícula (8-12 pos) + CPF (11 pos) + nome (30-50 pos) + data nasc + salário
        has_ativos_pattern = any(
            re.search(r'^\d{8,12}.{11}.{30,60}\d{8}.*\d{10,15}', line) for line in lines[:20]
        )

        if has_ativos_pattern:
            return 'PASSIVOS_ATUARIAIS_ATIVOS', layouts_cache.get('PASSIVOS_ATUARIAIS_ATIVOS')

        # APOPEN - aposentados/pensionistas têm padrão similar mas com data aposentadoria
        # e valores de benefício ao invés de salário
        has_apopen_pattern = any(
            re.search(r'^\d{8,12}.{11}.{30,60}\d{8}.*[AP].*\d{8}', line) for line in lines[:20]
        )

        if has_apopen_pattern:
            return 'PASSIVOS_ATUARIAIS_APOPEN', layouts_cache.get('PASSIVOS_ATUARIAIS_APOPEN')

        # PROCED - procedimentos têm códigos e valores de cálculo
        has_proced_pattern = any(
            re.search(r'^[A-Z0-9]{5,15}.*\d{8}.*\d{12,18}', line) for line in lines[:20]
        )

        if has_proced_pattern:
            return 'PASSIVOS_ATUARIAIS_PROCED', layouts_cache.get('PASSIVOS_ATUARIAIS_PROCED')

        # Se tem tamanho de passivos atuariais mas não identificou padrão específico
        # Verificar por palavras-chave nos primeiros caracteres
        first_chars_analysis = ''.join(lines[:5]).upper()

        if any(keyword in first_chars_analysis for keyword in ['ATIVO', 'SALARIO', 'ADMISSAO']):
            return 'PASSIVOS_ATUARIAIS_ATIVOS', layouts_cache.get('PASSIVOS_ATUARIAIS_ATIVOS')
        elif any(keyword in first_chars_analysis for keyword in ['APOSEN', 'PENSAO', 'BENEFIC']):
            return 'PASSIVOS_ATUARIAIS_APOPEN', layouts_cache.get('PASSIVOS_ATUARIAIS_APOPEN')
        elif any(keyword in first_chars_analysis for keyword in ['PROCED', 'CALC', 'TAXA']):
            return 'PASSIVOS_ATUARIAIS_PROCED', layouts_cache.get('PASSIVOS_ATUARIAIS_PROCED')

    # 2. Layouts de 37-38 posições
    elif most_common_length >= 35 and most_common_length <= 40:

        # Verificar se tem sinal no meio (posição 13)
        has_sign_at_13 = any(len(line) > 12 and line[12] in ['+', '-'] for line in lines)

        if most_common_length == 37:
            # Transpetro (37 posições, sem sinal)
            return 'FOLHA-TRANSPETRO', layouts_cache.get('FOLHA-TRANSPETRO')

        elif most_common_length == 38 and has_sign_at_13:
            # PBIO, Termobahia ou TBG (38 posições com sinal)
            # Por padrão retorna PBIO (pode ser refinado depois)
            return 'FOLHA-PBIO', layouts_cache.get('FOLHA-PBIO')

    return None


def validate_layout_with_content(layout_data: Dict, lines: List[str]) -> bool:
    """Validação rápida se o layout bate com o conteúdo"""
    if not lines or not layout_data:
        return False

    # Verificar tamanho
    expected_size = layout_data.get('record_size', 0)
    if expected_size > 0:
        avg_size = sum(len(line) for line in lines) / len(lines)
        # Tolerância de ±5 caracteres
        if abs(avg_size - expected_size) > 5:
            return False

    return True


def analisar_arquivo_estruturado(arquivo1_path, arquivo2_path, layout_info):
    """
    Analisa diferenças entre arquivos estruturados usando um layout.

    Args:
        arquivo1_path: Caminho para o primeiro arquivo
        arquivo2_path: Caminho para o segundo arquivo
        layout_info: Informações do layout (dicionário de campos ou informações de campos)

    Returns:
        Lista de diferenças encontradas
    """


    diferencas = []

    try:
        # Verificação rápida se os arquivos são idênticos
        if filecmp.cmp(arquivo1_path, arquivo2_path, shallow=False):
            return []  # Arquivos idênticos, retorna lista vazia

        # Obter campos do layout
        campos_layout = {}
        if isinstance(layout_info, dict) and 'fields' in layout_info:
            # Layout novo formato
            campos_layout = layout_info['fields']
        else:
            # Layout legado ou formato simplificado
            campos_layout = layout_info

        # Carregar arquivos
        try:
            with open(arquivo1_path, 'r', encoding='utf-8', errors='replace') as f1:
                linhas1 = f1.readlines()
        except Exception as e:
            print(f"Erro ao ler o arquivo 1: {str(e)}")
            return [{'linha': 0, 'tipo': 'erro', 'mensagem': f'Erro ao ler o arquivo 1: {str(e)}'}]

        try:
            with open(arquivo2_path, 'r', encoding='utf-8', errors='replace') as f2:
                linhas2 = f2.readlines()
        except Exception as e:
            print(f"Erro ao ler o arquivo 2: {str(e)}")
            return [{'linha': 0, 'tipo': 'erro', 'mensagem': f'Erro ao ler o arquivo 2: {str(e)}'}]



        # Pré-processamento do layout por tipo de registro (otimização)
        layout_por_tipo = {}
        for campo, info in campos_layout.items():
            tipo = info.get('tipo_registro')
            if tipo not in layout_por_tipo:
                layout_por_tipo[tipo] = []
            layout_por_tipo[tipo].append((campo, info))

        # Otimização: processo apenas números mínimos de linhas
        min_linhas = min(len(linhas1), len(linhas2))

        print(f"Layout recebido: {type(layout_info)}")
        print(f"Campos do layout: {list(campos_layout.keys()) if campos_layout else 'Nenhum'}")
        print(f"Layout por tipo: {layout_por_tipo}")
        # Usar uma estratégia mais eficiente para comparar linhas
        # Para arquivos grandes, comparamos linha a linha ao invés de usar difflib
        # que pode consumir muita memória
        # --- INÍCIO: NOVA LÓGICA SEM EFEITO CASCATA ---
        from difflib import SequenceMatcher

        matcher = SequenceMatcher(a=linhas1, b=linhas2, autojunk=False)
        opcodes = matcher.get_opcodes()

        MAX_DETAILED_LINES = 500
        detalhadas = 0

        def comparar_por_layout(idx1, idx2):
            nonlocal detalhadas
            if detalhadas >= MAX_DETAILED_LINES:
                return

            linha1 = linhas1[idx1].rstrip('\n')
            linha2 = linhas2[idx2].rstrip('\n')

            tipo_primo = linha1[0:2] if len(linha1) >= 2 else ""

            # Quando não há layout específico por tipo (ou é tipo "01"), compara campos
            if None in layout_por_tipo and (tipo_primo == "01" or not layout_por_tipo.get(tipo_primo)):
                linha_diferencas = []
                for campo, info in layout_por_tipo[None]:
                    pos_inicial = info.get('pos_inicial', 1) - 1
                    tamanho = info.get('tamanho', 0)
                    if (pos_inicial + tamanho <= len(linha1) and
                            pos_inicial + tamanho <= len(linha2)):
                        v1 = linha1[pos_inicial:pos_inicial + tamanho].strip()
                        v2 = linha2[pos_inicial:pos_inicial + tamanho].strip()
                        if v1 != v2:
                            linha_diferencas.append({
                                'campo': campo,
                                'descricao': info.get('description', ''),
                                'valor_antigo': v1,
                                'valor_novo': v2
                            })
                if linha_diferencas:
                    diferencas.append({
                        'linha': idx1 + 1,  # numeração baseada no arquivo 1
                        'tipo': 'campos_alterados',
                        'diferenca': linha_diferencas
                    })
                else:
                    diferencas.append({
                        'linha': idx1 + 1,
                        'tipo': 'linha_alterada',
                        'valor_antigo': linha1,
                        'valor_novo': linha2
                    })
            else:
                # Sem layout específico por tipo → registra como linha alterada
                diferencas.append({
                    'linha': idx1 + 1,
                    'tipo': 'linha_alterada',
                    'valor_antigo': linha1,
                    'valor_novo': linha2
                })
            detalhadas += 1

        for tag, i1, i2, j1, j2 in opcodes:
            if detalhadas >= MAX_DETAILED_LINES:
                diferencas.append({
                    'linha': 0,
                    'tipo': 'info',
                    'mensagem': f'Análise limitada às primeiras {MAX_DETAILED_LINES} linhas diferentes por questões de desempenho.'
                })
                break

            if tag == 'equal':
                continue

            elif tag == 'delete':
                # Linhas presentes só no arquivo 1 (removidas)
                for i in range(i1, i2):
                    if detalhadas >= MAX_DETAILED_LINES: break
                    diferencas.append({
                        'linha': i + 1,
                        'tipo': 'linha_removida',
                        'valor_antigo': linhas1[i].rstrip('\n'),
                        'valor_novo': ''
                    })
                    detalhadas += 1

            elif tag == 'insert':
                # Linhas presentes só no arquivo 2 (adicionadas)
                for j in range(j1, j2):
                    if detalhadas >= MAX_DETAILED_LINES: break
                    diferencas.append({
                        'linha': j + 1,  # pode usar a numeração do novo arquivo
                        'tipo': 'linha_adicionada',
                        'valor_antigo': '',
                        'valor_novo': linhas2[j].rstrip('\n')
                    })
                    detalhadas += 1

            elif tag == 'replace':
                # 1) compara par-a-par onde há linhas em ambos os lados
                common = min(i2 - i1, j2 - j1)
                for k in range(common):
                    comparar_por_layout(i1 + k, j1 + k)

                # 2) sobras no arquivo 1 → removidas
                for i in range(i1 + common, i2):
                    if detalhadas >= MAX_DETAILED_LINES: break
                    diferencas.append({
                        'linha': i + 1,
                        'tipo': 'linha_removida',
                        'valor_antigo': linhas1[i].rstrip('\n'),
                        'valor_novo': ''
                    })
                    detalhadas += 1

                # 3) sobras no arquivo 2 → adicionadas
                for j in range(j1 + common, j2):
                    if detalhadas >= MAX_DETAILED_LINES: break
                    diferencas.append({
                        'linha': j + 1,
                        'tipo': 'linha_adicionada',
                        'valor_antigo': '',
                        'valor_novo': linhas2[j].rstrip('\n')
                    })
                    detalhadas += 1
        # --- FIM: NOVA LÓGICA ---


        # Adicionar linhas exclusivas de cada arquivo
        # Limitar análise também por desempenho
        MAX_DISTINCT_LINES = 100

        # Linhas exclusivas do arquivo 1


        return diferencas

    except Exception as e:
        print(f"Erro ao analisar arquivos estruturados: {str(e)}")
        import traceback
        traceback.print_exc()
        return [{
            'linha': 0,
            'tipo': 'erro',
            'mensagem': f'Erro ao analisar arquivos: {str(e)}'
        }]


def extrair_dicionario_layout():
    """
    Função que cria o dicionário de layout para arquivos SECA.
    MODIFICADA: Exclui campos que não serão mais gerados no novo sistema (marcados em azul).
    """

    # Lista de campos que devem ser EXCLUÍDOS (campos marcados em azul nas imagens)
    campos_excluidos = [
        "ORPA_CD_ORG_PAG",  # Linha com ORPA-CD-ORG-PAG
        "CONL_CD_LOCALIDADE",  # Linha com CONL-CD-LOCALIDADE
        "CRAM_CD_INDIC_PAG",  # Linha com CRAM-CD-INDIC-PAG
        "RSPI_VL_IR_CALCULADO",  # Linha com RSPI-VL-IR-CALCULADO
        "RSPI_VL_INSS_CALCULADO",  # Linha com RSPI-VL-INSS-CALCULADO
        "RSPI_VL_ISS_CALCULADO",  # Linha com RSPI-VL-ISS-CALCULADO
        "VALOR_DE_COFINS",  # Linha com VALOR-DE-COFINS
        "PERCENTUAL_DE_COFINS",  # Linha com PERCENTUAL-DE-COFINS
        "VALOR_DE_PIS",  # Linha com VALOR-DE-PIS
        "PERCENTUAL_DE_PIS",  # Linha com PERCENTUAL-DE-PIS
        "VALOR_DE_CSLL",  # Linha com VALOR-DE-CSLL
        "PERCENTUAL_DE_CSLL",  # Linha com PERCENTUAL-DE-CSLL
        "VALOR_DE_IR",  # Linha com VALOR-DE-IR
        "PERCENTUAL_DE_IR",  # Linha com PERCENTUAL-DE-IR
        "VALOR_LEI_10833_COOPERATIVA",  # Linha com VALOR-LEI-10833-COOPERATIVA
        "PERCENTUAL_LEI_10833",  # Linha com PERCENTUAL-LEI-10833
        "VALOR_LEI_10833_NAO_COOPERATIVA",  # Linha com VALOR-LEI-10833-NÃO-COOPERATIVA
        "PERCENTUAL_LEI_10833_NAO_COOPERATIVA",  # Linha com PERCENTUAL-LEI-10833-NÃO-COOPERATIVA
        "VALOR_IR_PESSOA_FISICA"  # Linha com VALOR-IR-PESSOA-FISICA
    ]

    layout = {
        # Header (tipo 00)
        "TIPO_REGISTRO_HEADER": {"tipo_registro": "00", "pos_inicial": 1, "tamanho": 2,
                                 "descricao": "Tipo do Registro (Header)"},
        "TAB_ENVIADA": {"tipo_registro": "00", "pos_inicial": 3, "tamanho": 22, "descricao": "SERVICO CALCULADO"},
        "ID_REGISTRO": {"tipo_registro": "00", "pos_inicial": 25, "tamanho": 18,
                        "descricao": "Número da Versão do arquivo"},
        "NOME_INTERFACE": {"tipo_registro": "00", "pos_inicial": 43, "tamanho": 40, "descricao": "MIGRACAO TEMPO"},
        "EMPRESA": {"tipo_registro": "00", "pos_inicial": 83, "tamanho": 2, "descricao": "Empresa"},
        "ORGAO": {"tipo_registro": "00", "pos_inicial": 85, "tamanho": 22, "descricao": "TIC/CPSW/IST-I/PTEC-RJ"},
        "ORIGEM": {"tipo_registro": "00", "pos_inicial": 107, "tamanho": 10, "descricao": "SAM"},
        "DESTINO": {"tipo_registro": "00", "pos_inicial": 117, "tamanho": 10, "descricao": "CRC"},
        "DT_INI_CRIACAO": {"tipo_registro": "00", "pos_inicial": 127, "tamanho": 8,
                           "descricao": "Data da geração do arquivo"},
        "HH_INI_CRIACAO": {"tipo_registro": "00", "pos_inicial": 135, "tamanho": 6, "descricao": "Hora da geração"},
        "DATA_DE_REFERENCIA": {"tipo_registro": "00", "pos_inicial": 142, "tamanho": 6,
                               "descricao": "Data de referência (AAAAMM)"},

        # Registro de Serviço Calculado (tipo 01)
        "TIPO_REGISTRO": {"tipo_registro": "01", "pos_inicial": 1, "tamanho": 2,
                          "descricao": "Tipo do Registro (Serviço Calculado)"},
        "NUMERO_LINHA": {"tipo_registro": "01", "pos_inicial": 3, "tamanho": 8, "descricao": "Número da linha gravada"},
        "CRAM_CD_CREDENCIAD": {"tipo_registro": "01", "pos_inicial": 11, "tamanho": 14,
                               "descricao": "CPF/CNPJ do Credenciado"},
        "CRAM_NM_CREDENCIAD": {"tipo_registro": "01", "pos_inicial": 25, "tamanho": 50,
                               "descricao": "Nome do Prestador ou Credenc."},

        # CAMPO REMOVIDO: "ORPA_CD_ORG_PAG" (posição 75-78)
        # CAMPO REMOVIDO: "CONL_CD_LOCALIDADE" (posição 79-82)

        "EMPA_CD_EMPRESA": {"tipo_registro": "01", "pos_inicial": 83, "tamanho": 2,
                            "descricao": "Empresa do Beneficiário"},
        "BEAM_NR_MATRICULA": {"tipo_registro": "01", "pos_inicial": 85, "tamanho": 8,
                              "descricao": "Matricula do Beneficiário"},
        "NUMERO_BENEFICIARIO": {"tipo_registro": "01", "pos_inicial": 93, "tamanho": 2,
                                "descricao": "Número do Beneficiário"},
        "CRAM_CD_CBI": {"tipo_registro": "01", "pos_inicial": 95, "tamanho": 4, "descricao": "CBI do Credenciado"},
        "BEAM_CD_CBI": {"tipo_registro": "01", "pos_inicial": 99, "tamanho": 4, "descricao": "CBI do beneficiário"},
        "COBI_CD_CBI": {"tipo_registro": "01", "pos_inicial": 103, "tamanho": 4,
                        "descricao": "Código do CBI responsável pelo documento"},
        "RSPA_CD_RISCO": {"tipo_registro": "01", "pos_inicial": 107, "tamanho": 1,
                          "descricao": "Código do risco do documento"},
        "RSPA_NR_AUTORIZA": {"tipo_registro": "01", "pos_inicial": 108, "tamanho": 5,
                             "descricao": "Numero da Autorizacao"},
        "RSPA_CD_CBI": {"tipo_registro": "01", "pos_inicial": 113, "tamanho": 4,
                        "descricao": "CBI responsável pela autorização"},
        "RSPA_DT_ENTRADA": {"tipo_registro": "01", "pos_inicial": 117, "tamanho": 10,
                            "descricao": "Data de início da internação/atendimento"},
        "RSPA_DT_SAIDA": {"tipo_registro": "01", "pos_inicial": 127, "tamanho": 10,
                          "descricao": "Data de saída da internação/atendimento"},
        "RSPI_DT_REALIZACAO": {"tipo_registro": "01", "pos_inicial": 137, "tamanho": 10,
                               "descricao": "Data de realização do serviço"},
        "RSPI_QN_OCORRENCIAS": {"tipo_registro": "01", "pos_inicial": 147, "tamanho": 3,
                                "descricao": "Quantidade de ocorrências pagas do serviço"},
        "RSPI_QN_OCOR_COBR": {"tipo_registro": "01", "pos_inicial": 150, "tamanho": 3,
                              "descricao": "Quantidade de ocorrências cobradas do serviço"},
        "RSPI_CD_PERC_ADIC": {"tipo_registro": "01", "pos_inicial": 153, "tamanho": 1,
                              "descricao": "Código percentual adicional"},
        "RSPI_CD_USO_INTERN": {"tipo_registro": "01", "pos_inicial": 154, "tamanho": 1,
                               "descricao": "Código de uso interno"},
        "RSPI_VL_SERV_CALCU": {"tipo_registro": "01", "pos_inicial": 155, "tamanho": 11,
                               "descricao": "Valor serviço calculado"},
        "RSPI_VL_APRESENTAD": {"tipo_registro": "01", "pos_inicial": 166, "tamanho": 11,
                               "descricao": "Valor apresentado para pagamento"},
        "VASE_VL_SERVICO": {"tipo_registro": "01", "pos_inicial": 177, "tamanho": 11,
                            "descricao": "Valor serviço em usm"},
        "VALOR_PARC_BENEF_USM": {"tipo_registro": "01", "pos_inicial": 188, "tamanho": 11,
                                 "descricao": "Valor parcela em usm"},
        "RSPI_DT_ATUALIZADO": {"tipo_registro": "01", "pos_inicial": 199, "tamanho": 10,
                               "descricao": "Data de Atualização"},
        "RSPA_DT_EMISSAO": {"tipo_registro": "01", "pos_inicial": 209, "tamanho": 10, "descricao": "Data de Emissão"},
        "SEHI_CD_SERVICO": {"tipo_registro": "01", "pos_inicial": 219, "tamanho": 8,
                            "descricao": "Código Serviço CBHPM"},
        "RSPA_NR_DOCUM_EST": {"tipo_registro": "01", "pos_inicial": 227, "tamanho": 6,
                              "descricao": "Número Documento Estornado"},

        # CAMPOS REMOVIDOS (marcados em azul nas imagens):
        # - Todos os campos relacionados a cálculos de impostos (IR, INSS, ISS, COFINS, PIS, CSLL)
        # - Campos relacionados à Lei 10833
        # - Campo de IR Pessoa Física
        # Estes campos não serão mais gerados no novo sistema

        # Trailer (tipo 99)
        "TIPO_REGISTRO_TRAILER": {"tipo_registro": "99", "pos_inicial": 1, "tamanho": 2,
                                  "descricao": "Tipo do Registro (TRAILER)"},
        "DT_FIM_CRIACAO": {"tipo_registro": "99", "pos_inicial": 3, "tamanho": 8, "descricao": "TRAILER"},
        "HH_FIM_CRIACAO": {"tipo_registro": "99", "pos_inicial": 11, "tamanho": 6, "descricao": "Formato: HHMMSS"},
        "QUANTIDADE_REG_01": {"tipo_registro": "99", "pos_inicial": 17, "tamanho": 8,
                              "descricao": "Quantidade Total de Registros 01"},
        "TOTAL_REG_DE_DADOS": {"tipo_registro": "99", "pos_inicial": 25, "tamanho": 8,
                               "descricao": "Quantidade Total de Registros no Arquivo"}
    }

    # FILTRAR campos excluídos do layout final
    layout_filtrado = {}
    for campo, info in layout.items():
        if campo not in campos_excluidos:
            layout_filtrado[campo] = info

    print(f"Layout SECA: {len(campos_excluidos)} campos excluídos do novo sistema")
    print(f"Campos ativos no layout: {len(layout_filtrado)}")

    return layout_filtrado

def gerar_html_visao_negocio_estruturada(diferencas):
    """
    Gera HTML para visualização de negócio das diferenças em arquivos estruturados.
    Versão melhorada que produz resultado similar à imagem 1.
    """
    # Obter o nome do layout da sessão, se disponível
    layout_name = session.get('layout_name', 'Layout Identificado Automaticamente')
    print('CAIU NO LAYOUT NAME' +layout_name)

    if not diferencas:
        return f'<div class="business-summary"><p>Não foram encontradas diferenças entre os arquivos usando o layout: {layout_name}.</p></div>'

    # Verificar se temos objeto de diferenças válido
    if not isinstance(diferencas, list):
        return '<div class="business-summary"><p>Erro na análise: formato de diferenças inválido.</p></div>'

    # Contagem de tipos de diferenças
    contador = {
        'campos_alterados': 0,
        'linha_alterada': 0,
        'linha_adicionada': 0,
        'linha_removida': 0,
        'erro': 0,
        'info': 0
    }
    campos_total = 0

    # Iniciar a contagem
    for diff in diferencas:
        tipo = diff.get('tipo', '')
        if tipo in contador:
            contador[tipo] += 1

        if tipo == 'campos_alterados':
            campos_total += len(diff.get('diferenca', []))

    # Verificar se há erros na análise
    has_error = contador['erro'] > 0
    if has_error:
        html_parts = ['<div class="error-message">']
        html_parts.append('<h3>Erros na Análise</h3>')
        for diff in diferencas:
            if diff.get('tipo') == 'erro':
                html_parts.append(f'<p>{diff.get("mensagem", "Erro desconhecido")}</p>')
        html_parts.append('</div>')
        return ''.join(html_parts)

    # Lista para construção do HTML
    html_parts = []

    # Cabeçalho principal
    html_parts.append('<div class="business-summary">')
    html_parts.append(f'<h3>Análise Detalhada - Layout: {layout_name}</h3>')
    html_parts.append(
        f'<p>Total de diferenças encontradas: <strong>{len([d for d in diferencas if d.get("tipo") not in ["info", "erro"]])}</strong></p>')
    html_parts.append('</div>')

    # Processar informações
    for diff in diferencas:
        if diff.get('tipo') == 'info':
            html_parts.append(f'<div class="info-message"><p>{diff.get("mensagem", "")}</p></div>')

    # Filtrar apenas diferenças reais (excluir info e erro)
    real_diffs = [d for d in diferencas if d.get('tipo') not in ['info', 'erro']]

    # Agrupar por linha para criar uma visualização mais organizada
    lines_with_changes = {}
    for diff in real_diffs:
        linha = diff.get('linha', 0)
        if linha not in lines_with_changes:
            lines_with_changes[linha] = []
        lines_with_changes[linha].append(diff)

    # Processar cada linha com alterações
    for linha in sorted(lines_with_changes.keys()):
        changes = lines_with_changes[linha]
        print('VALIDANDO LINHAS' + str(changes))

        # Verificar se há campos alterados nesta linha
        campos_alterados_diff = next((d for d in changes if d.get('tipo') == 'campos_alterados'), None)
        # Debug mais detalhado
        print(f'CHANGES na linha {linha}: {changes}')
        print(f'TIPOS encontrados: {[d.get("tipo") for d in changes]}')
        print(f'CAMPOS ALTERADOS DIFF: {campos_alterados_diff}')

        if campos_alterados_diff:
            # Renderizar como a imagem 1: cabeçalho + explicação + tabela
            html_parts.append('<div class="change-container">')

            # Cabeçalho da linha
            html_parts.append(f'<div class="line-header">Linha {linha} - Campos alterados</div>')

            # Explicação em destaque (cor amarela como na imagem)
            num_campos = len(campos_alterados_diff.get('diferenca', []))
            html_parts.append('<div class="change-explanation-highlight">')
            html_parts.append(f'<em>Foram identificadas alterações em {num_campos} campos nesta linha.</em>')
            html_parts.append('</div>')

            # Tabela estruturada como na imagem 1
            html_parts.append('<table class="fields-comparison-table">')
            html_parts.append('<thead>')
            html_parts.append('<tr>')
            html_parts.append('<th>Campo</th>')
            html_parts.append('<th>Descrição</th>')
            html_parts.append('<th>Valor Original</th>')
            html_parts.append('<th>Novo Valor</th>')
            html_parts.append('</tr>')
            html_parts.append('</thead>')
            html_parts.append('<tbody>')

            for campo_diff in campos_alterados_diff.get('diferenca', []):
                campo_nome = campo_diff.get("campo", "")
                campo_desc = campo_diff.get("descricao", "") or campo_diff.get("description", "")
                valor_antigo = campo_diff.get("valor_antigo", "")
                valor_novo = campo_diff.get("valor_novo", "")

                html_parts.append('<tr>')
                html_parts.append(f'<td class="field-name-cell">{campo_nome}</td>')
                html_parts.append(f'<td class="field-desc-cell">{campo_desc}</td>')
                html_parts.append(f'<td class="old-value-cell">{valor_antigo}</td>')
                html_parts.append(f'<td class="new-value-cell">{valor_novo}</td>')
                html_parts.append('</tr>')

            html_parts.append('</tbody>')
            html_parts.append('</table>')
            html_parts.append('</div>')

        else:
            # Para outros tipos de diferenças (linha alterada, removida, adicionada)
            for diff in changes:
                tipo = diff.get('tipo', '')

                html_parts.append('<div class="change-container">')

                if tipo == 'linha_alterada':
                    html_parts.append(f'<div class="line-header">Linha {linha} - Linha Alterada</div>')
                    html_parts.append('<div class="change-explanation-highlight">')
                    html_parts.append('<em>A linha inteira foi alterada.</em>')
                    html_parts.append('</div>')

                    html_parts.append('<div class="content-comparison">')
                    html_parts.append('<div class="version-section original-version">')
                    html_parts.append('<h4>Versão Original:</h4>')
                    html_parts.append(f'<div class="content-box">{diff.get("valor_antigo", "")}</div>')
                    html_parts.append('</div>')

                    html_parts.append('<div class="version-section new-version">')
                    html_parts.append('<h4>Nova Versão:</h4>')
                    html_parts.append(f'<div class="content-box">{diff.get("valor_novo", "")}</div>')
                    html_parts.append('</div>')
                    html_parts.append('</div>')

                elif tipo == 'linha_adicionada':
                    html_parts.append(f'<div class="line-header">Linha {linha} - Linha Adicionada</div>')
                    html_parts.append('<div class="change-explanation-highlight">')
                    html_parts.append('<em>Esta linha existe apenas na nova versão.</em>')
                    html_parts.append('</div>')

                    html_parts.append('<div class="content-comparison">')
                    html_parts.append('<div class="version-section new-version full-width">')
                    html_parts.append('<h4>Nova Linha:</h4>')
                    html_parts.append(f'<div class="content-box">{diff.get("valor_novo", "")}</div>')
                    html_parts.append('</div>')
                    html_parts.append('</div>')

                elif tipo == 'linha_removida':
                    html_parts.append(f'<div class="line-header">Linha {linha} - Linha Removida</div>')
                    html_parts.append('<div class="change-explanation-highlight">')
                    html_parts.append('<em>Esta linha existe apenas na versão original.</em>')
                    html_parts.append('</div>')

                    html_parts.append('<div class="content-comparison">')
                    html_parts.append('<div class="version-section original-version full-width">')
                    html_parts.append('<h4>Linha Removida:</h4>')
                    html_parts.append(f'<div class="content-box">{diff.get("valor_antigo", "")}</div>')
                    html_parts.append('</div>')
                    html_parts.append('</div>')

                html_parts.append('</div>')

    # Adicionar CSS específico para reproduzir o estilo da imagem 1
    html_parts.append("""
    <style>
    .change-container {
        margin-bottom: 25px;
        border: 1px solid #dee2e6;
        border-radius: 8px;
        overflow: hidden;
        background-color: #fff;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
    }

    .line-header {
        background-color: #e9ecef;
        padding: 12px 15px;
        font-weight: bold;
        color: #495057;
        font-size: 16px;
        border-bottom: 1px solid #dee2e6;
    }

    .change-explanation-highlight {
        background-color: #fff3cd;
        padding: 12px 15px;
        border-bottom: 1px solid #ffeaa7;
        color: #856404;
        font-style: italic;
        font-size: 14px;
    }

    .fields-comparison-table {
        width: 100%;
        border-collapse: collapse;
        margin: 0;
        background-color: #fff;
    }

    .fields-comparison-table thead {
        background-color: #f8f9fa;
    }

    .fields-comparison-table th {
        padding: 12px 15px;
        text-align: left;
        font-weight: bold;
        color: #495057;
        background-color: #f8f9fa;
        border-bottom: 2px solid #dee2e6;
        font-size: 14px;
    }

    .fields-comparison-table td {
        padding: 12px 15px;
        border-bottom: 1px solid #dee2e6;
        vertical-align: top;
        font-size: 13px;
    }

    .field-name-cell {
        font-family: 'Courier New', monospace;
        font-weight: bold;
        background-color: #f8f9fa;
        color: #495057;
    }

    .field-desc-cell {
        color: #6c757d;
        font-style: italic;
    }

    .old-value-cell {
        background-color: #ffeef0;
        color: #721c24;
        font-family: 'Courier New', monospace;
        font-weight: bold;
    }

    .new-value-cell {
        background-color: #e6ffed;
        color: #155724;
        font-family: 'Courier New', monospace;
        font-weight: bold;
    }

    .content-comparison {
        display: flex;
        flex-direction: column;
    }

    @media (min-width: 768px) {
        .content-comparison {
            flex-direction: row;
        }
    }

    .version-section {
        padding: 15px;
        flex: 1;
    }

    .original-version {
        background-color: #ffeef0;
        border-right: 1px solid #f5c6cb;
    }

    .new-version {
        background-color: #e6ffed;
    }

    .full-width {
        width: 100%;
    }

    .content-box {
        font-family: 'Courier New', monospace;
        padding: 10px;
        background-color: rgba(255, 255, 255, 0.7);
        border-radius: 4px;
        white-space: pre-wrap;
        word-break: break-word;
        border: 1px solid rgba(0,0,0,0.1);
        font-size: 12px;
    }

    .business-summary {
        background-color: #e7f3ff;
        padding: 20px;
        border-radius: 8px;
        margin-bottom: 25px;
        border-left: 4px solid #007bff;
    }

    .business-summary h3 {
        margin-top: 0;
        color: #004085;
        font-size: 18px;
    }

    .info-message {
        padding: 12px 15px;
        margin: 15px 0;
        border-radius: 6px;
        border-left: 4px solid #17a2b8;
        background-color: #d1ecf1;
        color: #0c5460;
    }

    .error-message {
        background-color: #f8d7da;
        border-left-color: #dc3545;
        color: #721c24;
        padding: 12px 15px;
        margin: 15px 0;
        border-radius: 6px;
        border-left: 4px solid;
    }

    /* Responsividade para tabelas */
    @media (max-width: 768px) {
        .fields-comparison-table {
            font-size: 12px;
        }

        .fields-comparison-table th,
        .fields-comparison-table td {
            padding: 8px 10px;
        }

        .field-name-cell {
            font-size: 11px;
        }

        .field-desc-cell {
            font-size: 11px;
        }

        .old-value-cell, .new-value-cell {
            font-size: 11px;
        }
    }
    </style>
    """)

    # Juntar todas as partes HTML
    return ''.join(html_parts)

INDEX_HTML = r"""<!DOCTYPE html>
<html lang="pt-br">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Validador de Arquivos - Versão Melhorada</title>
    <link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.4.0/css/all.min.css">
    <style>
        :root {
            --primary-color: #4CAF50;
            --primary-hover: #43a047;
            --secondary-color: #3498db;
            --light-gray: #f4f4f4;
            --medium-gray: #e0e0e0;
            --dark-gray: #757575;
            --danger-color: #e74c3c;
            --warning-color: #f39c12;
            --success-color: #2ecc71;
            --border-radius: 8px;
            --box-shadow: 0 4px 12px rgba(0, 0, 0, 0.08);
            --transition-speed: 0.3s;
        }

        * {
            box-sizing: border-box;
            margin: 0;
            padding: 0;
        }

        body {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            background-color: #f9f9f9;
            color: #333;
            line-height: 1.6;
            padding: 20px;
        }

        .container {
            margin: 0 auto;
            background: white;
            border-radius: var(--border-radius);
            box-shadow: var(--box-shadow);
            overflow: hidden;
        }

        header {
            background-color: #fff;
            padding: 20px;
            text-align: center;
            border-bottom: 1px solid var(--medium-gray);
        }

        h1 {
            font-weight: 600;
            color: #333;
            margin-bottom: 10px;
            font-size: 28px;
        }

        .tab-container {
            display: flex;
            border-bottom: 1px solid var(--medium-gray);
        }

        .tab {
            flex: 1;
            padding: 15px 0;
            text-align: center;
            font-weight: 600;
            background-color: #f8f8f8;
            cursor: pointer;
            transition: background-color var(--transition-speed);
            border-bottom: 3px solid transparent;
        }

        .tab.active {
            background-color: #fff;
            border-bottom-color: var(--primary-color);
            color: var(--primary-color);
        }

        .tab:hover:not(.active) {
            background-color: #f0f0f0;
        }

        .form-container {
            padding: 30px;
        }

        .tab-content {
            display: none;
        }

        .tab-content.active {
            display: block;
        }

        .section-title {
            font-size: 20px;
            margin-bottom: 20px;
            color: #333;
            font-weight: 600;
        }

        .form-group {
            margin-bottom: 25px;
            position: relative;
        }

        .form-label {
            display: block;
            margin-bottom: 8px;
            font-weight: 500;
            color: #333;
        }

        .tooltip-icon {
            display: inline-block;
            width: 18px;
            height: 18px;
            background-color: var(--dark-gray);
            color: white;
            border-radius: 50%;
            text-align: center;
            font-size: 12px;
            line-height: 18px;
            margin-left: 6px;
            cursor: help;
            position: relative;
        }

        .tooltip-content {
            position: absolute;
            top: calc(100% + 10px);
            left: 50%;
            transform: translateX(-50%);
            width: 250px;
            background-color: #333;
            color: white;
            padding: 10px;
            border-radius: var(--border-radius);
            font-size: 13px;
            opacity: 0;
            pointer-events: none;
            transition: opacity var(--transition-speed);
            z-index: 10;
            box-shadow: 0 2px 10px rgba(0,0,0,0.2);
        }

        .tooltip-content::before {
            content: '';
            position: absolute;
            bottom: 100%;
            left: 50%;
            transform: translateX(-50%);
            border-width: 6px;
            border-style: solid;
            border-color: transparent transparent #333 transparent;
        }

        .tooltip-icon:hover .tooltip-content {
            opacity: 1;
        }

        .file-selector {
            position: relative;
        }

        .file-input {
            width: 100%;
            padding: 12px;
            border: 1px dashed var(--medium-gray);
            border-radius: var(--border-radius);
            cursor: pointer;
            background-color: #f9f9f9;
            color: #555;
            transition: all var(--transition-speed);
            position: relative;
        }

        .file-input:hover {
            background-color: #f0f0f0;
            border-color: var(--dark-gray);
        }

        .file-input input[type="file"] {
            position: absolute;
            width: 100%;
            height: 100%;
            top: 0;
            left: 0;
            opacity: 0;
            cursor: pointer;
        }

        .drop-zone {
            padding: 30px;
            text-align: center;
            border: 2px dashed var(--medium-gray);
            border-radius: var(--border-radius);
            background-color: #f9f9f9;
            color: #555;
            margin-bottom: 20px;
            transition: all var(--transition-speed);
            position: relative;
        }

        .drop-zone.active {
            border-color: var(--primary-color);
            background-color: rgba(76, 175, 80, 0.05);
        }

        .drop-zone i {
            font-size: 34px;
            margin-bottom: 10px;
            color: var(--dark-gray);
        }

        .drop-zone.active i {
            color: var(--primary-color);
        }

        .drop-zone-text {
            font-size: 16px;
            margin-bottom: 10px;
        }

        .drop-zone-hint {
            font-size: 13px;
            color: #888;
        }

        .selected-file {
            display: flex;
            align-items: center;
            background-color: #f9f9f9;
            border: 1px solid var(--medium-gray);
            padding: 8px 12px;
            border-radius: var(--border-radius);
            margin-top: 10px;
            font-size: 14px;
        }

        .selected-file i {
            margin-right: 10px;
            color: var(--dark-gray);
        }

        .remove-file {
            margin-left: auto;
            color: var(--danger-color);
            cursor: pointer;
            font-size: 18px;
        }

        .validation-method {
            display: block;
            width: 100%;
            padding: 12px;
            border: 1px solid var(--medium-gray);
            border-radius: var(--border-radius);
            background-color: white;
            font-size: 16px;
            color: #333;
            cursor: pointer;
        }

        .validation-method:focus {
            border-color: var(--secondary-color);
            outline: none;
        }

        .method-description {
            margin-top: 10px;
            padding: 12px;
            background-color: #f9f9f9;
            border-left: 3px solid var(--primary-color);
            border-radius: 0 var(--border-radius) var(--border-radius) 0;
            font-size: 14px;
            color: #666;
        }

        .advanced-options {
            margin-top: 20px;
            border: 1px solid var(--medium-gray);
            border-radius: var(--border-radius);
            overflow: hidden;
        }

        .advanced-options-header {
            display: flex;
            justify-content: space-between;
            align-items: center;
            padding: 15px;
            background-color: #f9f9f9;
            cursor: pointer;
            border-bottom: 1px solid var(--medium-gray);
        }

        .advanced-options-header h3 {
            font-size: 16px;
            font-weight: 600;
            color: #444;
        }

        .advanced-options-content {
            padding: 15px;
            display: none;
        }

        .advanced-options-content.open {
            display: block;
        }

        .advanced-options-toggle {
            font-size: 16px;
            transition: transform var(--transition-speed);
        }

        .advanced-options.open .advanced-options-toggle {
            transform: rotate(180deg);
        }

        .progress-container {
            height: 8px;
            background-color: var(--medium-gray);
            border-radius: 4px;
            margin-top: 10px;
            overflow: hidden;
            display: none;
        }

        .progress-bar {
            height: 100%;
            background-color: var(--primary-color);
            width: 0;
            transition: width 0.3s;
        }

        .button {
            display: block;
            width: 100%;
            padding: 14px;
            background-color: var(--primary-color);
            color: white;
            border: none;
            border-radius: var(--border-radius);
            font-size: 16px;
            font-weight: 600;
            cursor: pointer;
            transition: background-color var(--transition-speed);
            text-align: center;
        }

        .button:hover {
            background-color: var(--primary-hover);
        }

        .button:focus {
            outline: none;
            box-shadow: 0 0 0 3px rgba(76, 175, 80, 0.3);
        }

        .button i {
            margin-right: 8px;
        }

        .recent-comparison {
            padding: 15px;
            margin-top: 30px;
            border-top: 1px solid var(--medium-gray);
        }

        .recent-comparison h3 {
            margin-bottom: 15px;
            font-size: 18px;
            color: #444;
        }

        .comparison-list {
            max-height: 200px;
            overflow-y: auto;
        }

        .comparison-item {
            display: flex;
            padding: 12px;
            border-bottom: 1px solid var(--medium-gray);
            font-size: 14px;
            cursor: pointer;
            transition: background-color var(--transition-speed);
        }

        .comparison-item:hover {
            background-color: #f5f5f5;
        }

        .comparison-item:last-child {
            border-bottom: none;
        }

        .comparison-files {
            flex: 1;
        }

        .comparison-status {
            margin-left: 15px;
            font-weight: 600;
        }

        .comparison-status.identical {
            color: var(--success-color);
        }

        .comparison-status.different {
            color: var(--danger-color);
        }

        .comparison-timestamp {
            margin-left: 15px;
            color: var(--dark-gray);
        }

        .file-preview {
            position: fixed;
            top: 0;
            left: 0;
            width: 100%;
            height: 100%;
            background-color: rgba(0, 0, 0, 0.8);
            display: flex;
            align-items: center;
            justify-content: center;
            z-index: 1000;
            opacity: 0;
            pointer-events: none;
            transition: opacity var(--transition-speed);
        }

        .file-preview.open {
            opacity: 1;
            pointer-events: auto;
        }

        .preview-content {
            background-color: white;
            border-radius: var(--border-radius);
            max-width: 800px;
            width: 90%;
            max-height: 80vh;
            overflow: auto;
            position: relative;
        }

        .preview-header {
            display: flex;
            justify-content: space-between;
            align-items: center;
            padding: 15px;
            border-bottom: 1px solid var(--medium-gray);
        }

        .preview-title {
            font-size: 18px;
            font-weight: 600;
        }

        .preview-close {
            font-size: 20px;
            cursor: pointer;
            color: var(--dark-gray);
            transition: color var(--transition-speed);
        }

        .preview-close:hover {
            color: var(--danger-color);
        }

        .preview-body {
            padding: 20px;
            font-family: 'Courier New', monospace;
            font-size: 14px;
            line-height: 1.5;
            white-space: pre-wrap;
            overflow-x: auto;
        }

        /* Responsividade */
        @media (max-width: 768px) {
            .container {
                margin: 0;
            }

            .form-container {
                padding: 20px;
            }
        }

        /* Loading Overlay */
        #loading-overlay {
            position: fixed;
            top: 0;
            left: 0;
            width: 100%;
            height: 100%;
            background-color: rgba(255, 255, 255, 0.9);
            z-index: 2000;
            display: none;
            justify-content: center;
            align-items: center;
            flex-direction: column;
        }

        .spinner {
            width: 50px;
            height: 50px;
            border: 5px solid var(--medium-gray);
            border-radius: 50%;
            border-top-color: var(--primary-color);
            animation: spin 1s linear infinite;
        }

        .loading-text {
            margin-top: 15px;
            font-size: 16px;
            color: #444;
        }

        @keyframes spin {
            100% {
                transform: rotate(360deg);
            }
        }
    </style>
</head>
<body>
    <div class="container">
        <header>
            <h1>Validador de Arquivos</h1>
        </header>

        <div class="tab-container">
            <div class="tab active" data-tab="individual">Modo Individual</div>
            <div class="tab" data-tab="batch">Modo Lote</div>
            <div class="tab" data-tab="layouts">Gerenciar Layouts</div>
            <div class="tab" data-tab="quickcompare">Validação Rápida</div>
        </div>

        <div class="form-container">
            <!-- Modo Individual -->
            <div class="tab-content active" id="individual-tab">
                <h2 class="section-title">Comparar dois arquivos</h2>

                <div class="drop-zone" id="individual-drop-zone">
                    <i class="fas fa-cloud-upload-alt"></i>
                    <div class="drop-zone-text">Arraste e solte seus arquivos aqui</div>
                    <div class="drop-zone-hint">ou clique para selecionar arquivos</div>
                </div>

                <form action="/upload" method="post" enctype="multipart/form-data" id="individual-form">
                    <div class="form-group">
                        <label for="file1" class="form-label">
                            Arquivo 1
                            <span class="tooltip-icon">?
                                <div class="tooltip-content">
                                    Selecione o primeiro arquivo para comparação. Arquivos de até 2GB são suportados.
                                </div>
                            </span>
                        </label>
                        <div class="file-selector">
                            <div class="file-input">
                                <input type="file" id="file1" name="file1" required>
                                <span id="file1-placeholder">Selecionar arquivo...</span>
                            </div>
                            <div class="progress-container" id="progress-file1">
                                <div class="progress-bar"></div>
                            </div>
                        </div>
                        <div id="file1-preview-container"></div>
                        <button type="button" class="preview-button" id="preview-file1" style="display:none;">Pré-visualizar</button>
                    </div>

                    <div class="form-group">
                        <label for="file2" class="form-label">
                            Arquivo 2
                            <span class="tooltip-icon">?
                                <div class="tooltip-content">
                                    Selecione o segundo arquivo para comparação com o primeiro.
                                </div>
                            </span>
                        </label>
                        <div class="file-selector">
                            <div class="file-input">
                                <input type="file" id="file2" name="file2" required>
                                <span id="file2-placeholder">Selecionar arquivo...</span>
                            </div>
                            <div class="progress-container" id="progress-file2">
                                <div class="progress-bar"></div>
                            </div>
                        </div>
                        <div id="file2-preview-container"></div>
                        <button type="button" class="preview-button" id="preview-file2" style="display:none;">Pré-visualizar</button>
                    </div>

                    <div class="form-group">
                        <label for="method" class="form-label">
                            Método de Validação
                            <span class="tooltip-icon">?
                                <div class="tooltip-content">
                                    Escolha o método de comparação mais adequado para seus arquivos. O método será sugerido automaticamente com base no tipo de arquivo.
                                </div>
                            </span>
                        </label>
                        <select id="method" name="method" class="validation-method">
                            <option value="hash">Hash (rápido)</option>
                            <option value="content">Conteúdo (para arquivos de texto)</option>
                            <option value="binary">Binário (para qualquer tipo de arquivo)</option>
                        </select>
                        <div id="method-info" class="method-description">
                            Método de hash: Rápido e eficiente. Apenas informa se os arquivos são idênticos ou diferentes.
                        </div>
                    </div>

                    <div class="advanced-options">
                        <div class="advanced-options-header" id="advanced-toggle">
                            <h3>Opções Avançadas</h3>
                            <span class="advanced-options-toggle">
                                <i class="fas fa-chevron-down"></i>
                            </span>
                        </div>
                        <div class="advanced-options-content" id="advanced-content">
                            <div class="form-group">
                                <label for="ignore-whitespace" class="form-label">
                                    <input type="checkbox" id="ignore-whitespace" name="ignore_whitespace"> 
                                    Ignorar espaços em branco
                                </label>
                            </div>
                            <div class="form-group">
                                <label for="ignore-case" class="form-label">
                                    <input type="checkbox" id="ignore-case" name="ignore_case"> 
                                    Ignorar diferenças entre maiúsculas e minúsculas
                                </label>
                            </div>
                            <div class="form-group">
                                <label for="context-lines" class="form-label">
                                    Linhas de contexto
                                    <span class="tooltip-icon">?
                                        <div class="tooltip-content">
                                            Número de linhas mostradas antes e depois de cada diferença encontrada.
                                        </div>
                                    </span>
                                </label>
                                <input type="number" id="context-lines" name="context_lines" min="0" max="10" value="3" style="width: 100%; padding: 10px; border: 1px solid var(--medium-gray); border-radius: var(--border-radius);">
                            </div>
                        </div>
                    </div>

                    <div class="form-group" style="margin-top: 25px;">
                        <button type="submit" class="button">
                            <i class="fas fa-sync-alt"></i> Comparar Arquivos
                        </button>
                    </div>
                </form>

                

            <!-- Modo Lote -->
            <div class="tab-content" id="batch-tab">
                <h2 class="section-title">Comparar lotes de arquivos</h2>
                <p style="margin-bottom: 20px;">Envie dois arquivos ZIP contendo os arquivos a serem comparados. Os arquivos serão combinados por nome.</p>

                <div class="drop-zone" id="batch-drop-zone">
                    <i class="fas fa-cloud-upload-alt"></i>
                    <div class="drop-zone-text">Arraste e solte seus arquivos ZIP aqui</div>
                    <div class="drop-zone-hint">ou clique para selecionar arquivos</div>
                </div>

                <form action="/batch-upload" method="post" enctype="multipart/form-data" id="batch-form">
                    <div class="form-group">
                        <label for="source1" class="form-label">
                            Fonte 1 (ZIP)
                            <span class="tooltip-icon">?
                                <div class="tooltip-content">
                                    Selecione o primeiro arquivo ZIP contendo os arquivos a serem comparados.
                                </div>
                            </span>
                        </label>
                        <div class="file-selector">
                            <div class="file-input">
                                <input type="file" id="source1" name="source1" accept=".zip" required>
                                <span id="source1-placeholder">Selecionar arquivo ZIP...</span>
                            </div>
                            <div class="progress-container" id="progress-source1">
                                <div class="progress-bar"></div>
                            </div>
                        </div>
                        <div id="source1-preview-container"></div>
                    </div>

                    <div class="form-group">
                        <label for="source2" class="form-label">
                            Fonte 2 (ZIP)
                            <span class="tooltip-icon">?
                                <div class="tooltip-content">
                                    Selecione o segundo arquivo ZIP para comparação com o primeiro.
                                </div>
                            </span>
                        </label>
                        <div class="file-selector">
                            <div class="file-input">
                                <input type="file" id="source2" name="source2" accept=".zip" required>
                                <span id="source2-placeholder">Selecionar arquivo ZIP...</span>
                            </div>
                            <div class="progress-container" id="progress-source2">
                                <div class="progress-bar"></div>
                            </div>
                        </div>
                        <div id="source2-preview-container"></div>
                        <div class="form-group">
    <label for="batch-layout" class="form-label">
        Layout (opcional)
        <span class="tooltip-icon">?
            <div class="tooltip-content">
                Selecione um layout previamente importado para usar na comparação em lote.
            </div>
        </span>
    </label>
    <select id="batch-layout" name="layout_name" class="validation-method">
        <option value=""></option>
        <!-- Opções de layout serão inseridas dinamicamente pelo Python -->
    </select>
</div>
                    </div>

                    <div class="form-group">

                        <label for="batch-method" class="form-label">
                            Método de Validação
                            <span class="tooltip-icon">?
                                <div class="tooltip-content">
                                    Escolha o método de comparação para todos os arquivos no lote.
                                </div>
                            </span>
                        </label>
                        <select id="batch-method" name="method" class="validation-method">
                            <option value="hash">Hash (rápido)</option>
                            <option value="content">Conteúdo (para arquivos de texto)</option>
                            <option value="binary">Binário (para qualquer tipo de arquivo)</option>
                        </select>
                        <div id="batch-method-info" class="method-description">
                            Método de hash: Rápido e eficiente. Apenas informa se os arquivos são idênticos ou diferentes.
                        </div>
                    </div>

                    <div class="advanced-options">
                        <div class="advanced-options-header" id="batch-advanced-toggle">
                            <h3>Opções Avançadas</h3>
                            <span class="advanced-options-toggle">
                                <i class="fas fa-chevron-down"></i>
                            </span>
                        </div>
                        <div class="advanced-options-content" id="batch-advanced-content">
                            <div class="form-group">
                                <label for="file-extensions" class="form-label">
                                    Extensões de arquivo (separadas por vírgula)
                                    <span class="tooltip-icon">?
                                        <div class="tooltip-content">
                                            Inclua apenas os arquivos com estas extensões (ex: txt,csv,pdf)
                                        </div>
                                    </span>
                                </label>
                                <input type="text" id="file-extensions" name="file_extensions" placeholder="Exemplo: txt,pdf,docx" style="width: 100%; padding: 10px; border: 1px solid var(--medium-gray); border-radius: var(--border-radius);">
                            </div>

                            <div class="form-group">
                                <label for="name-pattern" class="form-label">
                                    Padrão de nome
                                    <span class="tooltip-icon">?
                                        <div class="tooltip-content">
                                            Use * e ? como curingas. Por exemplo: *.txt ou dados_*.csv
                                        </div>
                                    </span>
                                </label>
                                <input type="text" id="name-pattern" name="name_pattern" placeholder="Exemplo: *.txt" style="width: 100%; padding: 10px; border: 1px solid var(--medium-gray); border-radius: var(--border-radius);">
                            </div>

                            <div class="form-group">
                                <label for="exclude-pattern" class="form-label">
                                    Padrões a excluir (separados por vírgula)
                                </label>
                                <input type="text" id="exclude-pattern" name="exclude_pattern" placeholder="Exemplo: temp_*,~*" style="width: 100%; padding: 10px; border: 1px solid var(--medium-gray); border-radius: var(--border-radius);">
                            </div>

                            <div class="form-group">
                                <label for="ignore-batch-whitespace" class="form-label">
                                    <input type="checkbox" id="ignore-batch-whitespace" name="ignore_whitespace"> 
                                    Ignorar espaços em branco
                                </label>
                            </div>

                            <div class="form-group">
                                <label for="ignore-batch-case" class="form-label">
                                    <input type="checkbox" id="ignore-batch-case" name="ignore_case"> 
                                    Ignorar diferenças entre maiúsculas e minúsculas
                                </label>
                            </div>
                        </div>
                    </div>

                    <div class="form-group" style="margin-top: 25px;">
                        <button type="submit" class="button">
                            <i class="fas fa-sync-alt"></i> Comparar Lotes
                        </button>
                    </div>
                </form>

              
            </div>
        </div>
    </div>

    <!-- Pré-visualização de Arquivo -->
    <div class="file-preview" id="file-preview">
        <div class="preview-content">
            <div class="preview-header">
                <div class="preview-title" id="preview-title">Pré-visualização do Arquivo</div>
                <div class="preview-close" id="preview-close">
                    <i class="fas fa-times"></i>
                </div>
            </div>
            <div class="preview-body" id="preview-body">
                <!-- O conteúdo será inserido via JavaScript -->
            </div>
        </div>
    </div>

    <!-- Loading Overlay -->
    <div id="loading-overlay">
        <div class="spinner"></div>
        <div class="loading-text">Processando arquivos...</div>
    </div>

    <script>
        document.addEventListener('DOMContentLoaded', function() {
         document.querySelector('[data-tab="layouts"]').addEventListener('click', function() {
                window.location.href = '/layouts';
            });
            document.querySelector('[data-tab="quickcompare"]').addEventListener('click', function() {
                window.location.href = '/upload-and-compare';
            });
            // Trocar entre abas
            const tabs = document.querySelectorAll('.tab');
            tabs.forEach(tab => {
                tab.addEventListener('click', function() {
                    const tabId = this.getAttribute('data-tab');

                    // Atualizar abas ativas
                    document.querySelectorAll('.tab').forEach(t => t.classList.remove('active'));
                    document.querySelectorAll('.tab-content').forEach(c => c.classList.remove('active'));

                    this.classList.add('active');
                    document.getElementById(tabId + '-tab').classList.add('active');
                });
            });

            // Atualizar descrição do método conforme seleção
            const methodSelect = document.getElementById('method');
            const methodInfo = document.getElementById('method-info');
            const batchMethodSelect = document.getElementById('batch-method');
            const batchMethodInfo = document.getElementById('batch-method-info');

            function updateMethodDescription(select, infoElement) {
                const descriptions = {
                    'hash': 'Método de hash: Rápido e eficiente. Apenas informa se os arquivos são idênticos ou diferentes.',
                    'content': 'Método de conteúdo: Compara os arquivos linha por linha. Ideal para arquivos de texto.',
                    'binary': 'Método binário: Compara os arquivos byte a byte. Funciona para qualquer tipo de arquivo.'
                };
                infoElement.textContent = descriptions[select.value];
            }

            methodSelect.addEventListener('change', function() {
                updateMethodDescription(this, methodInfo);
            });

            batchMethodSelect.addEventListener('change', function() {
                updateMethodDescription(this, batchMethodInfo);
            });

            // Toggle opções avançadas
            document.getElementById('advanced-toggle').addEventListener('click', function() {
                const content = document.getElementById('advanced-content');
                content.classList.toggle('open');
                this.parentElement.classList.toggle('open');
            });

            document.getElementById('batch-advanced-toggle').addEventListener('click', function() {
                const content = document.getElementById('batch-advanced-content');
                content.classList.toggle('open');
                this.parentElement.classList.toggle('open');
            });

            // Manipulação de input de arquivo
            function handleFileSelect(fileInput, placeholderElement, progressContainerId, previewButtonId, previewContainerId) {
                fileInput.addEventListener('change', function() {
                    const file = this.files[0];
                    if (file) {
                        // Atualizar o texto placeholder
                        placeholderElement.textContent = file.name;

                        // Mostrar o contêiner de progresso para arquivos grandes
                        if (file.size > 10 * 1024 * 1024) { // 10MB
                            const progressContainer = document.getElementById(progressContainerId);
                            progressContainer.style.display = 'block';
                            simulateFileUploadProgress(progressContainerId);
                        }

                        // Verificar se o arquivo é visualizável
                        const isTextFile = /\.(txt|html|css|js|json|xml|md|csv|log)$/i.test(file.name);
                        if (isTextFile) {
                            document.getElementById(previewButtonId).style.display = 'inline-block';

                            // Adicionar container de arquivo selecionado
                            const container = document.getElementById(previewContainerId);
                            container.innerHTML = `
                                <div class="selected-file">
                                    <i class="fas fa-file-alt"></i>
                                    <span>${file.name} (${formatFileSize(file.size)})</span>
                                    <i class="fas fa-times remove-file" data-input="${fileInput.id}"></i>
                                </div>
                            `;
                        }

                        // Autodetectar método de validação recomendado
                        suggestValidationMethod(file);
                    }
                });
            }

            // Configurar seleção de arquivos
            const file1Input = document.getElementById('file1');
            const file2Input = document.getElementById('file2');
            const source1Input = document.getElementById('source1');
            const source2Input = document.getElementById('source2');

            handleFileSelect(file1Input, document.getElementById('file1-placeholder'), 'progress-file1', 'preview-file1', 'file1-preview-container');
            handleFileSelect(file2Input, document.getElementById('file2-placeholder'), 'progress-file2', 'preview-file2', 'file2-preview-container');
            handleFileSelect(source1Input, document.getElementById('source1-placeholder'), 'progress-source1', '', 'source1-preview-container');
            handleFileSelect(source2Input, document.getElementById('source2-placeholder'), 'progress-source2', '', 'source2-preview-container');

            // Remover arquivos selecionados
            document.addEventListener('click', function(e) {
                if (e.target.classList.contains('remove-file')) {
                    const inputId = e.target.getAttribute('data-input');
                    const input = document.getElementById(inputId);
                    input.value = ''; // Limpar o input de arquivo

                    // Resetar o texto placeholder
                    const placeholderId = inputId + '-placeholder';
                    document.getElementById(placeholderId).textContent = 'Selecionar arquivo...';

                    // Remover o contêiner de arquivo selecionado
                    e.target.closest('.selected-file').remove();

                    // Esconder o botão de pré-visualização
                    const previewButtonId = 'preview-' + inputId;
                    const previewButton = document.getElementById(previewButtonId);
                    if (previewButton) {
                        previewButton.style.display = 'none';
                    }
                }
            });

            // Função para sugerir método de validação com base no tipo de arquivo
            function suggestValidationMethod(file) {
                const textExtensions = ['.txt', '.html', '.css', '.js', '.json', '.xml', '.md', '.csv', '.log'];
                const binaryExtensions = ['.pdf', '.doc', '.docx', '.xls', '.xlsx', '.zip', '.rar', '.jpg', '.png', '.gif'];

                const extension = '.' + file.name.split('.').pop().toLowerCase();

                // Selecionar o método apropriado para o tipo de arquivo
                let recommendedMethod = 'hash'; // Método padrão

                if (textExtensions.includes(extension)) {
                    recommendedMethod = 'content';
                } else if (binaryExtensions.includes(extension)) {
                    recommendedMethod = 'binary';
                }

                // Atualizar o select com o método recomendado
                const tabId = document.querySelector('.tab.active').getAttribute('data-tab');
                const methodSelect = tabId === 'individual' 
                    ? document.getElementById('method') 
                    : document.getElementById('batch-method');

                methodSelect.value = recommendedMethod;

                // Atualizar a descrição
                const methodInfo = tabId === 'individual' 
                    ? document.getElementById('method-info') 
                    : document.getElementById('batch-method-info');

                updateMethodDescription(methodSelect, methodInfo);

                // Notificar o usuário sobre a seleção automática
                const methodMessage = `Método ${recommendedMethod} selecionado automaticamente com base no tipo de arquivo.`;

                // Adicionar temporariamente mensagem de notificação
                const notificationDiv = document.createElement('div');
                notificationDiv.textContent = methodMessage;
                notificationDiv.style.color = '#3498db';
                notificationDiv.style.fontSize = '14px';
                notificationDiv.style.marginTop = '5px';

                methodInfo.parentNode.appendChild(notificationDiv);

                // Remover a mensagem após 3 segundos
                setTimeout(() => {
                    notificationDiv.remove();
                }, 3000);
            }

            // Simulação de upload para arquivos grandes
            function simulateFileUploadProgress(progressContainerId) {
                const progressBar = document.querySelector(`#${progressContainerId} .progress-bar`);
                let width = 0;

                const interval = setInterval(() => {
                    if (width >= 100) {
                        clearInterval(interval);
                        // Esconder a barra depois de completada
                        setTimeout(() => {
                            document.getElementById(progressContainerId).style.display = 'none';
                        }, 1000);
                    } else {
                        width += Math.random() * 10;
                        if (width > 100) width = 100;
                        progressBar.style.width = width + '%';
                    }
                }, 200);
            }

            // Formatar tamanho de arquivo para exibição
            function formatFileSize(bytes) {
                if (bytes === 0) return '0 Bytes';
                const k = 1024;
                const sizes = ['Bytes', 'KB', 'MB', 'GB', 'TB'];
                const i = Math.floor(Math.log(bytes) / Math.log(k));
                return parseFloat((bytes / Math.pow(k, i)).toFixed(2)) + ' ' + sizes[i];
            }

            // Pré-visualização de arquivo
            const filePreview = document.getElementById('file-preview');
            const previewClose = document.getElementById('preview-close');
            const previewTitle = document.getElementById('preview-title');
            const previewBody = document.getElementById('preview-body');

            // Fechar pré-visualização
            previewClose.addEventListener('click', function() {
                filePreview.classList.remove('open');
            });

            // Também fechar ao clicar fora da área de conteúdo
            filePreview.addEventListener('click', function(e) {
                if (e.target === filePreview) {
                    filePreview.classList.remove('open');
                }
            });

            // Configurar botões de pré-visualização
            document.getElementById('preview-file1').addEventListener('click', function() {
                showFilePreview(file1Input.files[0]);
            });

            document.getElementById('preview-file2').addEventListener('click', function() {
                showFilePreview(file2Input.files[0]);
            });

            // Função para mostrar a pré-visualização do arquivo
            function showFilePreview(file) {
                if (!file) return;

                previewTitle.textContent = 'Pré-visualização: ' + file.name;
                previewBody.textContent = 'Carregando...';
                filePreview.classList.add('open');

                // Ler o conteúdo do arquivo
                const reader = new FileReader();

                reader.onload = function(e) {
                    let content = e.target.result;

                    // Limitar o tamanho para arquivos muito grandes
                    const maxLength = 100000; // ~100KB
                    if (content.length > maxLength) {
                        content = content.substring(0, maxLength) + '\n\n[Arquivo truncado para visualização. O arquivo completo será processado na comparação.]';
                    }

                    previewBody.textContent = content;
                };

                reader.onerror = function() {
                    previewBody.textContent = 'Erro ao ler o arquivo. Pode ser que o arquivo seja muito grande ou não seja um arquivo de texto.';
                };

                reader.readAsText(file);
            }

            // Implementação de drag and drop para upload
            setupDragAndDrop('individual-drop-zone', [file1Input, file2Input]);
            setupDragAndDrop('batch-drop-zone', [source1Input, source2Input]);

            function setupDragAndDrop(zoneId, inputs) {
                const dropZone = document.getElementById(zoneId);

                ['dragenter', 'dragover', 'dragleave', 'drop'].forEach(eventName => {
                    dropZone.addEventListener(eventName, preventDefaults, false);
                });

                function preventDefaults(e) {
                    e.preventDefault();
                    e.stopPropagation();
                }

                ['dragenter', 'dragover'].forEach(eventName => {
                    dropZone.addEventListener(eventName, highlight, false);
                });

                ['dragleave', 'drop'].forEach(eventName => {
                    dropZone.addEventListener(eventName, unhighlight, false);
                });

                function highlight() {
                    dropZone.classList.add('active');
                }

                function unhighlight() {
                    dropZone.classList.remove('active');
                }

                // Área de clique também funcionará para selecionar arquivos
                dropZone.addEventListener('click', () => {
                    inputs[0].click();
                });

                // Processar arquivos dropados
                dropZone.addEventListener('drop', function(e) {
                    const dt = e.dataTransfer;
                    const files = dt.files;

                    if (files.length === 0) return;

                    // Se for só um arquivo, usamos o primeiro input
                    if (files.length === 1) {
                        inputs[0].files = files;
                        inputs[0].dispatchEvent(new Event('change'));
                    } 
                    // Se houver dois ou mais arquivos, usamos os dois inputs
                    else if (files.length >= 2) {
                        const dataTransfer1 = new DataTransfer();
                        dataTransfer1.items.add(files[0]);
                        inputs[0].files = dataTransfer1.files;
                        inputs[0].dispatchEvent(new Event('change'));

                        const dataTransfer2 = new DataTransfer();
                        dataTransfer2.items.add(files[1]);
                        inputs[1].files = dataTransfer2.files;
                        inputs[1].dispatchEvent(new Event('change'));
                    }
                });
            }

            

            // Adicionar evento de submit para o formulário com efeito de loading
            document.getElementById('individual-form').addEventListener('submit', showLoading);
            document.getElementById('batch-form').addEventListener('submit', showLoading);

            function showLoading() {
                document.getElementById('loading-overlay').style.display = 'flex';

                // Simular atalhos de teclado
                setupKeyboardShortcuts();
            }

            // Implementar atalhos de teclado
            function setupKeyboardShortcuts() {
                // Adicionar área de ajuda para atalhos
                const shortcutsHelp = document.createElement('div');
                shortcutsHelp.style.position = 'fixed';
                shortcutsHelp.style.bottom = '10px';
                shortcutsHelp.style.right = '10px';
                shortcutsHelp.style.background = '#333';
                shortcutsHelp.style.color = 'white';
                shortcutsHelp.style.padding = '10px';
                shortcutsHelp.style.borderRadius = '5px';
                shortcutsHelp.style.fontSize = '12px';
                shortcutsHelp.style.opacity = '0';
                shortcutsHelp.style.transition = 'opacity 0.3s';
                shortcutsHelp.innerHTML = '<strong>Atalhos de Teclado:</strong><br>' +
                    'Alt+1: Modo Individual<br>' +
                    'Alt+2: Modo Lote<br>' +
                    'Alt+C: Comparar<br>' +
                    'Alt+H: Mostrar/ocultar essa ajuda<br>';
                document.body.appendChild(shortcutsHelp);

                document.addEventListener('keydown', function(e) {
                    // Alt+1: Mudar para modo individual
                    if (e.altKey && e.key === '1') {
                        document.querySelector('[data-tab="individual"]').click();
                    }

                    // Alt+2: Mudar para modo lote
                    if (e.altKey && e.key === '2') {
                        document.querySelector('[data-tab="batch"]').click();
                    }

                    // Alt+C: Enviar formulário atual
                    if (e.altKey && e.key === 'c') {
                        const activeTab = document.querySelector('.tab-content.active');
                        if (activeTab.id === 'individual-tab') {
                            document.getElementById('individual-form').submit();
                        } else {
                            document.getElementById('batch-form').submit();
                        }
                    }

                    // Alt+H: Mostrar/ocultar ajuda de atalhos
                    if (e.altKey && e.key === 'h') {
                        shortcutsHelp.style.opacity = shortcutsHelp.style.opacity === '1' ? '0' : '1';
                    }
                });
            }
        });
    </script>
</body>
</html>
"""

# Modificação no BATCH_RESULT_HTML_TEMPLATE para incluir a visão amigável
BATCH_RESULT_HTML_TEMPLATE = """<!DOCTYPE html>
<html lang="pt-br">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Resultado da Validação em Lote</title>
    <style>
        body {
            font-family: Arial, sans-serif;
            line-height: 1.6;
            margin: 0;
            padding: 20px;
            background-color: #f4f4f4;
        }
        .container {
            max-width: 1100px;
            margin: 0 auto;
            background: white;
            padding: 20px;
            border-radius: 5px;
            box-shadow: 0 0 10px rgba(0, 0, 0, 0.1);
        }
        h1, h2, h3, h4 {
            color: #333;
        }
        .source-info {
            margin-bottom: 20px;
            padding: 15px;
            background-color: #e9ecef;
            border-radius: 5px;
        }
        .source-name {
            font-weight: bold;
        }
        .filter-info {
            margin-bottom: 20px;
            padding: 15px;
            background-color: #f0f7fb;
            border-radius: 5px;
            border-left: 5px solid #3498db;
        }
        .filter-info ul {
            margin: 10px 0 0 0;
            padding-left: 20px;
        }
        .filter-info h4 {
            margin: 0;
            color: #3498db;
        }
        .stats-container {
            margin-top: 20px;
            padding: 15px;
            background-color: #e9ecef;
            border-radius: 5px;
            display: flex;
            flex-wrap: wrap;
            justify-content: space-between;
        }
        .stats-item {
            margin: 5px 10px;
            flex-grow: 1;
            text-align: center;
        }
        .stats-highlighted {
            font-weight: bold;
            font-size: 24px;
            color: #28a745;
            display: block;
        }
        .stats-warning {
            font-weight: bold;
            font-size: 24px;
            color: #dc3545;
            display: block;
        }
        .stats-label {
            font-size: 14px;
            color: #666;
        }
        table {
            width: 100%;
            border-collapse: collapse;
            margin-top: 20px;
            box-shadow: 0 2px 3px rgba(0,0,0,0.1);
        }
        th, td {
            padding: 12px 15px;
            text-align: left;
            border-bottom: 1px solid #ddd;
        }
        th {
            background-color: #f2f2f2;
            font-weight: bold;
            position: sticky;
            top: 0;
            z-index: 10;
        }
        tr:hover {
            background-color: #f5f5f5;
        }
        .status-identical {
            background-color: #dff0d8;
            padding: 5px 10px;
            border-radius: 4px;
            color: #3c763d;
            font-weight: bold;
        }
        .status-different {
            background-color: #f2dede;
            padding: 5px 10px;
            border-radius: 4px;
            color: #a94442;
            font-weight: bold;
        }
        .status-missing {
            background-color: #fcf8e3;
            padding: 5px 10px;
            border-radius: 4px;
            color: #8a6d3b;
            font-weight: bold;
        }
        .controls-container {
            display: flex;
            justify-content: space-between;
            align-items: center;
            margin: 20px 0;
            padding: 15px;
            background-color: #f8f9fa;
            border-radius: 5px;
        }
        .filter-container {
            flex-grow: 1;
        }
        .filter-container select {
            padding: 8px;
            border: 1px solid #ddd;
            border-radius: 4px;
            background-color: #fff;
        }
        .export-container {
            display: flex;
            gap: 10px;
        }
        .export-btn {
            display: inline-block;
            padding: 8px 15px;
            background: #28a745;
            color: white;
            text-decoration: none;
            border-radius: 4px;
            font-weight: bold;
            transition: background-color 0.3s;
        }
        .export-btn:hover {
            background: #218838;
        }
        .details-btn {
            padding: 5px 10px;
            background: #337ab7;
            color: white;
            border: none;
            border-radius: 4px;
            cursor: pointer;
            font-size: 14px;
            transition: background-color 0.3s;
        }
        .details-btn:hover {
            background: #286090;
        }
        .modal {
            display: none;
            position: fixed;
            z-index: 100;
            left: 0;
            top: 0;
            width: 100%;
            height: 100%;
            overflow: auto;
            background-color: rgba(0,0,0,0.4);
        }
        .modal-content {
            background-color: #fefefe;
            margin: 5% auto;
            padding: 20px;
            border: 1px solid #888;
            width: 90%;
            max-width: 1000px;
            border-radius: 5px;
            max-height: 80vh;
            overflow: auto;
        }
        .close {
            color: #aaa;
            float: right;
            font-size: 28px;
            font-weight: bold;
            cursor: pointer;
        }
        .close:hover,
        .close:focus {
            color: black;
            text-decoration: none;
        }
        .file-info-section {
            display: flex;
            justify-content: space-between;
            margin-bottom: 15px;
            padding: 10px;
            background-color: #f8f9fa;
            border-radius: 5px;
        }
        .file-info-item {
            flex-grow: 1;
        }
        .file-info-label {
            font-weight: bold;
            display: block;
        }
        .diff-controls {
            display: flex;
            justify-content: space-between;
            margin-bottom: 15px;
            padding: 10px;
            background-color: #f0f0f0;
            border-radius: 4px;
        }
        .diff-view-btn {
            background-color: #337ab7;
            color: white;
            border: none;
            padding: 5px 10px;
            border-radius: 4px;
            cursor: pointer;
        }
        .diff-view-btn:hover {
            background-color: #286090;
        }
        .diff-view-btn.active {
            background-color: #1e486b;
        }
        .diff-container {
            background-color: #f8f8f8;
            padding: 15px;
            border-radius: 5px;
            overflow-x: auto;
            max-height: 500px;
            border: 1px solid #ddd;
        }
        .advanced-diff-container {
            display: flex;
            margin-top: 20px;
            border: 1px solid #ddd;
            border-radius: 4px;
            overflow: hidden;
        }
        .file-column {
            flex: 1;
            padding: 10px;
            max-height: 500px;
            overflow: auto;
        }
        .file1-column {
            background-color: #f8f9fa;
            border-right: 1px solid #ddd;
        }
        .file2-column {
            background-color: #f8f9fa;
        }
        .line-number {
            color: #6c757d;
            user-select: none;
            text-align: right;
            padding-right: 10px;
            width: 40px;
            display: inline-block;
            font-family: monospace;
            border-right: 1px solid #ddd;
            margin-right: 10px;
        }
        .code-line {
            font-family: 'Courier New', monospace;
            white-space: pre;
            padding: 2px 0;
            line-height: 1.4;
        }
        .diff-line-added {
            background-color: #e6ffed;
        }
        .diff-line-removed {
            background-color: #ffeef0;
        }
        .diff-line-unchanged {
            background-color: transparent;
        }
        .diff-line {
            font-family: 'Courier New', monospace;
            white-space: pre;
            margin: 0;
            padding: 1px 5px;
        }
        .diff-added {
            background-color: #e6ffed;
            color: #22863a;
        }
        .diff-removed {
            background-color: #ffeef0;
            color: #cb2431;
        }
        .binary-diff {
            font-family: 'Courier New', monospace;
            margin: 2px 0;
        }
        .diff-header {
            color: #6c757d;
            padding: 5px;
            margin-top: 10px;
        }
        .diff-context {
            color: #6c757d;
        }
        .btn-home {
            display: block;
            width: 200px;
            margin: 20px auto;
            text-align: center;
            text-decoration: none;
        }
        .empty-message {
            text-align: center;
            padding: 40px;
            color: #666;
            font-style: italic;
            background-color: #f9f9f9;
            border-radius: 5px;
            margin: 20px 0;
        }
        .search-box {
            margin-left: 10px;
            padding: 8px;
            border: 1px solid #ddd;
            border-radius: 4px;
            width: 200px;
        }

        /* Estilos para a visualização de negócio */
        .view-controls {
            display: flex;
            justify-content: center;
            margin: 20px 0;
        }

        .view-button {
            padding: 8px 12px;
            margin: 0 5px;
            background-color: #f0f0f0;
            border: 1px solid #ddd;
            border-radius: 4px;
            cursor: pointer;
        }

        .view-button.active {
            background-color: #4CAF50;
            color: white;
            border-color: #45a049;
        }

        .view-content {
            display: none;
        }

        .view-content.active {
            display: block;
        }

        .difference-summary {
            background-color: #f5f5f5;
            padding: 15px;
            border-radius: 8px;
            margin-bottom: 20px;
            border-left: 4px solid #3498db;
        }

        .change-item {
            margin-bottom: 20px;
            border: 1px solid #ddd;
            border-radius: 8px;
            overflow: hidden;
        }

        .change-header {
            background-color: #f0f0f0;
            padding: 10px;
            font-weight: bold;
            border-bottom: 1px solid #ddd;
        }

        .change-content {
            display: flex;
            flex-direction: column;
        }

        @media (min-width: 768px) {
            .change-content {
                flex-direction: row;
            }
        }

        .before-change, .after-change {
            padding: 15px;
            flex: 1;
        }

        .before-change {
            background-color: #ffeeee;
            border-right: 1px solid #ddd;
        }

        .after-change {
            background-color: #eeffee;
        }

        .code-block {
            font-family: monospace;
            padding: 10px;
            background-color: rgba(255, 255, 255, 0.5);
            border-radius: 4px;
            white-space: pre-wrap;
            word-break: break-all;
        }

        .change-explanation {
            padding: 10px;
            background-color: #fffde7;
            border-bottom: 1px solid #ddd;
            font-style: italic;
        }

        /* Estilos para destacar alterações de caracteres */
        .highlight-remove {
            background-color: #ffb6c1;
            font-weight: bold;
            padding: 0 2px;
            border-radius: 3px;
            text-decoration: line-through;
        }

        .highlight-add {
            background-color: #90ee90;
            font-weight: bold;
            padding: 0 2px;
            border-radius: 3px;
        }

        /* Estilos para o resumo de negócios */
        .business-summary {
            background-color: #e9f7fe;
            padding: 15px;
            border-radius: 8px;
            margin-bottom: 20px;
            border-left: 4px solid #3498db;
        }

        .business-stats {
            display: flex;
            flex-wrap: wrap;
            justify-content: space-between;
            margin-bottom: 20px;
        }

        .business-stat-item {
            flex: 1;
            min-width: 200px;
            margin: 10px;
            padding: 15px;
            background-color: #fff;
            border-radius: 8px;
            box-shadow: 0 2px 5px rgba(0,0,0,0.1);
            text-align: center;
        }

        .business-stat-number {
            font-size: 24px;
            font-weight: bold;
            color: #3498db;
            margin-bottom: 5px;
        }

        .business-stat-label {
            color: #666;
        }

        .business-file-list {
            background-color: #fff;
            border-radius: 8px;
            box-shadow: 0 2px 5px rgba(0,0,0,0.1);
            overflow: hidden;
        }

        .business-file-row {
            display: flex;
            padding: 12px 15px;
            border-bottom: 1px solid #eee;
        }

        .business-file-row:last-child {
            border-bottom: none;
        }

        .business-file-row:hover {
            background-color: #f9f9f9;
        }

        .business-file-name {
            flex: 2;
            font-weight: bold;
        }

        .business-file-status {
            flex: 1;
            text-align: center;
        }

        .business-file-action {
            flex: 1;
            text-align: right;
        }
    </style>
</head>
<body>
    <div class="container">
        <h1>Resultado da Validação em Lote</h1>

        <div class="source-info">
            <p><span class="source-name">Fonte 1:</span> {source1}</p>
            <p><span class="source-name">Fonte 2:</span> {source2}</p>
            <p><span class="source-name">Método:</span> {method}</p>
        </div>

        {filter_info}

        <!-- Controles para alternar entre visualizações -->
      
        <!-- Visualização para Negócios -->
        <div id="business-view" class="view-content">
            <div class="business-summary">
                <h3>Resumo da Validação em Linguagem Simplificada</h3>
                <p>O sistema comparou {total_files} arquivos entre as duas fontes e encontrou:</p>
            </div>

            <div class="business-stats">
                <div class="business-stat-item">
                    <div class="business-stat-number">{identical_files}</div>
                    <div class="business-stat-label">Arquivos Idênticos</div>
                </div>
                <div class="business-stat-item">
                    <div class="business-stat-number" style="color: #e74c3c;">{different_files}</div>
                    <div class="business-stat-label">Arquivos com Diferenças</div>
                </div>
                <div class="business-stat-item">
                    <div class="business-stat-number">{only_in_source1}</div>
                    <div class="business-stat-label">Arquivos Exclusivos da Fonte 1</div>
                </div>
                <div class="business-stat-item">
                    <div class="business-stat-number">{only_in_source2}</div>
                    <div class="business-stat-label">Arquivos Exclusivos da Fonte 2</div>
                </div>
            </div>

            <h3>Lista de Arquivos Diferentes</h3>

            <div class="business-file-list">
                {business_file_rows}
            </div>
        </div>

        <div id="details-modal" class="modal">
            <div class="modal-content">
                <span class="close">&times;</span>
                <h2 id="modal-title">Detalhes da Validação</h2>

                <!-- Controles para alternar entre visualizações no modal -->
               <!-- Controles para alternar entre visualizações -->
<div class="view-controls">
    <button class="view-button active" onclick="switchView('business')">Visão Negocial</button>
</div>

                <div id="file-info-section" class="file-info-section">
                    <!-- Informações do arquivo serão inseridas aqui -->
                </div>

                <!-- Visualização Técnica no Modal -->
                <div id="modal-technical-view" class="view-content active">
                    <div id="diff-controls" class="diff-controls" style="display: none;">
                        <div>
                            <button type="button" class="diff-view-btn active" id="unified-view-btn" onclick="switchView('unified')">Visão Unificada</button>
                            <button type="button" class="diff-view-btn" id="split-view-btn" onclick="switchView('split')">Visão Lado a Lado</button>
                        </div>
                        <div>
                            <label for="show-unchanged"><input type="checkbox" id="show-unchanged" checked onchange="toggleUnchanged()"> Mostrar linhas não alteradas</label>
                        </div>
                    </div>

                    <div id="unified-diff-view">
                        <!-- Conteúdo da visualização unificada -->
                    </div>

                    <div id="split-diff-view" style="display: none;">
                        <div class="advanced-diff-container">
                            <div class="file-column file1-column">
                                <h3 id="file1-name"></h3>
                                <div id="file1-content"></div>
                            </div>
                            <div class="file-column file2-column">
                                <h3 id="file2-name"></h3>
                                <div id="file2-content"></div>
                            </div>
                        </div>
                    </div>
                </div>

                <!-- Visualização de Negócio no Modal -->
                <div id="modal-business-view" class="view-content active">
                    <!-- O conteúdo da visualização de negócio será inserido via JavaScript -->
                </div>
            </div>
        </div>

        <a href="/" class="btn-home">
            <button>Nova Validação</button>
        </a>
    </div>

    <script>
        // Filtrar a tabela por status e pesquisa
        function filterTable() {
            var filter = document.getElementById('filter-status').value;
            var searchText = document.getElementById('search-box').value.toLowerCase();
            var table = document.getElementById('results-table');
            var rows = table.getElementsByTagName('tr');
            var noResults = true;

            for (var i = 1; i < rows.length; i++) {
                var statusCell = rows[i].getElementsByTagName('td')[1];
                var filenameCell = rows[i].getElementsByTagName('td')[0];
                var showRow = true;

                if (statusCell) {
                    var statusText = statusCell.textContent || statusCell.innerText;
                    // Filtro por status
                    if (filter !== 'all' && !statusText.includes(filter)) {
                        showRow = false;
                    }

                    // Filtro por pesquisa de texto
                    if (searchText) {
                        var filename = filenameCell.textContent || filenameCell.innerText;
                        if (!filename.toLowerCase().includes(searchText)) {
                            showRow = false;
                        }
                    }

                    if (showRow) {
                        rows[i].style.display = '';
                        noResults = false;
                    } else {
                        rows[i].style.display = 'none';
                    }
                }
            }

            // Exibir mensagem se não houver resultados
            var emptyMessage = document.getElementById('empty-results-message');
            if (noResults) {
                if (!emptyMessage) {
                    emptyMessage = document.createElement('div');
                    emptyMessage.id = 'empty-results-message';
                    emptyMessage.className = 'empty-message';
                    emptyMessage.innerHTML = 'Nenhum arquivo corresponde aos filtros aplicados.';
                    table.parentNode.insertBefore(emptyMessage, table.nextSibling);
                }
                table.style.display = 'none';
            } else {
                if (emptyMessage) {
                    emptyMessage.remove();
                }
                table.style.display = '';
            }
        }

        // Modal para exibir detalhes
        var modal = document.getElementById('details-modal');
        var span = document.getElementsByClassName('close')[0];

        // Quando o usuário clica no X, fecha o modal
        span.onclick = function() {
            modal.style.display = 'none';
        }

        // Quando o usuário clica fora do modal, fecha-o
        window.onclick = function(event) {
            if (event.target == modal) {
                modal.style.display = 'none';
            }
        }

        // Alternar entre visualizações técnica e de negócio
        function switchView(view) {
            document.querySelectorAll('.view-content').forEach(el => {
                el.classList.remove('active');
            });

            document.querySelectorAll('.view-button').forEach(el => {
                el.classList.remove('active');
            });

            document.getElementById(view + '-view').classList.add('active');

            document.querySelectorAll('.view-button').forEach(el => {
                if (el.textContent.toLowerCase().includes(view)) {
                    el.classList.add('active');
                }
            });
        }

        // Alternar entre visualizações técnica e de negócio no modal
        function switchModalView(view) {
            document.querySelectorAll('#details-modal .view-content').forEach(el => {
                el.classList.remove('active');
            });

            document.querySelectorAll('#details-modal .view-button').forEach(el => {
                el.classList.remove('active');
            });

            document.getElementById('modal-' + view + '-view').classList.add('active');

            document.querySelectorAll('#details-modal .view-button').forEach(el => {
                if (el.textContent.toLowerCase().includes(view)) {
                    el.classList.add('active');
                }
            });
        }

        // Função para mostrar detalhes no modal
        function showDetails(index) {
            var detailsData = fileDetails[index];
            var modalTitle = document.getElementById('modal-title');
            var fileInfoSection = document.getElementById('file-info-section');
            var unifiedDiffView = document.getElementById('unified-diff-view');
            var diffControls = document.getElementById('diff-controls');
            var modalBusinessView = document.getElementById('modal-business-view');

            // Atualizar o título
            modalTitle.innerHTML = 'Detalhes: ' + detailsData.filename;

            // Exibir informações do arquivo
            var fileInfo = detailsData.file_info || {};
            fileInfoSection.innerHTML = '';

            if (detailsData.comparison_result === 'identical' || detailsData.comparison_result === 'different') {
                fileInfoSection.innerHTML = `
                    <div class="file-info-item">
                        <span class="file-info-label">Fonte 1:</span>
                        Tamanho: ${formatSize(fileInfo.size1 || 0)}<br>
                        Modificado: ${fileInfo.modified1 || 'N/A'}
                    </div>
                    <div class="file-info-item">
                        <span class="file-info-label">Fonte 2:</span>
                        Tamanho: ${formatSize(fileInfo.size2 || 0)}<br>
                        Modificado: ${fileInfo.modified2 || 'N/A'}
                    </div>
                `;
            } else if (detailsData.comparison_result === 'only_in_source1') {
                fileInfoSection.innerHTML = `
                    <div class="file-info-item">
                        <span class="file-info-label">Fonte 1:</span>
                        Tamanho: ${formatSize(fileInfo.size1 || 0)}<br>
                        Modificado: ${fileInfo.modified1 || 'N/A'}
                    </div>
                    <div class="file-info-item">
                        <span class="file-info-label">Fonte 2:</span>
                        Arquivo não existe
                    </div>
                `;
            } else if (detailsData.comparison_result === 'only_in_source2') {
                fileInfoSection.innerHTML = `
                    <div class="file-info-item">
                        <span class="file-info-label">Fonte 1:</span>
                        Arquivo não existe
                    </div>
                    <div class="file-info-item">
                        <span class="file-info-label">Fonte 2:</span>
                        Tamanho: ${formatSize(fileInfo.size2 || 0)}<br>
                        Modificado: ${fileInfo.modified2 || 'N/A'}
                    </div>
                `;
            }

            // Reset as visualizações
            document.getElementById('modal-technical-view').classList.add('active');
            document.getElementById('modal-business-view').classList.remove('active');
            document.querySelectorAll('#details-modal .view-button').forEach(el => {
                el.classList.remove('active');
                if (el.textContent.toLowerCase().includes('technical')) {
                    el.classList.add('active');
                }
            });

            // Conteúdo do modal baseado no resultado da comparação
            unifiedDiffView.innerHTML = '';
            modalBusinessView.innerHTML = '';

            // Preparar visualização de negócio
            let businessViewHtml = '';

            if (detailsData.comparison_result === 'only_in_source1') {
                unifiedDiffView.innerHTML = '<div class="diff-container"><p>Este arquivo existe apenas na Fonte 1.</p></div>';
                diffControls.style.display = 'none';

                businessViewHtml = `
                    <div class="difference-summary">
                        <h3>Explicação Simplificada</h3>
                        <p>Este arquivo está presente apenas na primeira fonte de dados.</p>
                        <p>Nome do arquivo: <strong>${detailsData.filename}</strong></p>
                        <p>Isso pode significar que o arquivo foi removido na segunda fonte ou foi adicionado apenas na primeira fonte.</p>
                    </div>
                `;
            } else if (detailsData.comparison_result === 'only_in_source2') {
                unifiedDiffView.innerHTML = '<div class="diff-container"><p>Este arquivo existe apenas na Fonte 2.</p></div>';
                diffControls.style.display = 'none';

                businessViewHtml = `
                    <div class="difference-summary">
                        <h3>Explicação Simplificada</h3>
                        <p>Este arquivo está presente apenas na segunda fonte de dados.</p>
                        <p>Nome do arquivo: <strong>${detailsData.filename}</strong></p>
                        <p>Isso pode significar que o arquivo foi adicionado recentemente ou foi removido da primeira fonte.</p>
                    </div>
                `;
            } else if (detailsData.comparison_result === 'identical') {
                unifiedDiffView.innerHTML = '<div class="diff-container"><p>Os arquivos são idênticos.</p></div>';
                diffControls.style.display = 'none';

                businessViewHtml = `
                    <div class="difference-summary">
                        <h3>Explicação Simplificada</h3>
                        <p>Os arquivos são exatamente iguais em ambas as fontes.</p>
                        <p>Nome do arquivo: <strong>${detailsData.filename}</strong></p>
                        <p>Não há diferenças no conteúdo entre as duas versões deste arquivo.</p>
                    </div>
                `;
            } else {
                diffControls.style.display = 'flex';

                // Preparar visualização unificada
                var detailsHtml = '<div class="diff-container">';

                if (typeof detailsData.details === 'string') {
                    detailsHtml += '<p>' + detailsData.details + '</p>';
                } else if (Array.isArray(detailsData.details)) {
                    if (detailsData.method === 'content') {
                        for (var i = 0; i < detailsData.details.length; i++) {
                            var line = detailsData.details[i];
                            if (line.startsWith('+') && !line.startsWith('+++')) {
                                detailsHtml += '<pre class="diff-line diff-added">' + escapeHtml(line) + '</pre>';
                            } else if (line.startsWith('-') && !line.startsWith('---')) {
                                detailsHtml += '<pre class="diff-line diff-removed">' + escapeHtml(line) + '</pre>';
                            } else if (line.startsWith('@@')) {
                                detailsHtml += '<div class="diff-header">' + escapeHtml(line) + '</div>';
                            } else {
                                detailsHtml += '<pre class="diff-line diff-context">' + escapeHtml(line) + '</pre>';
                            }
                        }
                    } else if (detailsData.method === 'binary') {
                        for (var i = 0; i < detailsData.details.length; i++) {
                            detailsHtml += '<div class="binary-diff">' + escapeHtml(detailsData.details[i]) + '</div>';
                        }
                    }
                }

                detailsHtml += '</div>';
                unifiedDiffView.innerHTML = detailsHtml;

                // Preparar visualização de negócio
                // Verificar se temos detalhes para analisar
                if (Array.isArray(detailsData.details) && detailsData.details.length > 0) {
                    const removedLines = detailsData.details.filter(line => 
                        line.startsWith('-') && !line.startsWith('---'));
                    const addedLines = detailsData.details.filter(line => 
                        line.startsWith('+') && !line.startsWith('+++'));

                    businessViewHtml = `
                        <div class="difference-summary">
                            <h3>Explicação Simplificada</h3>
                            <p>Os arquivos são diferentes. Foram encontradas alterações no conteúdo.</p>
                            <p>Nome do arquivo: <strong>${detailsData.filename}</strong></p>
                            <p>Total de modificações: ${removedLines.length + addedLines.length}</p>
                        </div>
                    `;

                    // Mostrar diferenças de forma amigável
                    if (removedLines.length > 0 || addedLines.length > 0) {
                        businessViewHtml += '<div class="changes-container"><h3>Alterações Detalhadas</h3>';

                        // Se temos o mesmo número de linhas removidas e adicionadas, podemos tentar parear
                        if (removedLines.length > 0 && addedLines.length > 0) {
                            // Agrupar por possíveis pares de alterações
                            const groups = [];
                            let currentGroup = { removed: [], added: [] };

                            for (let i = 0; i < detailsData.details.length; i++) {
                                const line = detailsData.details[i];

                                if (line.startsWith('@@')) {
                                    if (currentGroup.removed.length > 0 || currentGroup.added.length > 0) {
                                        groups.push(currentGroup);
                                    }
                                    currentGroup = { removed: [], added: [] };
                                } else if (line.startsWith('-') && !line.startsWith('---')) {
                                    currentGroup.removed.push(line.substring(1));
                                } else if (line.startsWith('+') && !line.startsWith('+++')) {
                                    currentGroup.added.push(line.substring(1));
                                }
                            }

                            if (currentGroup.removed.length > 0 || currentGroup.added.length > 0) {
                                groups.push(currentGroup);
                            }

                            // Gerar HTML para cada grupo
                            groups.forEach((group, index) => {
                                businessViewHtml += `
                                    <div class="change-item">
                                        <div class="change-header">Alteração ${index + 1}</div>
                                        <div class="change-explanation">
                                            <p>${getSimpleExplanation(group.removed, group.added)}</p>
                                        </div>
                                        <div class="change-content">
                                            <div class="before-change">
                                                <h4>Versão Original:</h4>
                                                <div class="code-block">
                                `;

                                // Processar versão original
                                if (group.removed.length === group.added.length) {
                                    // Possível alteração de caracteres
                                    businessViewHtml += highlightChangedChars(group.removed, group.added);
                                } else {
                                    businessViewHtml += group.removed.length > 0 ? 
                                        group.removed.join('<br>') : "[Sem conteúdo]";
                                }

                                businessViewHtml += `
                                                </div>
                                            </div>
                                            <div class="after-change">
                                                <h4>Nova Versão:</h4>
                                                <div class="code-block">
                                `;

                                // Processar nova versão
                                if (group.removed.length === group.added.length) {
                                    // Para alteração de caracteres, já aplicamos o destaque
                                    businessViewHtml += group.added.length > 0 ? 
                                        group.added.join('<br>') : "[Sem conteúdo]";
                                } else {
                                    businessViewHtml += group.added.length > 0 ? 
                                        group.added.join('<br>') : "[Conteúdo removido]";
                                }

                                businessViewHtml += `
                                                </div>
                                            </div>
                                        </div>
                                    </div>
                                `;
                            });
                        } else {
                            // Mostrar todas as linhas removidas e adicionadas
                            businessViewHtml += `
                                <div class="change-item">
                                    <div class="change-header">Alterações</div>
                                    <div class="change-explanation">
                                        <p>O conteúdo do arquivo foi modificado. ${removedLines.length} linha(s) removida(s) e ${addedLines.length} linha(s) adicionada(s).</p>
                                    </div>
                                    <div class="change-content">
                                        <div class="before-change">
                                            <h4>Linhas Removidas:</h4>
                                            <div class="code-block">
                                                ${removedLines.length > 0 ? 
                                                    removedLines.map(line => line.substring(1)).join('<br>') : 
                                                    "Nenhuma linha removida."}
                                            </div>
                                        </div>
                                        <div class="after-change">
                                            <h4>Linhas Adicionadas:</h4>
                                            <div class="code-block">
                                                ${addedLines.length > 0 ? 
                                                    addedLines.map(line => line.substring(1)).join('<br>') : 
                                                    "Nenhuma linha adicionada."}
                                            </div>
                                        </div>
                                    </div>
                                </div>
                            `;
                        }

                        businessViewHtml += '</div>';
                    }
                } else {
                    businessViewHtml = `
                        <div class="difference-summary">
                            <h3>Explicação Simplificada</h3>
                            <p>Os arquivos são diferentes, mas não foi possível identificar as diferenças específicas.</p>
                            <p>Nome do arquivo: <strong>${detailsData.filename}</strong></p>
                            <p>Para ver mais detalhes, verifique a visualização técnica.</p>
                        </div>
                    `;
                }

                // Preparar visualização lado a lado se disponível
                if (detailsData.processed_diff) {
                    document.getElementById('file1-name').textContent = 'Fonte 1: ' + detailsData.filename;
                    document.getElementById('file2-name').textContent = 'Fonte 2: ' + detailsData.filename;

                    // Renderizar a visualização lado a lado ao clicar na aba
                    // A renderização real é feita na função renderSplitView
                    window.currentProcessedDiff = detailsData.processed_diff;
                } else {
                    // Se não houver dados processados para visualização lado a lado, desabilitar o botão
                    document.getElementById('split-view-btn').disabled = true;
                    document.getElementById('split-view-btn').style.opacity = '0.5';
                }
            }

            // Atualizar a visualização de negócio
            modalBusinessView.innerHTML = businessViewHtml;

            modal.style.display = 'block';

            // Começar com a visualização unificada
            switchView('unified');
        }

        // Funções para destacar diferenças de caracteres
        function highlightChangedChars(originalLines, newLines) {
            if (originalLines.length !== newLines.length) {
                return originalLines.join('<br>');
            }

            let result = '';

            for (let i = 0; i < originalLines.length; i++) {
                const original = originalLines[i];
                const newLine = newLines[i];

                if (original === newLine) {
                    result += original + (i < originalLines.length - 1 ? '<br>' : '');
                    continue;
                }

                // Identificar as diferenças de caracteres
                let highlighted = '';
                let j = 0;
                let k = 0;

                // Algoritmo simplificado para destacar diferenças
                while (j < original.length || k < newLine.length) {
                    if (j < original.length && k < newLine.length && original[j] === newLine[k]) {
                        highlighted += original[j];
                        j++;
                        k++;
                    } else {
                        // Encontrou diferença
                        let matchIndex = -1;
                        let bestMatchPos = -1;

                        // Procurar o próximo caractere em comum
                        for (let look = 1; look < 10 && j + look < original.length; look++) {
                            const originalChar = original[j + look];

                            for (let lookNew = 0; lookNew < 10 && k + lookNew < newLine.length; lookNew++) {
                                if (originalChar === newLine[k + lookNew]) {
                                    if (matchIndex === -1 || lookNew < bestMatchPos) {
                                        matchIndex = look;
                                        bestMatchPos = lookNew;
                                        break;
                                    }
                                }
                            }

                            if (matchIndex !== -1) break;
                        }

                        if (matchIndex !== -1) {
                            // Destacar caracteres diferentes
                            highlighted += '<span class="highlight-remove">' + original.substring(j, j + matchIndex) + '</span>';
                            j += matchIndex;
                        } else {
                            // Não achou correspondência, adiciona o resto
                            highlighted += '<span class="highlight-remove">' + original.substring(j) + '</span>';
                            j = original.length;
                        }
                    }
                }

                result += highlighted + (i < originalLines.length - 1 ? '<br>' : '');
            }

            return result;
        }

        // Gerar explicação simples do que mudou
        function getSimpleExplanation(removedLines, addedLines) {
            if (removedLines.length === 0 && addedLines.length === 0) {
                return "Não foram detectadas alterações específicas.";
            }

            if (removedLines.length === 0) {
                return `Foram adicionadas ${addedLines.length} nova(s) linha(s) ao arquivo.`;
            }

            if (addedLines.length === 0) {
                return `Foram removidas ${removedLines.length} linha(s) do arquivo original.`;
            }

            if (removedLines.length === addedLines.length) {
                return `${removedLines.length} linha(s) foram modificadas, possivelmente alterando alguns caracteres.`;
            }

            return `Foram removidas ${removedLines.length} linha(s) e adicionadas ${addedLines.length} linha(s).`;
        }

        // Renderizar visualização lado a lado
        function renderSplitView() {
            if (!window.currentProcessedDiff) return;

            const file1Content = document.getElementById('file1-content');
            const file2Content = document.getElementById('file2-content');

            file1Content.innerHTML = '';
            file2Content.innerHTML = '';

            window.currentProcessedDiff.forEach(item => {
                if (item.type === 'unchanged') {
                    const line1 = document.createElement('div');
                    line1.className = 'code-line diff-line-unchanged';
                    line1.innerHTML = `<span class="line-number">${item.lineNum1}</span>${escapeHtml(item.content)}`;
                    file1Content.appendChild(line1);

                    const line2 = document.createElement('div');
                    line2.className = 'code-line diff-line-unchanged';
                    line2.innerHTML = `<span class="line-number">${item.lineNum2}</span>${escapeHtml(item.content)}`;
                    file2Content.appendChild(line2);
                } else if (item.type === 'removed') {
                    const line = document.createElement('div');
                    line.className = 'code-line diff-line-removed';
                    line.innerHTML = `<span class="line-number">${item.lineNum1}</span>${escapeHtml(item.content)}`;
                    file1Content.appendChild(line);

                    // Adicionar espaço em branco no arquivo 2 se não for um par
                    if (!item.paired) {
                        const emptyLine = document.createElement('div');
                        emptyLine.className = 'code-line';
                        emptyLine.innerHTML = `<span class="line-number"></span>`;
                        file2Content.appendChild(emptyLine);
                    }
                } else if (item.type === 'added') {
                    // Adicionar espaço em branco no arquivo 1 se não for um par
                    if (!item.paired) {
                        const emptyLine = document.createElement('div');
                        emptyLine.className = 'code-line';
                        emptyLine.innerHTML = `<span class="line-number"></span>`;
                        file1Content.appendChild(emptyLine);
                    }

                    const line = document.createElement('div');
                    line.className = 'code-line diff-line-added';
                    line.innerHTML = `<span class="line-number">${item.lineNum2}</span>${escapeHtml(item.content)}`;
                    file2Content.appendChild(line);
                }
            });

            // Aplicar configuração de exibição de linhas inalteradas
            toggleUnchanged();
        }

        // Alternar entre visualizações unificada e lado a lado
        function switchView(view) {
            if (view === 'unified') {
                document.getElementById('unified-diff-view').style.display = 'block';
                document.getElementById('split-diff-view').style.display = 'none';
                document.getElementById('unified-view-btn').classList.add('active');
                document.getElementById('split-view-btn').classList.remove('active');
            } else {
                document.getElementById('unified-diff-view').style.display = 'none';
                document.getElementById('split-diff-view').style.display = 'block';
                document.getElementById('unified-view-btn').classList.remove('active');
                document.getElementById('split-view-btn').classList.add('active');
                renderSplitView();
            }
        }

        // Mostrar/ocultar linhas inalteradas
        function toggleUnchanged() {
            const showUnchanged = document.getElementById('show-unchanged').checked;
            const isSplitView = document.getElementById('split-diff-view').style.display !== 'none';

            if (isSplitView) {
                // Para visualização lado a lado
                const unchangedLines = document.querySelectorAll('.diff-line-unchanged');
                unchangedLines.forEach(line => {
                    line.style.display = showUnchanged ? 'block' : 'none';
                });
            } else {
                // Para visualização unificada
                const contextLines = document.querySelectorAll('.diff-context');
                contextLines.forEach(line => {
                    line.style.display = showUnchanged ? 'block' : 'none';
                });
            }
        }

        // Função para escapar HTML
        function escapeHtml(text) {
            if (!text) return '';
            const map = {
                '&': '&amp;',
                '<': '&lt;',
                '>': '&gt;',
                '"': '&quot;',
                "'": '&#039;'
            };
            return text.replace(/[&<>"']/g, function(m) { return map[m]; });
        }

        // Função para formatar tamanho de arquivo
        function formatSize(bytes) {
            if (bytes === 0) return "0 B";
            const sizes = ['B', 'KB', 'MB', 'GB', 'TB'];
            const i = Math.floor(Math.log(bytes) / Math.log(1024));
            return parseFloat((bytes / Math.pow(1024, i)).toFixed(2)) + ' ' + sizes[i];
        }

        // Dados de detalhes dos arquivos
        var fileDetails = {file_details_json};
    </script>
</body>
</html>"""

# Modificação no RESULT_HTML_TEMPLATE para incluir a visão amigável
RESULT_HTML_TEMPLATE = """<!DOCTYPE html>
<html lang="pt-br">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Resultado da Validação</title>
    <style>
        body {
            font-family: Arial, sans-serif;
            line-height: 1.6;
            margin: 0;
            padding: 20px;
            background-color: #f4f4f4;
        }
        .container {
            margin: 0 auto;
            background: white;
            padding: 20px;
            border-radius: 5px;
            box-shadow: 0 0 10px rgba(0, 0, 0, 0.1);
        }
        h1, h2, h3, h4 {
            color: #333;
        }
        .summary {
            margin: 20px 0;
            padding: 15px;
            border-radius: 5px;
        }
        .identical {
            background-color: #dff0d8;
            color: #3c763d;
        }
        .different {
            background-color: #f2dede;
            color: #a94442;
        }
        .error-message {
            background-color: #fcf8e3;
            color: #8a6d3b;
            padding: 10px;
            border-radius: 5px;
            margin-bottom: 20px;
        }
        .diff-container {
            background-color: #f8f8f8;
            padding: 15px;
            border-radius: 5px;
            overflow-x: auto;
            margin-bottom: 20px;
        }
        .diff-line {
            font-family: 'Courier New', monospace;
            white-space: pre;
            margin: 0;
            padding: 1px 5px;
        }
        .diff-added {
            background-color: #e6ffed;
            color: #22863a;
        }
        .diff-removed {
            background-color: #ffeef0;
            color: #cb2431;
        }
        .binary-diff {
            font-family: 'Courier New', monospace;
            margin: 2px 0;
        }
        .file-info {
            margin-bottom: 20px;
            padding: 10px;
            background-color: #e9ecef;
            border-radius: 5px;
        }
        button {
            display: inline-block;
            padding: 10px 15px;
            background: #4CAF50;
            color: white;
            border: none;
            border-radius: 4px;
            cursor: pointer;
            font-size: 16px;
            text-decoration: none;
            margin-top: 20px;
        }
        button:hover {
            background: #45a049;
        }
        .file-name {
            font-weight: bold;
        }
        .diff-header {
            color: #6c757d;
            padding: 5px;
            margin-top: 10px;
        }
        .diff-context {
            color: #6c757d;
        }

        /* Estilos para a visualização amigável */
        .difference-summary {
            background-color: #f5f5f5;
            padding: 15px;
            border-radius: 8px;
            margin-bottom: 20px;
            border-left: 4px solid #3498db;
        }

        .change-item {
            margin-bottom: 20px;
            border: 1px solid #ddd;
            border-radius: 8px;
            overflow: hidden;
        }

        .change-header {
            background-color: #f0f0f0;
            padding: 10px;
            font-weight: bold;
            border-bottom: 1px solid #ddd;
        }

        .change-content {
            display: flex;
            flex-direction: column;
        }

        @media (min-width: 768px) {
            .change-content {
                flex-direction: row;
            }
        }

        .before-change, .after-change {
            padding: 15px;
            flex: 1;
        }

        .before-change {
            background-color: #ffeeee;
            border-right: 1px solid #ddd;
        }

        .after-change {
            background-color: #eeffee;
        }

        /* Estilos para destacar alterações de caracteres */
.highlight-remove {
    background-color: #ffb6c1;
    font-weight: bold;
    padding: 0 2px;
    border-radius: 3px;
    text-decoration: line-through;
}

.highlight-add {
    background-color: #90ee90;
    font-weight: bold;
    padding: 0 2px;
    border-radius: 3px;
}

        .code-block {
            font-family: monospace;
            padding: 10px;
            background-color: rgba(255, 255, 255, 0.5);
            border-radius: 4px;
            white-space: pre-wrap;
            word-break: break-all;
        }

        .change-explanation {
            padding: 10px;
            background-color: #fffde7;
            border-bottom: 1px solid #ddd;
            font-style: italic;
        }

        /* Botões para alternar entre visualizações */
        .view-controls {
            display: flex;
            justify-content: center;
            margin: 20px 0;
        }

        .view-button {
            padding: 8px 12px;
            margin: 0 5px;
            background-color: #f0f0f0;
            border: 1px solid #ddd;
            border-radius: 4px;
            cursor: pointer;
        }

        .view-button.active {
            background-color: #4CAF50;
            color: white;
            border-color: #45a049;
        }

        .view-content {
            display: none;
        }

        .view-content.active {
            display: block;
        }

        /* Estilos para gráfico visual de alterações */
        .file-map {
            margin: 20px 0;
            padding: 10px;
            background-color: #f9f9f9;
            border-radius: 4px;
        }

        .file-representation {
            height: 30px;
            display: flex;
            background-color: #ddd;
            border-radius: 15px;
            overflow: hidden;
        }

        .file-segment {
            height: 100%;
        }

        .file-segment.unchanged {
            background-color: #aaaaaa;
        }

        .file-segment.changed {
            background-color: #e74c3c;
        }

        .file-legend {
            text-align: center;
            font-size: 12px;
            color: #777;
            margin-top: 5px;
        }
    </style>

    <!-- Adicionar este código para o visualizador de diferenças lado a lado -->
    <style>
        .advanced-diff-container {
            display: flex;
            margin-top: 20px;
            border: 1px solid #ddd;
            border-radius: 4px;
            overflow: hidden;
        }
        .file-column {
            flex: 1;
            padding: 10px;
            max-height: 600px;
            overflow: auto;
        }
        .file1-column {
            background-color: #f8f9fa;
            border-right: 1px solid #ddd;
        }
        .file2-column {
            background-color: #f8f9fa;
        }
        .line-number {
            color: #6c757d;
            user-select: none;
            text-align: right;
            padding-right: 10px;
            width: 40px;
            display: inline-block;
            font-family: monospace;
            border-right: 1px solid #ddd;
            margin-right: 10px;
        }
        .code-line {
            font-family: 'Courier New', monospace;
            white-space: pre;
            padding: 2px 0;
            line-height: 1.4;
        }
        .diff-line-added {
            background-color: #e6ffed;
        }
        .diff-line-removed {
            background-color: #ffeef0;
        }
        .diff-line-unchanged {
            background-color: transparent;
        }
        .diff-controls {
            display: flex;
            justify-content: space-between;
            margin-bottom: 10px;
            padding: 10px;
            background-color: #f0f0f0;
            border-radius: 4px;
        }
        .diff-view-btn {
            background-color: #337ab7;
            color: white;
            border: none;
            padding: 5px 10px;
            border-radius: 4px;
            cursor: pointer;
        }
        .diff-view-btn:hover {
            background-color: #286090;
        }
        .diff-view-btn.active {
            background-color: #1e486b;
        }
    </style>
</head>
<body>
    <div class="container">
        <h1>Resultado da Validação</h1>

        <div class="file-info">
            <p><span class="file-name">Arquivo 1:</span> {file1}</p>
            <p><span class="file-name">Arquivo 2:</span> {file2}</p>
            <p><span class="file-name">Método:</span> {method}</p>
        </div>

        {error_message}

        <div class="summary {summary_class}">
            {summary_message}
        </div>

        <!-- Controles para alternar entre visualizações -->
        <div class="view-controls">
        </div>

      
        <!-- Visualização para Negócios -->
        <div id="business-view" class="view-content active">
            {business_view_html}
        </div>

        <a href="/" role="button"><button>Nova Validação</button></a>
    </div>

    <script>
        // Diferenças processadas para visualização lado a lado
        const processedDiff = {processed_diff_json};

        function renderSplitView() {
            const file1Content = document.getElementById('file1-content');
            const file2Content = document.getElementById('file2-content');

            file1Content.innerHTML = '';
            file2Content.innerHTML = '';

            processedDiff.forEach(item => {
                if (item.type === 'unchanged') {
                    const line1 = document.createElement('div');
                    line1.className = 'code-line diff-line-unchanged';
                    line1.innerHTML = `<span class="line-number">${item.lineNum1}</span>${escapeHtml(item.content)}`;
                    file1Content.appendChild(line1);

                    const line2 = document.createElement('div');
                    line2.className = 'code-line diff-line-unchanged';
                    line2.innerHTML = `<span class="line-number">${item.lineNum2}</span>${escapeHtml(item.content)}`;
                    file2Content.appendChild(line2);
                } else if (item.type === 'removed') {
                    const line = document.createElement('div');
                    line.className = 'code-line diff-line-removed';
                    line.innerHTML = `<span class="line-number">${item.lineNum1}</span>${escapeHtml(item.content)}`;
                    file1Content.appendChild(line);

                    // Adicionar espaço em branco no arquivo 2
                    if (!item.paired) {
                        const emptyLine = document.createElement('div');
                        emptyLine.className = 'code-line';
                        emptyLine.innerHTML = `<span class="line-number"></span>`;
                        file2Content.appendChild(emptyLine);
                    }
                } else if (item.type === 'added') {
                    // Adicionar espaço em branco no arquivo 1
                    if (!item.paired) {
                        const emptyLine = document.createElement('div');
                        emptyLine.className = 'code-line';
                        emptyLine.innerHTML = `<span class="line-number"></span>`;
                        file1Content.appendChild(emptyLine);
                    }

                    const line = document.createElement('div');
                    line.className = 'code-line diff-line-added';
                    line.innerHTML = `<span class="line-number">${item.lineNum2}</span>${escapeHtml(item.content)}`;
                    file2Content.appendChild(line);
                }
            });

            // Aplicar a configuração de mostrar/ocultar linhas não alteradas
            toggleUnchanged();
        }

        function switchDiffView(view) {
            if (view === 'unified') {
                document.getElementById('unified-diff-view').style.display = 'block';
                document.getElementById('split-diff-view').style.display = 'none';
                document.getElementById('unified-view-btn').classList.add('active');
                document.getElementById('split-view-btn').classList.remove('active');
            } else {
                document.getElementById('unified-diff-view').style.display = 'none';
                document.getElementById('split-diff-view').style.display = 'block';
                document.getElementById('unified-view-btn').classList.remove('active');
                document.getElementById('split-view-btn').classList.add('active');
                renderSplitView();
            }
        }

        function switchView(view) {
            // Ocultar todas as visualizações
            document.querySelectorAll('.view-content').forEach(el => {
                el.classList.remove('active');
            });

            // Desmarcar todos os botões
            document.querySelectorAll('.view-button').forEach(el => {
                el.classList.remove('active');
            });

            // Ativar a visualização selecionada
            document.getElementById(view + '-view').classList.add('active');

            // Marcar o botão selecionado
            document.querySelectorAll('.view-button').forEach(el => {
                if (el.textContent.toLowerCase().includes(view)) {
                    el.classList.add('active');
                }
            });
        }

        function toggleUnchanged() {
            const showUnchanged = document.getElementById('show-unchanged').checked;

            // Verificar qual visualização está ativa
            const isSplitView = document.getElementById('split-diff-view').style.display !== 'none';

            if (isSplitView) {
                // Para visualização lado a lado
                const unchangedLines = document.querySelectorAll('.diff-line-unchanged');
                unchangedLines.forEach(line => {
                    line.style.display = showUnchanged ? 'block' : 'none';
                });
            } else {
                // Para visualização unificada
                const contextLines = document.querySelectorAll('.diff-context');
                contextLines.forEach(line => {
                    line.style.display = showUnchanged ? 'block' : 'none';
                });
            }
        }

        function escapeHtml(text) {
            if (!text) return '';
            const map = {
                '&': '&amp;',
                '<': '&lt;',
                '>': '&gt;',
                '"': '&quot;',
                "'": '&#039;'
            };
            return text.replace(/[&<>"']/g, function(m) { return map[m]; });
        }
    </script>
</body>
</html>"""

# Definir templates HTML como constantes
EDIT_LAYOUT_HTML_TEMPLATE = """<!DOCTYPE html>
<html lang="pt-br">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Editar Layout - {{ layout_name }}</title>
    <style>
        body {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            background-color: #f9f9f9;
            color: #333;
            line-height: 1.6;
            padding: 20px;
        }
        .container {
            max-width: 1000px;
            margin: 0 auto;
            background: white;
            border-radius: 8px;
            box-shadow: 0 4px 12px rgba(0, 0, 0, 0.08);
            overflow: hidden;
            padding: 20px;
        }
        .header {
            display: flex;
            justify-content: space-between;
            align-items: center;
            margin-bottom: 20px;
            padding-bottom: 10px;
            border-bottom: 1px solid #e0e0e0;
        }
        .form-group {
            margin-bottom: 20px;
        }
        label {
            display: block;
            margin-bottom: 5px;
            font-weight: bold;
        }
        input[type="text"], textarea {
            width: 100%;
            padding: 10px;
            border: 1px solid #ddd;
            border-radius: 4px;
            font-size: 14px;
        }
        .field-row {
            display: flex;
            margin-bottom: 10px;
            gap: 10px;
            align-items: center;
        }
        .field-row input {
            flex: 1;
        }
        .remove-button {
            background-color: #f44336;
            color: white;
            border: none;
            border-radius: 4px;
            padding: 5px 10px;
            cursor: pointer;
        }
        .add-button {
            background-color: #2196F3;
            color: white;
            border: none;
            border-radius: 4px;
            padding: 8px 15px;
            cursor: pointer;
            margin-bottom: 20px;
        }
        .form-controls {
            margin-top: 30px;
            display: flex;
            gap: 10px;
        }
        .button {
            display: inline-block;
            padding: 10px 20px;
            background-color: #4CAF50;
            color: white;
            text-decoration: none;
            border-radius: 4px;
            font-weight: bold;
            border: none;
            cursor: pointer;
        }
        .button.secondary {
            background-color: #f0f0f0;
            color: #333;
        }
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>Editar Layout: {{ layout_name }}</h1>
            <div>
                <a href="/layouts" class="button secondary">Cancelar</a>
            </div>
        </div>

        <form action="/edit-layout/{{ layout_name }}" method="post">
            <div class="form-group">
                <label for="layout_name">Nome do Layout:</label>
                <input type="text" id="layout_name" name="layout_name" value="{{ layout.name }}" required>
            </div>

            <div class="form-group">
                <label for="layout_description">Descrição:</label>
                <textarea id="layout_description" name="layout_description" rows="3">{{ layout.description }}</textarea>
            </div>

            <div class="form-group">
                <label>Tipos de Registro:</label>
                <div id="record-types-container">
                    {% for record_type in layout.record_types %}
                    <input type="text" name="record_types" value="{{ record_type }}" placeholder="Tipo de Registro">
                    {% endfor %}
                    <input type="text" name="record_types" placeholder="Novo Tipo de Registro">
                </div>
                <button type="button" onclick="addRecordType()">+ Adicionar Tipo</button>
            </div>

            <h2>Campos do Layout</h2>
            <div id="fields-container">
                {% for field_name, field_info in layout.fields.items() %}
                <div class="field-row">
                    <input type="text" name="field_name" value="{{ field_name }}" placeholder="Nome do Campo" required>
                    <input type="text" name="field_type" value="{{ field_info.tipo_registro if 'tipo_registro' in field_info else '' }}" placeholder="Tipo de Registro">
                    <input type="text" name="field_position" value="{{ field_info.pos_inicial if 'pos_inicial' in field_info else '' }}" placeholder="Posição Inicial">
                    <input type="text" name="field_size" value="{{ field_info.tamanho if 'tamanho' in field_info else '' }}" placeholder="Tamanho">
                    <input type="text" name="field_description" value="{{ field_info.descricao if 'descricao' in field_info else '' }}" placeholder="Descrição">
                    <button type="button" class="remove-button" onclick="removeField(this)">-</button>
                </div>
                {% endfor %}
                <div class="field-row">
                    <input type="text" name="field_name" placeholder="Nome do Campo">
                    <input type="text" name="field_type" placeholder="Tipo de Registro">
                    <input type="text" name="field_position" placeholder="Posição Inicial">
                    <input type="text" name="field_size" placeholder="Tamanho">
                    <input type="text" name="field_description" placeholder="Descrição">
                    <button type="button" class="remove-button" onclick="removeField(this)">-</button>
                </div>
            </div>

            <button type="button" onclick="addField()" class="add-button">+ Adicionar Campo</button>

            <div class="form-controls">
                <button type="submit" class="button">Salvar Alterações</button>
                <a href="/view-layout/{{ layout_name }}" class="button secondary">Cancelar</a>
            </div>
        </form>
    </div>

    <script>
        function addField() {
            const container = document.getElementById('fields-container');
            const fieldRow = document.createElement('div');
            fieldRow.className = 'field-row';
            fieldRow.innerHTML = `
                <input type="text" name="field_name" placeholder="Nome do Campo">
                <input type="text" name="field_type" placeholder="Tipo de Registro">
                <input type="text" name="field_position" placeholder="Posição Inicial">
                <input type="text" name="field_size" placeholder="Tamanho">
                <input type="text" name="field_description" placeholder="Descrição">
                <button type="button" class="remove-button" onclick="removeField(this)">-</button>
            `;
            container.appendChild(fieldRow);
        }

        function removeField(button) {
            const row = button.parentElement;
            row.remove();
        }

        function addRecordType() {
            const container = document.getElementById('record-types-container');
            const input = document.createElement('input');
            input.type = 'text';
            input.name = 'record_types';
            input.placeholder = 'Novo Tipo de Registro';
            container.appendChild(input);
        }
    </script>
</body>
</html>"""

UPLOAD_AND_COMPARE_TEMPLATE = """<!DOCTYPE html>
<html lang="pt-br">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Upload e Comparação Rápida</title>
    <style>
        body {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            background-color: #f9f9f9;
            color: #333;
            line-height: 1.6;
            padding: 20px;
        }
        .container {
            margin: 0 auto;
            background: white;
            padding: 20px;
            border-radius: 8px;
            box-shadow: 0 4px 12px rgba(0, 0, 0, 0.08);
        }
        .form-group {
            margin-bottom: 20px;
        }
        .form-label {
            display: block;
            margin-bottom: 5px;
            font-weight: bold;
        }
        .form-control {
            width: 100%;
            padding: 10px;
            border: 1px solid #ddd;
            border-radius: 4px;
            font-size: 16px;
        }
        .file-input {
            display: block;
            width: 100%;
            padding: 10px;
            background-color: #f5f5f5;
            border: 1px dashed #aaa;
            border-radius: 4px;
            text-align: center;
            cursor: pointer;
        }
        .file-input input[type="file"] {
            display: none;
        }
        .button {
            display: inline-block;
            padding: 10px 20px;
            background-color: #4CAF50;
            color: white;
            border: none;
            border-radius: 4px;
            cursor: pointer;
            font-size: 16px;
            font-weight: bold;
        }
        .message {
            padding: 10px;
            margin-bottom: 20px;
            border-radius: 4px;
            background-color: #f8d7da;
            color: #721c24;
            border: 1px solid #f5c6cb;
        }
        .step {
            background-color: #f8f9fa;
            padding: 15px;
            border-radius: 5px;
            margin-bottom: 20px;
            border-left: 4px solid #4CAF50;
        }
        .step-title {
            font-weight: bold;
            margin-bottom: 10px;
            color: #4CAF50;
        }
        .file-preview {
            padding: 10px;
            background-color: #f0f0f0;
            border-radius: 5px;
            margin-top: 5px;
            font-size: 14px;
            display: none;
        }
    </style>
</head>
<body>
    <div class="container">
        <h1>Upload e Comparação Rápida</h1>

        <p>Esta página permite fazer upload de um documento de layout e dois arquivos a serem comparados em uma única operação.</p>

        <!-- Mensagem de erro, se houver -->
        {% if request.args.get('message') %}
        <div class="message">
            {{ request.args.get('message') }}
        </div>
        {% endif %}

        <form action="/upload-and-compare" method="post" enctype="multipart/form-data">
            <div class="step">
                <div class="step-title">Passo 1: Selecione o Documento de Layout</div>
                <div class="form-group">
                    <label for="layout-file" class="form-label">
                        Documento de Layout (Excel, Word, TXT)
                    </label>
                    <label class="file-input">
                        <input type="file" id="layout-file" name="layout_file" accept=".xlsx,.xls,.docx,.doc,.txt" required onchange="updateFilePreview(this, 'layout-preview')">
                        <span>Clique para selecionar o documento de layout</span>
                    </label>
                    <div id="layout-preview" class="file-preview"></div>
                    <small>O sistema extrairá automaticamente o layout deste documento.</small>
                </div>
            </div>

            <div class="step">
                <div class="step-title">Passo 2: Selecione os Arquivos a Serem Comparados</div>
                <div class="form-group">
                    <label for="file1" class="form-label">Arquivo 1</label>
                    <label class="file-input">
                        <input type="file" id="file1" name="file1" required onchange="updateFilePreview(this, 'file1-preview')">
                        <span>Clique para selecionar o primeiro arquivo</span>
                    </label>
                    <div id="file1-preview" class="file-preview"></div>
                </div>

                <div class="form-group">
                    <label for="file2" class="form-label">Arquivo 2</label>
                    <label class="file-input">
                        <input type="file" id="file2" name="file2" required onchange="updateFilePreview(this, 'file2-preview')">
                        <span>Clique para selecionar o segundo arquivo</span>
                    </label>
                    <div id="file2-preview" class="file-preview"></div>
                </div>
            </div>

            <div class="form-group">
                <button type="submit" class="button">Iniciar Análise</button>
                <a href="/" class="button" style="background-color: #f0f0f0; color: #333; margin-left: 10px;">Voltar</a>
            </div>
        </form>
    </div>

    <script>
        function updateFilePreview(input, previewId) {
            const preview = document.getElementById(previewId);
            if (input.files && input.files[0]) {
                const file = input.files[0];
                preview.textContent = `Arquivo selecionado: ${file.name} (${formatFileSize(file.size)})`;
                preview.style.display = 'block';

                // Atualizar o texto do label
                input.parentElement.querySelector('span').textContent = file.name;
            }
        }

        function formatFileSize(bytes) {
            if (bytes === 0) return '0 Bytes';
            const k = 1024;
            const sizes = ['Bytes', 'KB', 'MB', 'GB', 'TB'];
            const i = Math.floor(Math.log(bytes) / Math.log(k));
            return parseFloat((bytes / Math.pow(k, i)).toFixed(2)) + ' ' + sizes[i];
        }
    </script>
</body>
</html>"""

UPLOAD_AND_COMPARE_RESULT_TEMPLATE = """<!DOCTYPE html>
<html lang="pt-br">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Processando Comparação</title>
    <style>
        body {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            background-color: #f9f9f9;
            color: #333;
            line-height: 1.6;
            padding: 20px;
        }
        .container {
            margin: 0 auto;
            background: white;
            padding: 20px;
            border-radius: 8px;
            box-shadow: 0 4px 12px rgba(0, 0, 0, 0.08);
            text-align: center;
        }
        .loading {
            margin: 40px 0;
        }
        .spinner {
            display: inline-block;
            width: 50px;
            height: 50px;
            border: 5px solid rgba(0, 0, 0, 0.1);
            border-radius: 50%;
            border-top-color: #4CAF50;
            animation: spin 1s ease-in-out infinite;
        }
        @keyframes spin {
            to { transform: rotate(360deg); }
        }
        .file-info {
            background-color: #f8f9fa;
            padding: 15px;
            border-radius: 5px;
            margin: 20px 0;
            text-align: left;
        }
        .file-name {
            font-weight: bold;
        }
        .button {
            display: inline-block;
            padding: 10px 20px;
            background-color: #4CAF50;
            color: white;
            text-decoration: none;
            border: none;
            border-radius: 4px;
            cursor: pointer;
            font-size: 16px;
            font-weight: bold;
            margin-top: 20px;
        }
        .button.secondary {
            background-color: #f0f0f0;
            color: #333;
        }
    </style>
</head>
<body>
    <div class="container">
        <h1>Processando Arquivos</h1>

        <div class="file-info">
            <p><span class="file-name">Documento de Layout:</span> {{ layout_name }}</p>
            <p><span class="file-name">Arquivo 1:</span> {{ file1 }}</p>
            <p><span class="file-name">Arquivo 2:</span> {{ file2 }}</p>
        </div>

        <div class="loading">
            <div class="spinner"></div>
            <p>O sistema está extraindo o layout e analisando os arquivos...</p>
        </div>

        <div>
            <a href="/process-comparison/{{ timestamp }}" class="button">Ver Resultados</a>
            <a href="/" class="button secondary">Cancelar</a>
        </div>
    </div>

    <script>
        // Redirecionar automaticamente após um breve atraso
        setTimeout(function() {
            window.location.href = "/process-comparison/{{ timestamp }}";
        }, 3000);
    </script>
</body>
</html>"""

LAYOUTS_HTML_TEMPLATE = """<!DOCTYPE html>
<html lang="pt-br">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Gerenciamento de Layouts - Validador de Arquivos</title>
    <style>
        body {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            background-color: #f9f9f9;
            color: #333;
            line-height: 1.6;
            padding: 20px;
        }
        .container {
            margin: 0 auto;
            background: white;
            border-radius: 8px;
            box-shadow: 0 4px 12px rgba(0, 0, 0, 0.08);
            overflow: hidden;
        }
        header {
            background-color: #fff;
            padding: 20px;
            text-align: center;
            border-bottom: 1px solid #e0e0e0;
        }
        h1, h2 {
            color: #333;
        }
        .form-container {
            padding: 20px;
        }
        .section-title {
            font-size: 20px;
            margin-bottom: 20px;
            color: #333;
            font-weight: 600;
        }
        .form-group {
            margin-bottom: 20px;
        }
        .form-label {
            display: block;
            margin-bottom: 8px;
            font-weight: 500;
            color: #333;
        }
        .form-control {
            width: 100%;
            padding: 10px;
            border: 1px solid #ddd;
            border-radius: 4px;
            font-size: 16px;
        }
        .file-input {
            width: 100%;
            padding: 12px;
            border: 1px dashed var(--medium-gray);
            border-radius: 8px;
            cursor: pointer;
            background-color: #f9f9f9;
            color: #555;
            transition: all 0.3s;
            position: relative;
        }
        .file-input input[type="file"] {
            position: absolute;
            width: 100%;
            height: 100%;
            top: 0;
            left: 0;
            opacity: 0;
            cursor: pointer;
        }
        .button {
            display: block;
            width: 100%;
            padding: 14px;
            background-color: #4CAF50;
            color: white;
            border: none;
            border-radius: 4px;
            font-size: 16px;
            font-weight: 600;
            cursor: pointer;
            transition: background-color 0.3s;
            text-align: center;
            text-decoration: none;
        }
        .button:hover {
            background-color: #45a049;
        }
        .layouts-list {
            margin-top: 20px;
        }
        .layout-item {
            border: 1px solid #ddd;
            border-radius: 8px;
            padding: 15px;
            margin-bottom: 15px;
            background-color: #f9f9f9;
        }
        .layout-header {
            display: flex;
            justify-content: space-between;
            align-items: center;
            margin-bottom: 10px;
        }
        .layout-header h3 {
            margin: 0;
            font-size: 18px;
            color: #333;
        }
        .layout-actions {
            display: flex;
            gap: 10px;
        }
        .btn-edit, .btn-delete {
            padding: 5px 10px;
            border: none;
            border-radius: 4px;
            cursor: pointer;
            font-weight: 500;
        }
        .btn-edit {
            background-color: #2196F3;
            color: white;
        }
        .btn-delete {
            background-color: #f44336;
            color: white;
        }
        .layout-description {
            color: #666;
            margin-bottom: 10px;
        }
        .layout-details {
            background-color: #fff;
            padding: 10px;
            border-radius: 4px;
            border: 1px solid #eee;
        }
        .layout-details ul {
            margin: 5px 0 0 0;
            padding-left: 20px;
        }
        .empty-message {
            text-align: center;
            padding: 30px 0;
            color: #666;
            font-style: italic;
        }
    </style>
</head>
<body>
    <div class="container">
        <header>
            <h1>Gerenciamento de Layouts</h1>
        </header>

        <div class="form-container">
            <h2 class="section-title">Upload de Novo Layout</h2>

            <form action="/upload-layout-automatico" method="post" enctype="multipart/form-data">
                <div class="form-group">
                    <label for="layout-file" class="form-label">Arquivo de Layout (Excel, Word, TXT)</label>
                    <div class="file-input">
                        <input type="file" id="layout-file" name="layout_file" accept=".xlsx,.xls,.docx,.doc,.txt" required>
                        <span>Selecionar arquivo...</span>
                    </div>
                </div>

                <div class="form-group">
                    <label for="layout-name" class="form-label">Nome do Layout</label>
                    <input type="text" id="layout-name" name="layout_name" class="form-control" required>
                </div>

                <div class="form-group">
                    <label for="layout-description" class="form-label">Descrição</label>
                    <textarea id="layout-description" name="layout_description" class="form-control" rows="3"></textarea>
                </div>

                <button type="submit" class="button">Adicionar Layout</button>
            </form>

            <h2 class="section-title" style="margin-top: 30px;">Layouts Disponíveis</h2>

            {% if layouts %}
            <div class="layouts-list">
                {% for layout in layouts %}
                <div class="layout-item">
                    <div class="layout-header">
                        <h3>{{ layout.name }}</h3>
                        <div class="layout-actions">
                            <button class="btn-edit" onclick="editLayout('{{ layout.name }}')">Editar</button>
                            <button class="btn-delete" onclick="deleteLayout('{{ layout.name }}')">Excluir</button>
                        </div>
                    </div>
                    <div class="layout-description">{{ layout.description }}</div>
                    <div class="layout-details">
                        <strong>Identificadores:</strong>
                        <ul>
                            {% for key, value in layout.identifier.items() %}
                            <li>{{ key }}: {{ value }}</li>
                            {% endfor %}
                        </ul>
                    </div>
                </div>
                {% endfor %}
            </div>
            {% else %}
            <div class="empty-message">
                Nenhum layout disponível. Adicione um layout usando o formulário acima.
            </div>
            {% endif %}

            <a href="/" class="button" style="margin-top: 20px;">Voltar para a Página Principal</a>
        </div>
    </div>

    <script>
        // Funções para gerenciar layouts
        function editLayout(name) {
            window.location.href = '/edit-layout/' + name;
        }

        function deleteLayout(name) {
            if (confirm('Tem certeza que deseja excluir o layout "' + name + '"?')) {
                window.location.href = '/delete-layout/' + name;
            }
        }

        // Atualizar nome do arquivo selecionado
        document.getElementById('layout-file').addEventListener('change', function() {
            const fileName = this.files[0] ? this.files[0].name : 'Selecionar arquivo...';
            this.nextElementSibling.textContent = fileName;
        });
    </script>
</body>
</html>"""


def extrair_layout_de_documento(arquivo_path):
    """
    Extrai informações de layout de um documento (Excel, Word, TXT).

    Args:
        arquivo_path: Caminho para o arquivo de layout

    Returns:
        Dicionário com a estrutura do layout ou None se não for possível extrair
    """
    extensao = os.path.splitext(arquivo_path)[1].lower()

    try:
        if extensao == '.xlsx' or extensao == '.xls':
            return extrair_layout_excel(arquivo_path)
        elif extensao == '.docx':
            return extrair_layout_docx(arquivo_path)
        elif extensao == '.doc':
            # Para arquivos .doc antigos, é necessário um tratamento especial
            # Pode ser necessário uma biblioteca adicional como pywin32
            return extrair_layout_doc(arquivo_path)
        elif extensao == '.txt':
            return extrair_layout_txt(arquivo_path)
        else:
            return None
    except Exception as e:
        print(f"Erro ao extrair layout do arquivo {arquivo_path}: {str(e)}")
        return None


def extrair_layout_excel(arquivo_path):
    """Extrai layout de um arquivo Excel"""
    # Carregar o arquivo Excel
    df = pd.read_excel(arquivo_path)

    # Tentar identificar colunas padrão para informações de layout
    colunas_esperadas = ['nome_campo', 'tipo_registro', 'posicao', 'tamanho', 'descricao']
    colunas_alternativas = {
        'nome_campo': ['campo', 'nome do campo', 'field', 'field_name'],
        'tipo_registro': ['tipo', 'type', 'record_type', 'tipo registro'],
        'posicao': ['pos', 'position', 'posicao', 'posição', 'inicio', 'start'],
        'tamanho': ['tam', 'size', 'length', 'comprimento'],
        'descricao': ['desc', 'description', 'descricao', 'descrição']
    }

    # Mapear as colunas do DataFrame para as colunas esperadas
    mapeamento_colunas = {}
    for coluna_esperada, alternativas in colunas_alternativas.items():
        for col in df.columns:
            if col.lower() in [alt.lower() for alt in alternativas + [coluna_esperada]]:
                mapeamento_colunas[coluna_esperada] = col
                break

    # Verificar se encontramos pelo menos o campo, a posição e o tamanho
    colunas_obrigatorias = ['nome_campo', 'posicao', 'tamanho']
    if not all(col in mapeamento_colunas for col in colunas_obrigatorias):
        # Tentar uma abordagem mais flexível, analisando os valores das células
        return extrair_layout_excel_flexivel(df)

    # Converter o DataFrame para o formato de layout
    layout = {'fields': {}, 'name': 'Layout Extraído de Excel', 'description': 'Layout extraído automaticamente'}
    layout['record_types'] = []

    for _, row in df.iterrows():
        nome_campo = row[mapeamento_colunas['nome_campo']]
        if pd.isna(nome_campo):  # Pular linhas com valores vazios
            continue

        campo_info = {}

        # Obter tipo de registro, se disponível
        if 'tipo_registro' in mapeamento_colunas:
            tipo_registro = row[mapeamento_colunas['tipo_registro']]
            if not pd.isna(tipo_registro):
                tipo_registro = str(tipo_registro).strip()
                campo_info['tipo_registro'] = tipo_registro
                if tipo_registro not in layout['record_types']:
                    layout['record_types'].append(tipo_registro)

        # Obter posição inicial
        posicao = row[mapeamento_colunas['posicao']]
        if not pd.isna(posicao):
            try:
                # Tentar converter para inteiro
                campo_info['pos_inicial'] = int(posicao)
            except:
                # Se falhar, verificar se é um intervalo (ex: "10-20")
                if isinstance(posicao, str) and '-' in posicao:
                    try:
                        inicio, fim = posicao.split('-')
                        campo_info['pos_inicial'] = int(inicio.strip())
                        campo_info['pos_final'] = int(fim.strip())
                    except:
                        campo_info['pos_inicial'] = 1  # Valor padrão
                else:
                    campo_info['pos_inicial'] = 1  # Valor padrão

        # Obter tamanho
        tamanho = row[mapeamento_colunas['tamanho']]
        if not pd.isna(tamanho):
            try:
                campo_info['tamanho'] = int(tamanho)
            except:
                # Se não puder converter para inteiro, verificar se tem 'pos_final'
                if 'pos_final' in campo_info and 'pos_inicial' in campo_info:
                    campo_info['tamanho'] = campo_info['pos_final'] - campo_info['pos_inicial'] + 1
                else:
                    campo_info['tamanho'] = 1  # Valor padrão
        elif 'pos_final' in campo_info and 'pos_inicial' in campo_info:
            campo_info['tamanho'] = campo_info['pos_final'] - campo_info['pos_inicial'] + 1

        # Obter descrição, se disponível
        if 'descricao' in mapeamento_colunas:
            descricao = row[mapeamento_colunas['descricao']]
            if not pd.isna(descricao):
                campo_info['descricao'] = str(descricao)

        # Adicionar o campo ao layout
        layout['fields'][str(nome_campo)] = campo_info

    # Definir padrões para identificação automática
    layout['identifier'] = {
        'record_types': layout['record_types']
    }

    # Se não identificou tipos de registro, tentar alguma heurística
    if not layout['record_types']:
        tipos_possiveis = set()
        for campo, info in layout['fields'].items():
            if 'tipo_registro' in info:
                tipos_possiveis.add(info['tipo_registro'])

        if tipos_possiveis:
            layout['record_types'] = list(tipos_possiveis)
            layout['identifier']['record_types'] = list(tipos_possiveis)

    return layout


def extrair_layout_excel_flexivel(df):
    """Tenta extrair layout de um Excel quando o formato não segue o padrão esperado"""
    # Implementar heurísticas mais avançadas para identificar as colunas
    # Essa é uma versão simplificada que analisa as células para encontrar padrões comuns

    # Verificar se temos alguma coluna com valores que parecem ser posições
    posicao_candidatos = []
    for col in df.columns:
        # Verificar se a coluna tem valores numéricos ou padrões como "10-20"
        valores = df[col].dropna().astype(str)
        if valores.str.match(r'^\d+(-\d+)?$').any():
            posicao_candidatos.append(col)

    # Se não encontramos candidatos para posição, provavelmente não é um layout
    if not posicao_candidatos:
        return None

    # Tentar identificar a coluna de campos
    campos_candidatos = []
    for col in df.columns:
        # Uma coluna de campos geralmente tem strings como valor
        valores = df[col].dropna()
        if valores.dtype == 'object' and len(valores) > 0:
            campos_candidatos.append(col)

    if not campos_candidatos:
        return None

    # Escolher as colunas mais prováveis
    coluna_campo = campos_candidatos[0]  # Primeira coluna de texto
    coluna_posicao = posicao_candidatos[0]  # Primeira coluna numérica

    # Tentar encontrar a coluna de tamanho
    tamanho_candidatos = []
    for col in df.columns:
        if col != coluna_posicao:
            valores = df[col].dropna()
            if valores.dtype in ['int64', 'float64'] or (
                    valores.dtype == 'object' and valores.str.match(r'^\d+$').any()):
                tamanho_candidatos.append(col)

    coluna_tamanho = tamanho_candidatos[0] if tamanho_candidatos else None

    # Construir o layout
    layout = {'fields': {}, 'name': 'Layout Extraído (Flexível)', 'description': 'Layout extraído automaticamente'}
    layout['record_types'] = []

    for _, row in df.iterrows():
        campo = row[coluna_campo]
        if pd.isna(campo):  # Pular linhas com valores vazios
            continue

        campo_info = {}

        # Obter posição
        posicao = row[coluna_posicao]
        if not pd.isna(posicao):
            try:
                # Verificar se é um intervalo (ex: "10-20")
                if isinstance(posicao, str) and '-' in posicao:
                    inicio, fim = posicao.split('-')
                    campo_info['pos_inicial'] = int(inicio.strip())
                    campo_info['pos_final'] = int(fim.strip())
                    campo_info['tamanho'] = campo_info['pos_final'] - campo_info['pos_inicial'] + 1
                else:
                    campo_info['pos_inicial'] = int(posicao)
            except:
                campo_info['pos_inicial'] = 1  # Valor padrão

        # Obter tamanho, se a coluna foi identificada
        if coluna_tamanho:
            tamanho = row[coluna_tamanho]
            if not pd.isna(tamanho):
                try:
                    campo_info['tamanho'] = int(tamanho)
                except:
                    pass

        # Se não tiver tamanho definido, mas tiver posição final, calcular
        if 'tamanho' not in campo_info and 'pos_final' in campo_info and 'pos_inicial' in campo_info:
            campo_info['tamanho'] = campo_info['pos_final'] - campo_info['pos_inicial'] + 1

        # Verificar se há alguma coluna que pode ser tipo de registro
        for col in df.columns:
            if col not in [coluna_campo, coluna_posicao, coluna_tamanho]:
                valor = row[col]
                if not pd.isna(valor) and isinstance(valor, str) and len(valor) <= 2:
                    # Possível tipo de registro
                    campo_info['tipo_registro'] = valor
                    if valor not in layout['record_types']:
                        layout['record_types'].append(valor)
                    break

        # Adicionar o campo ao layout
        layout['fields'][str(campo)] = campo_info

    # Configurar identificadores
    layout['identifier'] = {}
    if layout['record_types']:
        layout['identifier']['record_types'] = layout['record_types']

    return layout


def extrair_layout_docx(arquivo_path):
    """Extrai layout de um arquivo Word (.docx)"""
    doc = Document(arquivo_path)
    layout = {'fields': {}, 'name': 'Layout Extraído de Word', 'description': 'Layout extraído automaticamente'}

    # Tentar extrair de tabelas primeiro
    if doc.tables:
        return extrair_layout_de_tabela_word(doc.tables)

    # Se não houver tabelas, tentar extrair do texto
    texto_completo = "\n".join([para.text for para in doc.paragraphs])

    # Padrões comuns para definições de layout
    # 1. Nome do campo seguido por posição e tamanho
    padrao1 = r'([A-Za-z0-9_]+).*?posição:?\s*(\d+)[-\s,a]+(\d+).*?tamanho:?\s*(\d+)'
    # 2. Nome do campo com tipo, posição e tamanho
    padrao2 = r'([A-Za-z0-9_]+).*?tipo:?\s*([A-Za-z0-9]+).*?posição:?\s*(\d+).*?tamanho:?\s*(\d+)'
    # 3. Formato tabular
    padrao3 = r'([A-Za-z0-9_]+)\s+(\d+)\s+(\d+)\s+([A-Za-z0-9]+)'

    padroes = [padrao1, padrao2, padrao3]
    campos_encontrados = []

    for padrao in padroes:
        matches = re.finditer(padrao, texto_completo, re.IGNORECASE | re.MULTILINE)
        for match in matches:
            campos_encontrados.append(match.groups())

    # Se encontrou campos, processar conforme o padrão
    if campos_encontrados:
        for grupos in campos_encontrados:
            if len(grupos) == 4:  # padrao1 ou padrao3
                nome_campo, posicao, posicao_final, tamanho = grupos
                layout['fields'][nome_campo] = {
                    'pos_inicial': int(posicao),
                    'tamanho': int(tamanho)
                }
            elif len(grupos) == 4:  # padrao2
                nome_campo, tipo, posicao, tamanho = grupos
                layout['fields'][nome_campo] = {
                    'tipo_registro': tipo,
                    'pos_inicial': int(posicao),
                    'tamanho': int(tamanho)
                }

    # Extrair tipos de registro
    tipos_registro = set()
    for _, info in layout['fields'].items():
        if 'tipo_registro' in info:
            tipos_registro.add(info['tipo_registro'])

    layout['record_types'] = list(tipos_registro)

    # Configurar identificadores
    layout['identifier'] = {}
    if layout['record_types']:
        layout['identifier']['record_types'] = layout['record_types']

    return layout


def extrair_layout_de_tabela_word(tabelas):
    """Extrai layout de tabelas em documento Word"""
    layout = {'fields': {}, 'name': 'Layout Extraído de Tabela Word', 'description': 'Layout extraído automaticamente'}
    layout['record_types'] = []

    for tabela in tabelas:
        # Determinar cabeçalhos da tabela a partir da primeira linha
        if len(tabela.rows) < 2:
            continue

        cabecalhos = []
        for celula in tabela.rows[0].cells:
            cabecalhos.append(celula.text.strip().lower())

        # Verificar se a tabela tem colunas que parecem relevantes para um layout
        colunas_interesse = ['campo', 'nome', 'posição', 'posicao', 'inicio', 'tamanho', 'tipo']
        if not any(col in cabecalhos for col in colunas_interesse):
            continue

        # Mapear índices para as colunas de interesse
        indice_campo = next((i for i, h in enumerate(cabecalhos) if 'campo' in h or 'nome' in h), None)
        indice_posicao = next((i for i, h in enumerate(cabecalhos) if 'pos' in h or 'início' in h or 'inicio' in h),
                              None)
        indice_tamanho = next((i for i, h in enumerate(cabecalhos) if 'tam' in h or 'size' in h), None)
        indice_tipo = next((i for i, h in enumerate(cabecalhos) if 'tipo' in h or 'type' in h), None)

        # Se não encontrou colunas básicas, continuar para próxima tabela
        if indice_campo is None or (indice_posicao is None and indice_tamanho is None):
            continue

        # Processar linhas da tabela
        for i in range(1, len(tabela.rows)):
            row = tabela.rows[i]
            cells = [cell.text.strip() for cell in row.cells]

            # Ignorar linhas vazias
            if not any(cells):
                continue

            # Obter valores
            nome_campo = cells[indice_campo] if indice_campo < len(cells) else None
            if not nome_campo:
                continue

            campo_info = {}

            # Obter posição
            if indice_posicao is not None and indice_posicao < len(cells):
                pos_texto = cells[indice_posicao]
                try:
                    # Verificar se é um intervalo (ex: "10-20")
                    if '-' in pos_texto:
                        inicio, fim = pos_texto.split('-')
                        campo_info['pos_inicial'] = int(inicio.strip())
                        campo_info['pos_final'] = int(fim.strip())
                    else:
                        campo_info['pos_inicial'] = int(pos_texto)
                except:
                    pass

            # Obter tamanho
            if indice_tamanho is not None and indice_tamanho < len(cells):
                tam_texto = cells[indice_tamanho]
                try:
                    campo_info['tamanho'] = int(tam_texto)
                except:
                    # Se não puder converter para inteiro e tiver posição inicial e final
                    if 'pos_final' in campo_info and 'pos_inicial' in campo_info:
                        campo_info['tamanho'] = campo_info['pos_final'] - campo_info['pos_inicial'] + 1
            elif 'pos_final' in campo_info and 'pos_inicial' in campo_info:
                campo_info['tamanho'] = campo_info['pos_final'] - campo_info['pos_inicial'] + 1

            # Obter tipo de registro
            if indice_tipo is not None and indice_tipo < len(cells):
                tipo_texto = cells[indice_tipo]
                if tipo_texto:
                    campo_info['tipo_registro'] = tipo_texto
                    if tipo_texto not in layout['record_types']:
                        layout['record_types'].append(tipo_texto)

            # Adicionar o campo ao layout
            layout['fields'][nome_campo] = campo_info

    # Configurar identificadores
    layout['identifier'] = {}
    if layout['record_types']:
        layout['identifier']['record_types'] = layout['record_types']

    return layout


def extrair_layout_txt(arquivo_path):
    """Extrai layout de um arquivo TXT"""
    layout = {'fields': {}, 'name': 'Layout Extraído de TXT', 'description': 'Layout extraído automaticamente'}
    layout['record_types'] = []

    # Ler o arquivo
    with open(arquivo_path, 'r', encoding='utf-8', errors='replace') as f:
        linhas = f.readlines()

    # Padrões para reconhecer definições de layout em formato texto
    padrao_campo = r'([A-Za-z0-9_]+)[\s:]+.*?(\d+)[-\s,a]+(\d+).*?tamanho:?\s*(\d+)'
    padrao_campo_com_tipo = r'([A-Za-z0-9_]+)[\s:]+.*?tipo:?\s*([A-Za-z0-9]+).*?posição:?\s*(\d+).*?tamanho:?\s*(\d+)'

    # Verificar cada linha
    for linha in linhas:
        # Tentar padrão com tipo
        match = re.search(padrao_campo_com_tipo, linha, re.IGNORECASE)
        if match:
            nome_campo, tipo, posicao, tamanho = match.groups()
            layout['fields'][nome_campo] = {
                'tipo_registro': tipo,
                'pos_inicial': int(posicao),
                'tamanho': int(tamanho)
            }

            if tipo not in layout['record_types']:
                layout['record_types'].append(tipo)

            continue

        # Tentar padrão sem tipo
        match = re.search(padrao_campo, linha, re.IGNORECASE)
        if match:
            nome_campo, posicao, posicao_final, tamanho = match.groups()
            layout['fields'][nome_campo] = {
                'pos_inicial': int(posicao),
                'tamanho': int(tamanho)
            }

    # Configurar identificadores
    layout['identifier'] = {}
    if layout['record_types']:
        layout['identifier']['record_types'] = layout['record_types']

    return layout


def extrair_layout_doc(arquivo_path):
    """
    Extrai layout de um arquivo .doc (Word antigo)
    Requer implementação especializada devido ao formato binário
    """
    # Implementação necessitaria de biblioteca especializada como pywin32
    # Por enquanto, retornar None
    return None


# MÉTODO 1: upload_files CORRIGIDO
@app.route('/upload', methods=['POST'])
@login_required
def upload_files():
    if 'file1' not in request.files or 'file2' not in request.files:
        message = "Erro: Ambos os arquivos devem ser enviados"
        return redirect(f"/?message={urllib.parse.quote(message)}")

    file1 = request.files['file1']
    file2 = request.files['file2']

    # CORREÇÃO: Verificar se um layout foi selecionado
    layout_name = request.form.get('layout_name', '')
    layout = None

    if layout_name:
        layouts_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'layouts')
        layout_path = os.path.join(layouts_dir, f"{layout_name}.json")
        if os.path.exists(layout_path):
            try:
                # CORREÇÃO: Usar json.load() diretamente em vez de load_json_file()
                with open(layout_path, 'r', encoding='utf-8') as f:
                    layout = json.load(f)
                session['layout_name'] = layout_name.replace('_', ' ')
                session['selected_layout'] = layout  # NOVO: Armazenar layout na sessão
                print(f"Layout '{layout_name}' carregado com sucesso")
                print(
                    f"Campos do layout: {list(layout.get('fields', {}).keys()) if layout.get('fields') else 'Sem campos'}")
            except Exception as e:
                print(f"Erro ao carregar layout: {str(e)}")
                layout = None

    if file1.filename == '' or file2.filename == '':
        message = "Erro: Nenhum arquivo selecionado"
        return redirect(f"/?message={urllib.parse.quote(message)}")

    method = request.form.get('method', 'hash')

    # Salvar os arquivos com nomes únicos para evitar conflitos
    timestamp = str(int(time.time()))
    file1_name = f"{timestamp}_{secure_filename(file1.filename)}"
    file2_name = f"{timestamp}_{secure_filename(file2.filename)}"
    file1_path = os.path.join(app.config['UPLOAD_FOLDER'], file1_name)
    file2_path = os.path.join(app.config['UPLOAD_FOLDER'], file2_name)

    # Salvar os arquivos em blocos para lidar com arquivos grandes
    chunk_size = 8192

    # Salvar file1
    with open(file1_path, 'wb') as f:
        while True:
            chunk = file1.read(chunk_size)
            if not chunk:
                break
            f.write(chunk)

    # Salvar file2
    with open(file2_path, 'wb') as f:
        while True:
            chunk = file2.read(chunk_size)
            if not chunk:
                break
            f.write(chunk)

    # Comparar os arquivos
    result = {}
    result['file1'] = file1.filename
    result['file2'] = file2.filename
    result['method'] = method

    error_message = ""
    differences_html = ""
    summary_message = ""
    summary_class = ""
    processed_diff_json = "[]"  # Valor padrão para JSON de diferenças processadas
    business_view_html = ""  # HTML para a visualização amigável

    if method == 'hash':
        identical, message = compare_files_hash(file1_path, file2_path)
        if "Erro" in message:
            error_message = f'<div class="error-message">{message}</div>'
        elif identical:
            summary_message = "<h2>✓ Os arquivos são idênticos!</h2>"
            summary_class = "identical"
        else:
            summary_message = "<h2>✗ Os arquivos são diferentes!</h2>"
            summary_class = "different"
            business_view_html = '<div class="difference-summary"><h3>Resultado da Análise</h3><p>Os arquivos são diferentes, mas o método de hash não fornece detalhes específicos sobre as diferenças.</p><p>Para ver detalhes específicos, tente usar o método de "Conteúdo" ou "Binário".</p></div>'

    elif method == 'content':
        # CORREÇÃO: Usar layout selecionado se disponível
        if layout:
            print("Usando layout selecionado manualmente")
            differences, message = compare_files_content_with_layout(file1_path, file2_path, layout)
        else:
            print("Nenhum layout selecionado, usando detecção automática")
            differences, message = compare_files_content(file1_path, file2_path)

        if "Erro" in message:
            error_message = f'<div class="error-message">{message}</div>'
        elif not differences:
            summary_message = "<h2>✓ Os arquivos são idênticos!</h2>"
            summary_class = "identical"
        else:
            summary_message = "<h2>✗ Os arquivos são diferentes!</h2>"
            summary_class = "different"

            differences_html = '<h3>Diferenças Encontradas:</h3><div class="diff-container">'
            for line in differences:
                if line.startswith('+') and not line.startswith('+++'):
                    differences_html += f'<pre class="diff-line diff-added">{line}</pre>'
                elif line.startswith('-') and not line.startswith('---'):
                    differences_html += f'<pre class="diff-line diff-removed">{line}</pre>'
                elif line.startswith('@@'):
                    differences_html += f'<div class="diff-header">{line}</div>'
                else:
                    differences_html += f'<pre class="diff-line diff-context">{line}</pre>'
            differences_html += '</div>'

            # Gerar a visualização amigável para negócios
            business_view_html = generate_user_friendly_diff(differences, result['file1'], result['file2'])

            # Processar diferenças para visualizador avançado
            try:
                processed_diff = process_differences_for_advanced_view(file1_path, file2_path, differences)
                processed_diff_json = json.dumps(processed_diff)
            except Exception as e:
                # Se houver erro no processamento, apenas continua sem a visualização avançada
                pass

    elif method == 'binary':
        differences, message = compare_files_binary(file1_path, file2_path)
        if "Erro" in message:
            error_message = f'<div class="error-message">{message}</div>'
        elif not differences:
            summary_message = "<h2>✓ Os arquivos são idênticos!</h2>"
            summary_class = "identical"
        else:
            summary_message = "<h2>✗ Os arquivos são diferentes!</h2>"
            summary_class = "different"

            differences_html = '<h3>Diferenças Encontradas:</h3><div class="diff-container">'
            for diff in differences:
                differences_html += f'<div class="binary-diff">{diff}</div>'
            differences_html += '</div>'

            # Gerar a visualização amigável para negócios
            business_view_html = f"""
            <div class="difference-summary">
                <h3>Resultado da Análise Binária</h3>
                <p>Os arquivos são diferentes em seu conteúdo binário.</p>
                <p>Foram encontradas {len(differences)} diferenças a nível de bytes.</p>
            </div>

            <div class="changes-container">
                <h3>Explicação Simplificada</h3>
                <div class="change-item">
                    <div class="change-header">Informações Técnicas</div>
                    <div class="change-explanation">
                        <p>Os arquivos contêm dados diferentes. Isso pode significar que o conteúdo foi modificado 
                        ou que os arquivos são de tipos diferentes.</p>
                    </div>
                </div>
            </div>
            """

    # Limpar os arquivos após a comparação para não ocupar espaço desnecessariamente
    try:
        os.remove(file1_path)
        os.remove(file2_path)
    except Exception as e:
        # Se não conseguir remover, apenas continua
        pass

    # Renderizar o resultado usando substituição de string
    result_html = RESULT_HTML_TEMPLATE
    result_html = result_html.replace("{file1}", result['file1'])
    result_html = result_html.replace("{file2}", result['file2'])
    result_html = result_html.replace("{method}", result['method'])
    result_html = result_html.replace("{error_message}", error_message)
    result_html = result_html.replace("{summary_class}", summary_class)
    result_html = result_html.replace("{summary_message}", summary_message)
    result_html = result_html.replace("{differences_html}", differences_html)
    result_html = result_html.replace("{processed_diff_json}", processed_diff_json)
    result_html = result_html.replace("{business_view_html}", business_view_html)

    return result_html


@app.route('/info', methods=['GET'])
@login_required
def show_info():
    info_template = """<!DOCTYPE html>
    <html lang="pt-br">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Informações Importantes</title>
        <style>
            body {
                font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                background-color: #f9f9f9;
                color: #333;
                line-height: 1.6;
                padding: 20px;
                display: flex;
                justify-content: center;
                align-items: center;
                min-height: 100vh;
            }
            .info-container {
                max-width: 600px;
                background: white;
                border-radius: 8px;
                box-shadow: 0 4px 12px rgba(0, 0, 0, 0.1);
                padding: 30px;
                text-align: center;
            }
            .info-header {
                background-color: #4CAF50;
                color: white;
                padding: 15px;
                border-radius: 8px 8px 0 0;
                margin: -30px -30px 20px -30px;
            }
            .info-item {
                background-color: #f8f9fa;
                padding: 15px;
                margin: 15px 0;
                border-radius: 6px;
                border-left: 4px solid #4CAF50;
            }
            .button {
                background-color: #4CAF50;
                color: white;
                padding: 12px 24px;
                border: none;
                border-radius: 6px;
                cursor: pointer;
                text-decoration: none;
                display: inline-block;
                margin-top: 20px;
            }
        </style>
    </head>
    <body>
        <div class="info-container">
            <div class="info-header">
                <h2>Informações Importantes</h2>
            </div>

            <div class="info-item">
                <h3>🔍 Verificação Individual</h3>
                <p>Os arquivos <strong>devem ter nomes diferentes</strong> para comparação individual.</p>
            </div>

            <div class="info-item">
                <h3>📦 Verificação em Lote</h3>
                <p>Os nomes dos arquivos das <strong>duas origens deverão ser iguais</strong> para comparação em lote.</p>
            </div>

            <a href="/" class="button">Continuar para o Validador</a>
        </div>
    </body>
    </html>"""

    return info_template
# MÉTODO 2: NOVA FUNÇÃO compare_files_content_with_layout
def compare_files_content_with_layout(file1_path, file2_path, layout, context_lines=3):
    """
    Compara dois arquivos usando um layout específico fornecido.
    """
    if not os.path.exists(file1_path):
        return [], f"Erro: O arquivo '{file1_path}' não existe."
    if not os.path.exists(file2_path):
        return [], f"Erro: O arquivo '{file2_path}' não existe."

    # Verificação rápida usando filecmp
    if filecmp.cmp(file1_path, file2_path, shallow=False):
        return [], "Os arquivos são idênticos."

    try:
        print(f"Iniciando análise estruturada com layout: {layout.get('name', 'Layout sem nome')}")

        # Usar diretamente o layout fornecido para análise estruturada
        estrutura_diffs = analisar_arquivo_estruturado(file1_path, file2_path, layout)

        if estrutura_diffs is None:
            print("Aviso: Análise estruturada retornou None, usando análise padrão.")
            return super_compare_files(file1_path, file2_path, context_lines)

        # Armazenar para uso na visualização de negócio
        session['estrutura_diffs'] = estrutura_diffs
        session['layout_name'] = layout.get('name', 'Layout Selecionado')

        # Converter para o formato tradicional de diferenças
        traditional_diffs = []

        # Verificar se há erros na análise
        has_error = any(diff.get('tipo') == 'erro' for diff in estrutura_diffs)
        if has_error:
            error_diff = next(diff for diff in estrutura_diffs if diff.get('tipo') == 'erro')
            return [], f"Erro na análise estruturada: {error_diff.get('mensagem', 'Erro desconhecido')}"

        # Converter diferenças estruturadas para o formato tradicional
        for diff in estrutura_diffs:
            diff_type = diff.get('tipo', '')

            if diff_type == 'campos_alterados':
                traditional_diffs.append(f"@@ Linha {diff['linha']} - Alterações de campos @@")
                for campo in diff.get('diferenca', []):
                    traditional_diffs.append(f"- Campo {campo['campo']}: {campo['valor_antigo']}")
                    traditional_diffs.append(f"+ Campo {campo['campo']}: {campo['valor_novo']}")

            elif diff_type == 'linha_alterada':
                traditional_diffs.append(f"@@ Linha {diff['linha']} @@")
                traditional_diffs.append(f"- {diff.get('valor_antigo', '')}")
                traditional_diffs.append(f"+ {diff.get('valor_novo', '')}")

            elif diff_type == 'linha_removida':
                traditional_diffs.append(f"@@ Linha {diff['linha']} @@")
                traditional_diffs.append(f"- {diff.get('valor_antigo', '')}")

            elif diff_type == 'linha_adicionada':
                traditional_diffs.append(f"@@ Linha {diff['linha']} @@")
                traditional_diffs.append(f"+ {diff.get('valor_novo', '')}")

            elif diff_type == 'info':
                traditional_diffs.append(f"@@ {diff.get('mensagem', '')} @@")

        print(f"Análise estruturada concluída. Diferenças encontradas: {len(traditional_diffs)}")
        return traditional_diffs, f"Análise de arquivo estruturado completa (Layout: {layout.get('name', 'Selecionado')})."

    except Exception as e:
        import traceback
        traceback.print_exc()
        print(f"Erro ao processar arquivo estruturado: {str(e)}")
        return super_compare_files(file1_path, file2_path, context_lines)


def generate_user_friendly_diff(differences, file1_name, file2_name):
    """Gera uma visualização amigável das diferenças para usuários não técnicos."""
    print('CAIU NO AMIGÁVEL')
    # Verificar se temos diferenças estruturadas armazenadas na sessão
    estrutura_diffs = session.get('estrutura_diffs', None)

    # Verificar se estrutura_diffs é válido e não vazio
    if estrutura_diffs is not None and isinstance(estrutura_diffs, list):
        # Limpar da sessão para não interferir em comparações futuras
        session.pop('estrutura_diffs', None)

        # Usar a visualização estruturada
        try:
            return gerar_html_visao_negocio_estruturada(estrutura_diffs)
        except Exception as e:
            # Em caso de erro, registrar e continuar com a visualização padrão
            print(f"Erro ao gerar HTML estruturado: {str(e)}")
            import traceback
            traceback.print_exc()

    # Continuar com a visualização padrão
    # Contar o número de alterações
    removed_lines = [d for d in differences if d.startswith('-') and not d.startswith('---')]
    added_lines = [d for d in differences if d.startswith('+') and not d.startswith('+++')]
    total_changes = len(removed_lines) + len(added_lines)

    # Determinar o tipo de alteração
    change_type = "Conteúdo modificado"
    if removed_lines and not added_lines:
        change_type = "Conteúdo removido"
    elif added_lines and not removed_lines:
        change_type = "Conteúdo adicionado"

    html = f"""
    <div class="difference-summary">
        <h3>Resumo das Diferenças</h3>
        <ul>
            <li><strong>Status:</strong> Os arquivos são diferentes</li>
            <li><strong>Total de alterações:</strong> {total_changes} alterações encontradas</li>
            <li><strong>Tipo de alterações:</strong> {change_type}</li>
        </ul>
    </div>
    """

    # Processar as diferenças para um formato mais amigável
    change_items = []
    current_section = None
    section_changes = []

    for line in differences:
        if line.startswith('@@'):
            # Nova seção de diferenças
            if current_section and section_changes:
                change_items.append((current_section, section_changes))

            current_section = line
            section_changes = []
        elif line.startswith('+') or line.startswith('-'):
            if not line.startswith('+++') and not line.startswith('---'):
                section_changes.append(line)

    # Adicionar a última seção
    if current_section and section_changes:
        change_items.append((current_section, section_changes))

    # Gerar HTML para cada alteração
    html += '<div class="changes-container"><h3>Alterações Detalhadas</h3>'

    for i, (section, changes) in enumerate(change_items, 1):
        removed_lines = [c[1:] for c in changes if c.startswith('-')]
        added_lines = [c[1:] for c in changes if c.startswith('+')]

        # Gerar explicação simples
        explanation = get_simple_explanation(removed_lines, added_lines)

        html += f"""
        <div class="change-item">
            <div class="change-header">Alteração {i}</div>
            <div class="change-explanation">
                <p>{explanation}</p>
            </div>
            <div class="change-content">
                <div class="before-change">
                    <h4>Versão Original:</h4>
                    <div class="code-block">
        """

        # Processar linhas originais
        if len(removed_lines) == len(added_lines):
            for j, (removed, added) in enumerate(zip(removed_lines, added_lines)):
                removed_html, added_html = highlight_different_chars(removed, added)
                html += f"{removed_html}<br>" if j < len(removed_lines) - 1 else removed_html
        else:
            html += "<br>".join(removed_lines) if removed_lines else "[Sem conteúdo]"

        html += """
                    </div>
                </div>
                <div class="after-change">
                    <h4>Nova Versão:</h4>
                    <div class="code-block">
        """

        # Adicionar conteúdo para a nova versão
        if len(removed_lines) == len(added_lines):
            for j, (removed, added) in enumerate(zip(removed_lines, added_lines)):
                removed_html, added_html = highlight_different_chars(removed, added)
                html += f"{added_html}<br>" if j < len(added_lines) - 1 else added_html
        else:
            html += "<br>".join(added_lines) if added_lines else "[Linha removida]"

        html += """
                    </div>
                </div>
            </div>
        </div>
        """

    html += '</div>'

    return html


def highlight_different_chars(text1, text2):
    """
    Compara duas strings e retorna versões com os caracteres diferentes destacados em HTML.

    Args:
        text1: Primeira string a comparar
        text2: Segunda string a comparar

    Returns:
        Tupla com as duas strings com caracteres diferentes destacados em HTML
    """
    if not text1 and not text2:
        return "", ""

    if not text1:
        return "", f"<span class='highlight-add'>{text2}</span>"

    if not text2:
        return f"<span class='highlight-remove'>{text1}</span>", ""

    # Algoritmo de distância de edição modificado para rastrear diferenças
    import difflib

    # Usar o SequenceMatcher para identificar blocos similares
    matcher = difflib.SequenceMatcher(None, text1, text2)

    # Construir as versões destacadas
    text1_highlighted = ""
    text2_highlighted = ""

    for op, i1, i2, j1, j2 in matcher.get_opcodes():
        if op == 'equal':
            # Parte igual em ambas as strings
            text1_highlighted += text1[i1:i2]
            text2_highlighted += text2[j1:j2]
        elif op == 'replace':
            # Caracteres diferentes - destacar
            text1_highlighted += f"<span class='highlight-remove'>{text1[i1:i2]}</span>"
            text2_highlighted += f"<span class='highlight-add'>{text2[j1:j2]}</span>"
        elif op == 'delete':
            # Caracteres presentes apenas na primeira string
            text1_highlighted += f"<span class='highlight-remove'>{text1[i1:i2]}</span>"
        elif op == 'insert':
            # Caracteres presentes apenas na segunda string
            text2_highlighted += f"<span class='highlight-add'>{text2[j1:j2]}</span>"

    return text1_highlighted, text2_highlighted


def get_simple_explanation(before_lines, after_lines):
    """Gera uma explicação simples do que foi alterado."""
    if not before_lines and after_lines:
        return "Novo conteúdo foi adicionado ao arquivo."
    elif before_lines and not after_lines:
        return "Conteúdo foi removido do arquivo original."
    elif before_lines and after_lines:
        # Análise mais detalhada se for modificação
        if len(before_lines) == len(after_lines):
            return "O conteúdo foi modificado, mantendo a mesma quantidade de linhas."
        elif len(before_lines) > len(after_lines):
            return f"O conteúdo foi modificado, com {len(before_lines) - len(after_lines)} linha(s) removida(s)."
        else:
            return f"O conteúdo foi modificado, com {len(after_lines) - len(before_lines)} linha(s) adicionada(s)."
    return "Alteração detectada, mas tipo não identificado."


@app.route('/', methods=['GET'])
@login_required
def index():
    # Obter layouts disponíveis
    available_layouts = get_available_layouts()

    # Adicionar botão de logout no HTML
    logout_button = """
        <div style="position: fixed; top: 10px; right: 10px; z-index: 1000;">
            <a href="/logout" style="text-decoration: none;">
                <button style="padding: 8px 16px; background-color: #f44336; color: white; border: none; border-radius: 4px; cursor: pointer;">
                    <i class="fas fa-sign-out-alt" style="margin-right: 5px;"></i> Logout
                </button>
            </a>
        </div>
        """

    # Adicionar o botão de logout ao INDEX_HTML
    modified_index_html = INDEX_HTML.replace('<body>', '<body>' + logout_button)

    # Criar o seletor de layout para o modo individual
    layout_selector_individual = """
    <div class="form-group">
        <label for="layout" class="form-label">
            Layout (opcional)
            <span class="tooltip-icon">?
                <div class="tooltip-content">
                    Selecione um layout previamente importado para usar na comparação.
                </div>
            </span>
        </label>
        <select id="layout" name="layout_name" class="validation-method">
            <option value="">Autodetectar layout</option>
    """

    for layout_id, layout_name in available_layouts:
        layout_selector_individual += f'<option value="{layout_id}">{layout_name}</option>'

    layout_selector_individual += """
        </select>
    </div>
    """

    # Criar o seletor de layout para o modo lote
    layout_selector_batch = """
    <div class="form-group">
        <label for="batch-layout" class="form-label">
            Layout (opcional)
            <span class="tooltip-icon">?
                <div class="tooltip-content">
                    Selecione um layout previamente importado para usar na comparação em lote.
                </div>
            </span>
        </label>
        <select id="batch-layout" name="layout_name" class="validation-method">
            <option value="">Autodetectar layout</option>
    """

    for layout_id, layout_name in available_layouts:
        layout_selector_batch += f'<option value="{layout_id}">{layout_name}</option>'

    layout_selector_batch += """
        </select>
    </div>
    """

    # Inserir o seletor de layout no modo individual
    individual_insert_point = '<div class="form-group" style="margin-top: 25px;">'
    modified_index_html = modified_index_html.replace(
        individual_insert_point,
        layout_selector_individual + individual_insert_point
    )

    # Inserir o seletor de layout no modo lote
    batch_insert_point = '<div class="form-group">\n                        <label for="batch-method"'
    modified_index_html = modified_index_html.replace(
        batch_insert_point,
        layout_selector_batch + batch_insert_point
    )

    return modified_index_html


@app.route('/layouts', methods=['GET'])
@login_required
def manage_layouts():
    """
    Página para gerenciamento de layouts.
    """
    # Definir o diretório de layouts
    layouts_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'layouts')
    os.makedirs(layouts_dir, exist_ok=True)

    # Obter layouts disponíveis
    available_layouts = []

    try:
        for filename in os.listdir(layouts_dir):
            if filename.endswith('.json'):
                try:
                    layout_path = os.path.join(layouts_dir, filename)
                    with open(layout_path, 'r', encoding='utf-8') as file:
                        layout_data = json.load(file)
                        layout_name = os.path.splitext(filename)[0]

                        # Obter descrição e identificadores do layout
                        description = layout_data.get('description', 'Sem descrição')
                        identifier = layout_data.get('identifier', {})

                        available_layouts.append({
                            'name': layout_name,
                            'description': description,
                            'identifier': identifier
                        })
                except Exception as e:
                    print(f"Erro ao carregar layout {filename}: {str(e)}")
    except Exception as e:
        print(f"Erro ao listar layouts: {str(e)}")

    message = request.args.get('message', '')

    # Renderizar a página de gerenciamento de layouts
    return render_template_string(LAYOUTS_HTML_TEMPLATE, layouts=available_layouts, message=message)


def get_available_layouts():
    """Obter lista de layouts disponíveis"""
    layouts_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'layouts')
    try:
        layouts = []
        if os.path.exists(layouts_dir):
            for filename in os.listdir(layouts_dir):
                if filename.endswith('.json'):
                    layout_name = os.path.splitext(filename)[0]
                    # Torna o nome mais amigável para exibição
                    display_name = layout_name.replace('_', ' ').title()
                    layouts.append((layout_name, display_name))

        # ADICIONE esta linha para ordenar:
        layouts.sort(key=lambda x: x[1])  # Ordena pelo display_name

        return layouts
    except Exception as e:
        print(f"Erro ao obter layouts: {str(e)}")
        return []

@app.route('/upload-layout-automatico', methods=['POST'])
@login_required
def upload_layout_automatico():
    if 'layout_file' not in request.files:
        return redirect('/layouts?message=Nenhum arquivo selecionado')

    layout_file = request.files['layout_file']
    layout_name = request.form.get('layout_name', '')
    layout_description = request.form.get('layout_description', '')

    if layout_file.filename == '' or layout_name == '':
        return redirect('/layouts?message=Nome do layout e arquivo são obrigatórios')

    try:
        # Salvar o arquivo temporariamente
        temp_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'temp_layouts')
        os.makedirs(temp_dir, exist_ok=True)

        temp_file_path = os.path.join(temp_dir, secure_filename(layout_file.filename))
        layout_file.save(temp_file_path)

        # Extrair layout do arquivo
        layout_content = extrair_layout_de_documento(temp_file_path)

        # Remover arquivo temporário
        os.remove(temp_file_path)

        if layout_content is None:
            return redirect('/layouts?message=Não foi possível extrair layout do arquivo. Formato não reconhecido.')

        # Adicionar metadados
        layout_content['name'] = layout_name
        layout_content['description'] = layout_description

        # Salvar o layout extraído
        layouts_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'layouts')
        os.makedirs(layouts_dir, exist_ok=True)

        layout_path = os.path.join(layouts_dir, f"{layout_name.replace(' ', '_')}.json")
        with open(layout_path, 'w') as f:
            json.dump(layout_content, f, indent=2)

        return redirect('/layouts?message=Layout extraído e adicionado com sucesso')
    except Exception as e:
        return redirect(f'/layouts?message=Erro ao processar layout: {str(e)}')


@app.route('/edit-layout/<name>', methods=['GET', 'POST'])
@login_required
def edit_layout(name):
    layouts_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'layouts')
    layout_path = os.path.join(layouts_dir, f"{name}.json")

    if not os.path.exists(layout_path):
        return redirect('/layouts?message=Layout não encontrado')

    if request.method == 'GET':
        try:
            with open(layout_path, 'r') as f:
                layout = json.load(f)

            return render_template_string(EDIT_LAYOUT_HTML_TEMPLATE, layout=layout, layout_name=name)
        except Exception as e:
            return redirect(f'/layouts?message=Erro ao carregar layout: {str(e)}')
    else:  # POST
        try:
            # Obter os dados do formulário
            layout_data = {
                'name': request.form.get('layout_name', name),
                'description': request.form.get('layout_description', ''),
                'fields': {},
                'record_types': request.form.getlist('record_types'),
                'identifier': {
                    'record_types': request.form.getlist('record_types')
                }
            }

            # Processar os campos do layout
            field_names = request.form.getlist('field_name')
            field_types = request.form.getlist('field_type')
            field_positions = request.form.getlist('field_position')
            field_sizes = request.form.getlist('field_size')
            field_descriptions = request.form.getlist('field_description')

            for i in range(len(field_names)):
                if not field_names[i]:
                    continue

                field_info = {}

                # Tipo de registro
                if i < len(field_types) and field_types[i]:
                    field_info['tipo_registro'] = field_types[i]

                # Posição
                if i < len(field_positions) and field_positions[i]:
                    try:
                        if '-' in field_positions[i]:
                            inicio, fim = field_positions[i].split('-')
                            field_info['pos_inicial'] = int(inicio.strip())
                            field_info['pos_final'] = int(fim.strip())
                        else:
                            field_info['pos_inicial'] = int(field_positions[i])
                    except:
                        field_info['pos_inicial'] = 1

                # Tamanho
                if i < len(field_sizes) and field_sizes[i]:
                    try:
                        field_info['tamanho'] = int(field_sizes[i])
                    except:
                        if 'pos_final' in field_info and 'pos_inicial' in field_info:
                            field_info['tamanho'] = field_info['pos_final'] - field_info['pos_inicial'] + 1
                elif 'pos_final' in field_info and 'pos_inicial' in field_info:
                    field_info['tamanho'] = field_info['pos_final'] - field_info['pos_inicial'] + 1

                # Descrição
                if i < len(field_descriptions) and field_descriptions[i]:
                    field_info['descricao'] = field_descriptions[i]

                # Adicionar o campo ao layout
                layout_data['fields'][field_names[i]] = field_info

            # Salvar o layout atualizado
            with open(layout_path, 'w') as f:
                json.dump(layout_data, f, indent=2)

            return redirect(f'/layouts?message=Layout {name} atualizado com sucesso')
        except Exception as e:
            return redirect(f'/layouts?message=Erro ao atualizar layout: {str(e)}')


@app.route('/upload-and-compare', methods=['GET', 'POST'])
@login_required
def upload_and_compare():
    if request.method == 'POST':
        # Verificar se todos os arquivos foram enviados
        if 'layout_file' not in request.files or 'file1' not in request.files or 'file2' not in request.files:
            message = "Erro: Todos os arquivos (layout, arquivo 1 e arquivo 2) devem ser enviados"
            return redirect(f"/upload-and-compare?message={urllib.parse.quote(message)}")

        layout_file = request.files['layout_file']
        file1 = request.files['file1']
        file2 = request.files['file2']

        if layout_file.filename == '' or file1.filename == '' or file2.filename == '':
            message = "Erro: Todos os arquivos devem ser selecionados"
            return redirect(f"/upload-and-compare?message={urllib.parse.quote(message)}")

        method = request.form.get('method',
                                  'content')  # Usar 'content' como padrão, pois é melhor para arquivos estruturados

        try:
            # 1. Salvar os arquivos temporariamente
            timestamp = str(int(time.time()))
            temp_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'temp_' + timestamp)
            os.makedirs(temp_dir, exist_ok=True)

            layout_path = os.path.join(temp_dir, secure_filename(layout_file.filename))
            file1_path = os.path.join(temp_dir, secure_filename(file1.filename))
            file2_path = os.path.join(temp_dir, secure_filename(file2.filename))

            layout_file.save(layout_path)
            file1.save(file1_path)
            file2.save(file2_path)

            # 2. Extrair o layout
            layout_content = extrair_layout_de_documento(layout_path)

            if layout_content is None:
                return redirect(
                    f"/upload-and-compare?message={urllib.parse.quote('Não foi possível extrair layout do arquivo. Formato não reconhecido.')}")

            # 3. Salvar o layout temporariamente para uso na comparação
            layout_json_path = os.path.join(temp_dir, 'temp_layout.json')
            with open(layout_json_path, 'w') as f:
                json.dump(layout_content, f, indent=2)

            # 4. Comparar os arquivos usando o layout
            # Armazenar o layout na sessão para uso durante a comparação
            session['temp_layout'] = layout_content
            session['temp_layout_name'] = os.path.splitext(layout_file.filename)[0]

            # 5. Redirecionar para a visualização de resultados
            method = 'content'  # Forçar método 'content' para arquivos estruturados
            return render_template_string(UPLOAD_AND_COMPARE_RESULT_TEMPLATE,
                                          file1=file1.filename,
                                          file2=file2.filename,
                                          layout_name=layout_file.filename,
                                          temp_dir=temp_dir,
                                          timestamp=timestamp)

        except Exception as e:
            # Limpar arquivos temporários em caso de erro
            try:
                shutil.rmtree(temp_dir)
            except:
                pass

            message = f"Erro ao processar os arquivos: {str(e)}"
            return redirect(f"/upload-and-compare?message={urllib.parse.quote(message)}")

    # Método GET - exibir o formulário
    return render_template_string(UPLOAD_AND_COMPARE_TEMPLATE)


@app.route('/process-comparison/<timestamp>', methods=['GET'])
@login_required
def process_comparison(timestamp):
    temp_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'temp_' + timestamp)

    try:
        # Verificar se o diretório temporário existe
        if not os.path.exists(temp_dir):
            message = "Erro: Sessão de comparação expirada ou inválida"
            return redirect(f"/?message={urllib.parse.quote(message)}")

        # Obter os arquivos do diretório temporário
        files = [f for f in os.listdir(temp_dir) if os.path.isfile(os.path.join(temp_dir, f))]

        # Identificar o arquivo de layout (geralmente Excel, Word ou TXT)
        layout_file = next((f for f in files if f.endswith(('.xlsx', '.xls', '.docx', '.doc', '.txt'))
                            and "layout" in f.lower()), None)

        # Identificar os arquivos para comparação (excluindo o arquivo de layout)
        comparison_files = [f for f in files if f != layout_file and not f.startswith('.')]

        if len(comparison_files) < 2:
            message = "Erro: Número insuficiente de arquivos para comparação"
            return redirect(f"/?message={urllib.parse.quote(message)}")

        # Selecionar os primeiros dois arquivos para comparação
        file1 = comparison_files[0]
        file2 = comparison_files[1]

        # Obter caminhos completos dos arquivos
        file1_path = os.path.join(temp_dir, file1)
        file2_path = os.path.join(temp_dir, file2)

        # Se temos um arquivo de layout, extrair o layout dele
        layout = None
        layout_name = "Layout Detectado Automaticamente"

        if layout_file:
            layout_path = os.path.join(temp_dir, layout_file)
            layout_name = os.path.splitext(layout_file)[0]

            # Tentar extrair o layout
            try:
                layout = extrair_layout_de_documento(layout_path)
            except Exception as e:
                print(f"Erro ao extrair layout: {str(e)}")

        # Usar o layout da sessão se não conseguimos extrair
        if not layout:
            layout = session.get('temp_layout')
            layout_name = session.get('temp_layout_name', layout_name)

        # Configurar layout para a comparação
        session['layout_name'] = layout_name

        # Comparar os arquivos
        if layout:
            # Usar o layout detectado para comparação estruturada
            estrutura_diffs = analisar_arquivo_estruturado(file1_path, file2_path, layout)
            session['estrutura_diffs'] = estrutura_diffs

            # Converter diferenças estruturadas para formato tradicional
            differences = []
            for diff in estrutura_diffs:
                diff_type = diff.get('tipo', '')

                if diff_type == 'campos_alterados':
                    differences.append(f"@@ Linha {diff['linha']} - Alterações de campos @@")
                    for campo in diff.get('diferenca', []):
                        differences.append(f"- Campo {campo['campo']}: {campo['valor_antigo']}")
                        differences.append(f"+ Campo {campo['campo']}: {campo['valor_novo']}")

                elif diff_type == 'linha_alterada':
                    differences.append(f"@@ Linha {diff['linha']} @@")
                    differences.append(f"- {diff.get('valor_antigo', '')}")
                    differences.append(f"+ {diff.get('valor_novo', '')}")

                elif diff_type == 'linha_removida':
                    differences.append(f"@@ Linha {diff['linha']} @@")
                    differences.append(f"- {diff.get('valor_antigo', '')}")

                elif diff_type == 'linha_adicionada':
                    differences.append(f"@@ Linha {diff['linha']} @@")
                    differences.append(f"+ {diff.get('valor_novo', '')}")

            message = f"Análise de arquivo estruturado completa (Layout: {layout_name})."
        else:
            # Comparação básica sem layout específico
            differences, message = compare_files_content(file1_path, file2_path, context_lines=3)

        # Preparar os dados para o template
        error_message = ""
        differences_html = ""
        summary_message = ""
        summary_class = ""
        processed_diff_json = "[]"  # Valor padrão para JSON de diferenças processadas
        business_view_html = ""  # HTML para visualização amigável

        # Verificar resultado da comparação
        if "Erro" in message:
            error_message = f'<div class="error-message">{message}</div>'
        elif not differences:
            summary_message = "<h2>✓ Os arquivos são idênticos!</h2>"
            summary_class = "identical"
        else:
            summary_message = "<h2>✗ Os arquivos são diferentes!</h2>"
            summary_class = "different"

            # Gerar HTML para as diferenças
            differences_html = '<h3>Diferenças Encontradas:</h3><div class="diff-container">'
            for line in differences:
                if line.startswith('+') and not line.startswith('+++'):
                    differences_html += f'<pre class="diff-line diff-added">{line}</pre>'
                elif line.startswith('-') and not line.startswith('---'):
                    differences_html += f'<pre class="diff-line diff-removed">{line}</pre>'
                elif line.startswith('@@'):
                    differences_html += f'<div class="diff-header">{line}</div>'
                else:
                    differences_html += f'<pre class="diff-line diff-context">{line}</pre>'
            differences_html += '</div>'

            # Gerar a visualização amigável para negócios
            if 'estrutura_diffs' in session:
                business_view_html = gerar_html_visao_negocio_estruturada(session.get('estrutura_diffs', []))

            # Processar diferenças para visualizador avançado
            try:
                processed_diff = process_differences_for_advanced_view(file1_path, file2_path, differences)
                import json
                processed_diff_json = json.dumps(processed_diff)
            except Exception as e:
                print(f"Erro ao processar diferenças avançadas: {str(e)}")

        # Usar o mesmo template da função upload_files
        result_html = RESULT_HTML_TEMPLATE
        result_html = result_html.replace("{file1}", os.path.basename(file1_path))
        result_html = result_html.replace("{file2}", os.path.basename(file2_path))
        result_html = result_html.replace("{method}", "content")
        result_html = result_html.replace("{error_message}", error_message)
        result_html = result_html.replace("{summary_class}", summary_class)
        result_html = result_html.replace("{summary_message}", summary_message)
        result_html = result_html.replace("{differences_html}", differences_html)
        result_html = result_html.replace("{processed_diff_json}", processed_diff_json)
        result_html = result_html.replace("{business_view_html}", business_view_html)

        return result_html

    except Exception as e:
        # Limpar arquivos temporários em caso de erro
        try:
            shutil.rmtree(temp_dir)
        except:
            pass

        import traceback
        traceback.print_exc()
        message = f"Erro ao processar a comparação: {str(e)}"
        return redirect(f"/?message={urllib.parse.quote(message)}")


def compare_files_with_layout(file1_path, file2_path, layout):
    """
    Compara dois arquivos usando um layout temporário.
    """
    try:
        # Verificação rápida
        if filecmp.cmp(file1_path, file2_path, shallow=False):
            return [], "Os arquivos são idênticos."

        # Analisar usando o layout
        estrutura_diffs = analisar_arquivo_estruturado(file1_path, file2_path, layout)

        # Verificar resultado
        if estrutura_diffs is None:
            return super_compare_files(file1_path, file2_path, 3)

        # Armazenar para visualização
        if 'estrutura_diffs' not in session:
            session['estrutura_diffs'] = []
        session['estrutura_diffs'] = estrutura_diffs

        # Armazenar o nome do layout
        session['layout_name'] = layout.get('name', 'Layout Temporário')

        # Converter para o formato tradicional de diferenças
        traditional_diffs = []

        # Verificar se há erros
        has_error = any(diff.get('tipo') == 'erro' for diff in estrutura_diffs)
        if has_error:
            error_diff = next(diff for diff in estrutura_diffs if diff.get('tipo') == 'erro')
            return [], f"Erro na análise estruturada: {error_diff.get('mensagem', 'Erro desconhecido')}"

        # Converter diferenças estruturadas para o formato tradicional
        for diff in estrutura_diffs:
            diff_type = diff.get('tipo', '')

            if diff_type == 'campos_alterados':
                traditional_diffs.append(f"@@ Linha {diff['linha']} - Alterações de campos @@")
                for campo in diff.get('diferenca', []):
                    traditional_diffs.append(f"- Campo {campo['campo']}: {campo['valor_antigo']}")
                    traditional_diffs.append(f"+ Campo {campo['campo']}: {campo['valor_novo']}")

            elif diff_type == 'linha_alterada':
                traditional_diffs.append(f"@@ Linha {diff['linha']} @@")
                traditional_diffs.append(f"- {diff.get('valor_antigo', '')}")
                traditional_diffs.append(f"+ {diff.get('valor_novo', '')}")

            elif diff_type == 'linha_removida':
                traditional_diffs.append(f"@@ Linha {diff['linha']} @@")
                traditional_diffs.append(f"- {diff.get('valor_antigo', '')}")

            elif diff_type == 'linha_adicionada':
                traditional_diffs.append(f"@@ Linha {diff['linha']} @@")
                traditional_diffs.append(f"+ {diff.get('valor_novo', '')}")

            elif diff_type == 'info':
                traditional_diffs.append(f"@@ {diff.get('mensagem', '')} @@")

        return traditional_diffs, f"Análise de arquivo estruturado completa (Layout: {layout.get('name', 'Temporário')})."

    except Exception as e:
        import traceback
        traceback.print_exc()
        print(f"Erro ao processar arquivo estruturado: {str(e)}")
        return super_compare_files(file1_path, file2_path, 3)


@app.route('/batch-upload', methods=['POST'])
@login_required
def batch_upload_files():
    import json

    if 'source1' not in request.files or 'source2' not in request.files:
        message = "Erro: Ambos os arquivos ZIP devem ser enviados"
        return redirect(f"/?message={urllib.parse.quote(message)}")

    source1 = request.files['source1']
    source2 = request.files['source2']

    if source1.filename == '' or source2.filename == '':
        message = "Erro: Nenhum arquivo ZIP selecionado"
        return redirect(f"/?message={urllib.parse.quote(message)}")

    # Verificar se os arquivos são ZIPs
    if not source1.filename.lower().endswith('.zip') or not source2.filename.lower().endswith('.zip'):
        message = "Erro: Os arquivos devem estar no formato ZIP"
        return redirect(f"/?message={urllib.parse.quote(message)}")

    method = request.form.get('method', 'hash')

    # CORREÇÃO: Verificar se um layout foi selecionado
    layout_name = request.form.get('layout_name', '')
    layout = None

    if layout_name:
        layouts_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'layouts')
        layout_path = os.path.join(layouts_dir, f"{layout_name}.json")
        if os.path.exists(layout_path):
            try:
                # CORREÇÃO: Usar json.load() diretamente
                with open(layout_path, 'r', encoding='utf-8') as f:
                    layout = json.load(f)
                session['layout_name'] = layout_name.replace('_', ' ')
                session['selected_layout'] = layout  # NOVO: Armazenar layout na sessão
                print(f"Layout '{layout_name}' carregado para comparação em lote")
            except Exception as e:
                print(f"Erro ao carregar layout para lote: {str(e)}")

    # Criar diretórios temporários para extração
    timestamp = str(int(time.time()))
    temp_dir = os.path.join(app.config['UPLOAD_FOLDER'], timestamp)
    source1_dir = os.path.join(temp_dir, 'source1')
    source2_dir = os.path.join(temp_dir, 'source2')

    os.makedirs(source1_dir, exist_ok=True)
    os.makedirs(source2_dir, exist_ok=True)

    # Salvar os arquivos ZIP em blocos para lidar com arquivos grandes
    source1_path = os.path.join(temp_dir, secure_filename(source1.filename))
    source2_path = os.path.join(temp_dir, secure_filename(source2.filename))

    chunk_size = 8192  # 8KB chunks

    # Salvar source1 em chunks
    with open(source1_path, 'wb') as f:
        while True:
            chunk = source1.read(chunk_size)
            if not chunk:
                break
            f.write(chunk)

    # Salvar source2 em chunks
    with open(source2_path, 'wb') as f:
        while True:
            chunk = source2.read(chunk_size)
            if not chunk:
                break
            f.write(chunk)

    # Obter filtros, se houver
    filters = {}
    filter_fields = [
        'file_extensions', 'size_min', 'size_max', 'modified_after',
        'modified_before', 'name_pattern', 'exclude_pattern'
    ]

    for field in filter_fields:
        value = request.form.get(field)
        if value:
            filters[field] = value

    try:
        # Extrair os arquivos ZIP
        extract_zip(source1_path, source1_dir)
        extract_zip(source2_path, source2_dir)

        # CORREÇÃO: Comparar os diretórios, usando o layout se disponível e método for content
        if layout and method == 'content':
            print(f"Usando layout selecionado para comparação em lote: {layout.get('name', 'Sem nome')}")
            # Comparação em lote usando layout estruturado
            comparison_results = []

            # Obter lista de arquivos em cada diretório
            files1 = [f for f in os.listdir(source1_dir) if os.path.isfile(os.path.join(source1_dir, f))]
            files2 = [f for f in os.listdir(source2_dir) if os.path.isfile(os.path.join(source2_dir, f))]

            # Aplicar filtros se fornecidos
            if filters:
                files1 = apply_filters(files1, source1_dir, filters)
                files2 = apply_filters(files2, source2_dir, filters)

            # Encontrar arquivos comuns
            common_files = set(files1).intersection(set(files2))
            only_in_source1 = set(files1) - set(files2)
            only_in_source2 = set(files2) - set(files1)

            # Comparar arquivos comuns usando o layout
            for filename in common_files:
                file1_path = os.path.join(source1_dir, filename)
                file2_path = os.path.join(source2_dir, filename)

                result = {
                    'filename': filename,
                    'exists_in_source1': True,
                    'exists_in_source2': True,
                    'comparison_result': None,
                    'details': None,
                    'file_info': {
                        'size1': os.path.getsize(file1_path),
                        'size2': os.path.getsize(file2_path),
                        'modified1': datetime.datetime.fromtimestamp(os.path.getmtime(file1_path)).strftime(
                            '%Y-%m-%d %H:%M:%S'),
                        'modified2': datetime.datetime.fromtimestamp(os.path.getmtime(file2_path)).strftime(
                            '%Y-%m-%d %H:%M:%S')
                    }
                }

                # Usar o layout estruturado para comparação
                estrutura_diffs = analisar_arquivo_estruturado(file1_path, file2_path, layout)

                if estrutura_diffs is None or len(estrutura_diffs) == 0:
                    result['comparison_result'] = 'identical'
                    result['details'] = 'Os arquivos são idênticos'
                else:
                    result['comparison_result'] = 'different'

                    # Converter para formato tradicional
                    differences = []
                    for diff in estrutura_diffs:
                        diff_type = diff.get('tipo', '')

                        if diff_type == 'campos_alterados':
                            differences.append(f"@@ Linha {diff['linha']} - Alterações de campos @@")
                            for campo in diff.get('diferenca', []):
                                differences.append(f"- Campo {campo['campo']}: {campo['valor_antigo']}")
                                differences.append(f"+ Campo {campo['campo']}: {campo['valor_novo']}")

                        elif diff_type == 'linha_alterada':
                            differences.append(f"@@ Linha {diff['linha']} @@")
                            differences.append(f"- {diff.get('valor_antigo', '')}")
                            differences.append(f"+ {diff.get('valor_novo', '')}")

                        elif diff_type == 'linha_removida':
                            differences.append(f"@@ Linha {diff['linha']} @@")
                            differences.append(f"- {diff.get('valor_antigo', '')}")

                        elif diff_type == 'linha_adicionada':
                            differences.append(f"@@ Linha {diff['linha']} @@")
                            differences.append(f"+ {diff.get('valor_novo', '')}")

                    result['details'] = differences

                    # Processar diferenças para visualização avançada
                    try:
                        result['processed_diff'] = process_differences_for_advanced_view(file1_path, file2_path,
                                                                                         differences)
                    except Exception as process_error:
                        print(f"Erro ao processar diferenças avançadas: {str(process_error)}")
                        # Ignorar erros no processamento avançado

                comparison_results.append(result)

            # Adicionar os arquivos que existem apenas em uma fonte
            for filename in only_in_source1:
                file_path = os.path.join(source1_dir, filename)
                comparison_results.append({
                    'filename': filename,
                    'exists_in_source1': True,
                    'exists_in_source2': False,
                    'comparison_result': 'only_in_source1',
                    'details': 'Arquivo existe apenas na Fonte 1',
                    'file_info': {
                        'size1': os.path.getsize(file_path),
                        'size2': 0,
                        'modified1': datetime.datetime.fromtimestamp(os.path.getmtime(file_path)).strftime(
                            '%Y-%m-%d %H:%M:%S'),
                        'modified2': 'N/A'
                    }
                })

            for filename in only_in_source2:
                file_path = os.path.join(source2_dir, filename)
                comparison_results.append({
                    'filename': filename,
                    'exists_in_source1': False,
                    'exists_in_source2': True,
                    'comparison_result': 'only_in_source2',
                    'details': 'Arquivo existe apenas na Fonte 2',
                    'file_info': {
                        'size1': 0,
                        'size2': os.path.getsize(file_path),
                        'modified1': 'N/A',
                        'modified2': datetime.datetime.fromtimestamp(os.path.getmtime(file_path)).strftime(
                            '%Y-%m-%d %H:%M:%S')
                    }
                })
        else:
            # Usar a função de comparação em lote existente
            comparison_results = compare_file_batches_with_filters(source1_dir, source2_dir, method, filters)

        # Preparar estatísticas
        total_files = len(comparison_results)
        identical_files = sum(1 for r in comparison_results if r['comparison_result'] == 'identical')
        different_files = sum(1 for r in comparison_results if r['comparison_result'] == 'different')
        only_in_source1 = sum(1 for r in comparison_results if r['comparison_result'] == 'only_in_source1')
        only_in_source2 = sum(1 for r in comparison_results if r['comparison_result'] == 'only_in_source2')

        # Resto do código permanece igual...
        # (código de geração de HTML, tabelas, etc.)

        # Gerar as linhas da tabela técnica
        table_rows = ""
        # Gerar as linhas para visualização de negócios
        business_file_rows = ""
        file_details = {}

        for i, result in enumerate(comparison_results):
            filename = result['filename']
            status_class = ""
            status_text = ""
            business_status = ""

            if result['comparison_result'] == 'identical':
                status_class = "status-identical"
                status_text = "identical"
                business_status = "Idêntico"
            elif result['comparison_result'] == 'different':
                status_class = "status-different"
                status_text = "different"
                business_status = "Diferente"
            elif result['comparison_result'] == 'only_in_source1':
                status_class = "status-missing"
                status_text = "missing"
                business_status = "Somente na Fonte 1"
            elif result['comparison_result'] == 'only_in_source2':
                status_class = "status-missing"
                status_text = "missing"
                business_status = "Somente na Fonte 2"

            # Obter as informações de file_info do objeto result
            fileInfo = result.get('file_info', {})

            # Linha para visão técnica
            table_rows += f"""
            <tr>
                <td>{filename}</td>
                <td><span class="{status_class}">{status_text}</span></td>
                <td>{fileInfo.get('modified1', 'N/A')}</td>
                <td>{fileInfo.get('modified2', 'N/A')}</td>
                <td>{formatSize(fileInfo.get('size1', 0))}</td>
                <td>{formatSize(fileInfo.get('size2', 0))}</td>
                <td><button class="details-btn" onclick="showDetails({i})">Ver Detalhes</button></td>
            </tr>
            """

            # Linha para visão de negócios (apenas arquivos diferentes e ausentes)
            if result['comparison_result'] != 'identical':
                business_file_rows += f"""
                <div class="business-file-row">
                    <div class="business-file-name">{filename}</div>
                    <div class="business-file-status">{business_status}</div>
                    <div class="business-file-action">
                        <button class="details-btn" onclick="showDetails({i})">Ver Detalhes</button>
                    </div>
                </div>
                """

            # Adicionar detalhes para uso no JavaScript
            file_details[i] = {
                'filename': filename,
                'comparison_result': result['comparison_result'],
                'exists_in_source1': result['exists_in_source1'],
                'exists_in_source2': result['exists_in_source2'],
                'details': result['details'],
                'method': method,
                'file_info': fileInfo
            }

            # Adicionar dados processados para visualização avançada de diferenças, se disponível
            if result['comparison_result'] == 'different' and method == 'content' and 'processed_diff' in result:
                file_details[i]['processed_diff'] = result['processed_diff']

        # Renderizar o resultado
        import json
        file_details_json = json.dumps(file_details)

        # Armazenar os resultados na sessão para exportação posterior
        session[f'comparison_results_{timestamp}'] = comparison_results

        # Adicionar informações sobre exportação
        export_options = f"""
        <a href="/export-csv/{timestamp}" class="export-btn">Exportar CSV</a>
        <a href="/export-excel/{timestamp}" class="export-btn">Exportar Excel</a>
        """

        # Informações sobre os filtros aplicados, se houver
        filter_info = ""
        if filters:
            filter_info = '<div class="filter-info"><h4>Filtros Aplicados:</h4><ul>'

            if 'file_extensions' in filters:
                filter_info += f'<li>Extensões de arquivo: {filters["file_extensions"]}</li>'

            if 'size_min' in filters:
                filter_info += f'<li>Tamanho mínimo: {filters["size_min"]} KB</li>'

            if 'size_max' in filters:
                filter_info += f'<li>Tamanho máximo: {filters["size_max"]} KB</li>'

            if 'modified_after' in filters:
                filter_info += f'<li>Modificado após: {filters["modified_after"]}</li>'

            if 'modified_before' in filters:
                filter_info += f'<li>Modificado antes: {filters["modified_before"]}</li>'

            if 'name_pattern' in filters:
                filter_info += f'<li>Padrão de nome: {filters["name_pattern"]}</li>'

            if 'exclude_pattern' in filters:
                filter_info += f'<li>Padrão de exclusão: {filters["exclude_pattern"]}</li>'

            # Adicionar informação do layout utilizado, se aplicável
            if layout_name:
                filter_info += f'<li>Layout utilizado: {layout_name}</li>'

            filter_info += '</ul></div>'

        # Se não houver arquivos diferentes, adicionar mensagem
        if different_files == 0 and only_in_source1 == 0 and only_in_source2 == 0:
            business_file_rows = """
            <div class="business-file-row" style="justify-content: center; padding: 30px;">
                <div style="text-align: center; font-style: italic; color: #555;">
                    Não foram encontradas diferenças entre os arquivos.
                    Todos os arquivos são idênticos nas duas fontes.
                </div>
            </div>
            """
        elif different_files == 0 and (only_in_source1 > 0 or only_in_source2 > 0):
            # Há arquivos ausentes, mas não há diferenças
            business_file_rows = '<div class="difference-summary" style="margin-bottom: 15px;">' + \
                                 '<p>Não foram encontradas diferenças no conteúdo dos arquivos comuns.</p>' + \
                                 '<p>Porém existem arquivos que estão presentes em apenas uma das fontes.</p>' + \
                                 '</div>' + business_file_rows

        result_html = BATCH_RESULT_HTML_TEMPLATE
        result_html = result_html.replace("{export_options}", export_options)
        result_html = result_html.replace("{filter_info}", filter_info)
        result_html = result_html.replace("{source1}", os.path.basename(source1.filename))
        result_html = result_html.replace("{source2}", os.path.basename(source2.filename))
        result_html = result_html.replace("{method}", method)
        result_html = result_html.replace("{total_files}", str(total_files))
        result_html = result_html.replace("{identical_files}", str(identical_files))
        result_html = result_html.replace("{different_files}", str(different_files))
        result_html = result_html.replace("{only_in_source1}", str(only_in_source1))
        result_html = result_html.replace("{only_in_source2}", str(only_in_source2))
        result_html = result_html.replace("{table_rows}", table_rows)
        result_html = result_html.replace("{business_file_rows}", business_file_rows)
        result_html = result_html.replace("{file_details_json}", file_details_json)

        return result_html

    except Exception as e:
        message = f"Erro ao processar os arquivos: {str(e)}"
        import traceback
        traceback.print_exc()
        return redirect(f"/?message={urllib.parse.quote(message)}")
    finally:
        # Limpeza dos arquivos temporários
        try:
            shutil.rmtree(temp_dir)
        except Exception as e:
            # Apenas ignorar erros de limpeza
            print(f"Erro ao limpar arquivos temporários: {str(e)}")


def formatSize(size_bytes):
    """Formata o tamanho em bytes para uma representação mais legível."""
    if size_bytes == 0:
        return "N/A"

    size_names = ["B", "KB", "MB", "GB", "TB"]
    i = 0
    while size_bytes >= 1024 and i < len(size_names) - 1:
        size_bytes /= 1024
        i += 1

    return f"{size_bytes:.2f} {size_names[i]}"


if __name__ == '__main__':
    app.run(debug=True, port=5000)  # Rotas da aplicação web
